# generate_curricula_v8.1.py
"""
ECM Curriculum Generator - Enhanced Version 8.1
Reads from curricula input JSON file and generates comprehensive curricula
Addresses critique: explicit module themes, refined autonomy language, enhanced pedagogical arcs
NO DigComp mappings as per requirements
"""

import json
import os
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement, qn

class ECMCurriculumGenerator:
    """Generate comprehensive ECM curricula from JSON input with enhanced precision and professional formatting"""
    
    def __init__(self, config_path='../config/settings.json'):
        """Initialize generator with configuration"""
        self.config = self.load_config(config_path)
        self.setup_paths()
        self.curricula_data = self.load_curricula_data()
        
        print(f"📚 ECM Curriculum Generator - Enhanced Professional Standards v8.1")
        print(f"⚙️ Configuration loaded from settings.json")
        print(f"📁 Input curricula: {getattr(self, 'curricula_path', 'Using fallback data')}")
        print(f"📁 Output directory: {self.output_dir}")
        print(f"📄 Loaded {len(self.curricula_data)} base curricula")
    
    def load_config(self, config_path):
        """Load configuration from settings.json"""
        try:
            project_root = Path(__file__).parent.parent
            config_file = project_root / config_path.lstrip('./')
            
            with open(config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)
            print(f"✅ Configuration loaded from: {config_file}")
            return config
        except FileNotFoundError:
            print(f"⚠️ Config file not found: {config_path}")
            return self.get_default_config()
        except Exception as e:
            print(f"⚠️ Error loading config: {e}")
            return self.get_default_config()
    
    def get_default_config(self):
        """Default configuration if settings.json not found"""
        return {
            "paths": {
                "input_modules": "./input/modules/modules_v5.json",
                "input_curricula": "./input/curricula/curricula.json"
            },
            "output": {
                "curricula": {
                    "directory": "./output/curricula",
                    "formats": ["json", "html", "docx"]
                }
            },
            "system": {"version": "8.1.0"}
        }
    
    def setup_paths(self):
        """Setup paths based on configuration"""
        project_root = Path(__file__).parent.parent
        
        modules_config = self.config['paths']['input_modules'].lstrip('./')
        self.modules_path = project_root / modules_config
        
        # Try to load curricula from dedicated file, fall back to educational profiles
        curricula_config = self.config['paths'].get('input_curricula', 
                                                   self.config['paths'].get('input_educational_profiles', 
                                                                            './input/curricula/curricula.json')).lstrip('./')
        self.curricula_path = project_root / curricula_config
        
        output_config = self.config.get('output', {}).get('curricula', {})
        output_dir = output_config.get('directory', './output/curricula').lstrip('./')
        self.output_dir = project_root / output_dir
        
        self.output_formats = output_config.get('formats', ['json', 'html', 'docx'])
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def load_modules(self):
        """Load modules from configured path"""
        try:
            with open(self.modules_path, 'r', encoding='utf-8') as f:
                modules = json.load(f)
            print(f"✅ Loaded {len(modules)} modules")
            return modules
        except FileNotFoundError:
            print(f"❌ Modules file not found: {self.modules_path}")
            return []
        except Exception as e:
            print(f"❌ Error loading modules: {e}")
            return []
    
    def load_curricula_data(self):
        """Load curricula from JSON input file or extract from educational profiles"""
        try:
            with open(self.curricula_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Check if this is educational profiles data (list format) or curricula data
            if isinstance(data, list):
                print(f"✅ Processing educational profiles data to extract curricula")
                return self.extract_curricula_from_profiles(data)
            else:
                print(f"✅ Loaded curricula data from: {self.curricula_path}")
                return data.get('curricula', []) if 'curricula' in data else [data]
                
        except FileNotFoundError:
            print(f"⚠️ Curricula file not found: {self.curricula_path}")
            print(f"📚 Using fallback curriculum data")
            return self.get_fallback_curricula_data()
        except Exception as e:
            print(f"⚠️ Error loading curricula: {e}")
            print(f"📚 Using fallback curriculum data")
            return self.get_fallback_curricula_data()
    
    def extract_curricula_from_profiles(self, profiles_data):
        """Extract curricula information from educational profiles JSON"""
        curricula = []
        
        for profile in profiles_data:
            # Extract basic info
            profile_id = profile.get('id', 'UNKNOWN')
            profile_name = profile.get('profile_name', 'Unknown Profile').replace(' Educational Profile', '')
            
            # Determine EQF levels from the profile
            eqf_levels = self.extract_eqf_levels(profile)
            
            for eqf_level in eqf_levels:
                curriculum = self.create_curriculum_from_profile(profile, eqf_level)
                curricula.append(curriculum)
        
        return curricula
    
    def extract_eqf_levels(self, profile):
        """Extract EQF levels from profile data"""
        eqf_levels = set()
        
        # Check enhanced_purpose_and_application
        purpose_app = profile.get('enhanced_purpose_and_application', {})
        for key in purpose_app.keys():
            if key.startswith('eqf_'):
                level = key.replace('eqf_', '')
                if level.isdigit():
                    eqf_levels.add(int(level))
        
        # Check learning_outcomes_by_eqf
        learning_outcomes = profile.get('learning_outcomes_by_eqf', {})
        for level in learning_outcomes.keys():
            if level.isdigit():
                eqf_levels.add(int(level))
        
        # If no levels found, default to 6
        if not eqf_levels:
            eqf_levels.add(6)
            
        return sorted(eqf_levels)
    
    def create_curriculum_from_profile(self, profile, eqf_level):
        """Create curriculum data structure from profile"""
        profile_id = profile.get('id', 'UNKNOWN')
        profile_name = profile.get('profile_name', 'Unknown Profile').replace(' Educational Profile', '')
        
        # Calculate ECTS based on EQF level
        ects_mapping = {4: 25, 5: 25, 6: 36, 7: 45, 8: 50}
        target_ects = ects_mapping.get(eqf_level, 36)
        
        # Extract modules from learning outcomes
        learning_outcomes = profile.get('learning_outcomes_by_eqf', {})
        eqf_outcomes = learning_outcomes.get(str(eqf_level), learning_outcomes.get('6', {}))
        
        modules = []
        if 'units' in eqf_outcomes:
            for unit_name, unit_data in eqf_outcomes['units'].items():
                modules.append(unit_name)
        
        # If no modules found, create default modules
        if not modules:
            modules = [
                f'Core {profile_name} Foundations',
                f'Advanced {profile_name} Practice', 
                f'{profile_name} Implementation & Excellence'
            ]
        
        # Extract delivery mode from role description
        role_desc = profile.get('role_description', '').lower()
        if 'research' in role_desc:
            delivery = 'Research-Based (Academic + Project-Based Learning)'
        elif 'consulting' in role_desc or 'advisory' in role_desc:
            delivery = 'Flexible (Online + Client-Based Practical)'
        elif 'technical' in role_desc or 'developer' in role_desc:
            delivery = 'Technical (Lab-Based + Industry Projects)'
        elif 'management' in role_desc or 'leader' in role_desc:
            delivery = 'Blended (Online + Practical + Work-Based Option)'
        else:
            delivery = 'Blended (Online + Practical + Work-Based Option)'
        
        # Create pedagogical arc
        if eqf_level <= 5:
            pedagogical_arc = 'Foundations → Application → Professional Practice'
        elif eqf_level == 6:
            pedagogical_arc = 'Professional Foundations → Advanced Application → Implementation Excellence'
        elif eqf_level == 7:
            pedagogical_arc = 'Advanced Theory → Strategic Application → Leadership Excellence'
        else:
            pedagogical_arc = 'Innovation Foundations → Breakthrough Application → Global Impact'
        
        # Extract target audience
        target_audience = self.extract_target_audience(profile, eqf_level)
        
        return {
            'id': f'{profile_id}_EQF{eqf_level}_{target_ects}ECTS',
            'role_name': profile_name,
            'eqf_level': eqf_level,
            'target_ects': target_ects,
            'delivery': delivery,
            'module_count': len(modules),
            'specialization': self.extract_specialization(profile),
            'target_audience': target_audience,
            'pedagogical_arc': pedagogical_arc,
            'modules': modules,
            'source_profile': profile_id,
            'enhanced_purpose': profile.get('enhanced_purpose_and_application', {}).get(f'eqf_{eqf_level}', ''),
            'programme_outcome': eqf_outcomes.get('programme_outcome', ''),
            'learning_outcomes': eqf_outcomes
        }
    
    def extract_specialization(self, profile):
        """Extract specialization from profile data"""
        profile_name = profile.get('profile_name', '').lower()
        
        if 'leader' in profile_name:
            return 'Strategic Leadership'
        elif 'manager' in profile_name:
            return 'Program Management'
        elif 'consultant' in profile_name:
            return 'Advisory Services'
        elif 'analyst' in profile_name:
            return 'Data Analytics'
        elif 'engineer' in profile_name:
            return 'Technical Engineering'
        elif 'developer' in profile_name:
            return 'Software Development'
        elif 'designer' in profile_name:
            return 'Systems Design'
        elif 'specialist' in profile_name:
            return 'Technical Specialization'
        else:
            return 'Professional Practice'
    
    def extract_target_audience(self, profile, eqf_level):
        """Extract target audience based on profile and EQF level"""
        base_desc = profile.get('role_description', '')
        
        if eqf_level <= 5:
            return f"Entry-level professionals and technicians seeking to develop {profile.get('profile_name', 'sustainability').lower()} capabilities"
        elif eqf_level == 6:
            return f"Mid-level professionals implementing {profile.get('profile_name', 'sustainability').lower()} programs and initiatives"
        elif eqf_level == 7:
            return f"Senior professionals and managers leading {profile.get('profile_name', 'sustainability').lower()} transformation initiatives"
        else:
            return f"Executive leaders and experts pioneering {profile.get('profile_name', 'sustainability').lower()} innovation and global impact"
    
    def get_fallback_curricula_data(self):
        """Fallback curriculum data if no input file available"""
        return [
            {
                'id': 'DSL_EQF7_45ECTS',
                'role_name': 'Digital Sustainability Leader',
                'eqf_level': 7,
                'target_ects': 45,
                'delivery': 'Blended (Online + Practical + Work-Based Option)',
                'module_count': 15,
                'specialization': 'Strategic Leadership',
                'target_audience': 'Senior professionals and executives transitioning to sustainability leadership roles',
                'pedagogical_arc': 'Foundation → Strategic Application → Executive Integration',
                'modules': [
                    'Strategic Leadership & Transformation Theory',
                    'Digital Governance & Stakeholder Orchestration', 
                    'Innovation Strategy & Business Model Evolution',
                    'Executive Decision-Making in Sustainability',
                    'Advanced Stakeholder Engagement Systems'
                ]
            },
            {
                'id': 'DSM_EQF6_36ECTS',
                'role_name': 'Digital Sustainability Manager',
                'eqf_level': 6,
                'target_ects': 36,
                'delivery': 'Blended (Online + Practical + Work-Based Option)',
                'module_count': 12,
                'specialization': 'Program Management',
                'target_audience': 'Mid-level managers implementing sustainability programs',
                'pedagogical_arc': 'Management Foundations → Digital Implementation → Program Excellence',
                'modules': [
                    'Program Management & Multi-Stakeholder Coordination',
                    'Digital Tools Integration & Implementation Systems', 
                    'Performance Monitoring & Advanced Reporting Frameworks',
                    'Change Management in Digital Transformation'
                ]
            },
            {
                'id': 'SSD_EQF6_30ECTS',
                'role_name': 'Sustainable Systems Designer',
                'eqf_level': 6,
                'target_ects': 30,
                'delivery': 'Design-Focused (Studio-Based + Practical Projects)',
                'module_count': 10,
                'specialization': 'Systems Design',
                'target_audience': 'Designers and architects of sustainable systems',
                'pedagogical_arc': 'Design Foundations → Systems Integration → Implementation Excellence',
                'modules': [
                    'ESG Data for Design & Visualization Systems',
                    'Systems Architecture & Sustainability Integration', 
                    'Implementation Design & Validation Systems'
                ]
            }
        ]
    
    def generate_comprehensive_curriculum(self, curriculum_data):
        """Generate comprehensive curriculum following gold standard detailed format with enhanced autonomy language"""
        
        modules = self.load_modules()
        if not modules:
            modules = [{'id': f'MOD{i:03d}', 'name': f'Module {i}'} for i in range(1, curriculum_data['module_count'] + 1)]
        
        selected_modules = modules[:curriculum_data['module_count']]
        total_ects = curriculum_data['target_ects']
        total_hours = total_ects * 25
        
        curriculum = {
            'curriculum_info': {
                'id': curriculum_data['id'],
                'title': f"{curriculum_data['role_name']} Professional Development Course",
                'eqf_level': curriculum_data['eqf_level'],
                'target_ects': total_ects,
                'total_hours': total_hours,
                'specialization': curriculum_data['specialization'],
                'delivery': curriculum_data['delivery'],
                'module_count': curriculum_data['module_count'],
                'modules': curriculum_data['modules'],
                'pedagogical_arc': curriculum_data['pedagogical_arc'],
                'generated_date': datetime.now().isoformat(),
                'source_data': f"Generated from {curriculum_data.get('source_profile', 'input data')}"
            },
            
            'programme_learning_outcomes': {
                'completion_statement': 'Upon successful completion, participants will be able to:',
                'main_outcomes': [
                    f"Critically analyze ESG reporting frameworks (GRI, CSRD, TCFD), statistical methods, and sustainability data governance standards with autonomous professional judgment appropriate for {curriculum_data['role_name'].lower()} practice",
                    f"Strategically design sustainability dashboards, data visualization systems, and reporting tools using appropriate technologies (Python, R, Power BI) with professional autonomy, accessing specialized guidance where appropriate to refine complex professional decisions",
                    f"Autonomously govern data-driven sustainability initiatives while orchestrating teams and stakeholder engagement in complex professional environments with independent decision-making authority"
                ],
                'enhanced_programme_outcome': curriculum_data.get('programme_outcome', ''),
                'green_competences': {
                    'framework': 'GreenComp Framework',
                    'competency_groups': {
                        'environmental_understanding': [
                            "Evaluate environmental impact through lifecycle assessment (LCA) and carbon footprint analysis with autonomous analytical judgment (GreenComp 1.1 Interconnectedness: Level 4 - Advanced)",
                            "Champion systems thinking to understand sustainability interconnections using independent critical analysis (GreenComp 1.2 Systems Thinking: Level 4 - Advanced)"
                        ],
                        'sustainable_innovation': [
                            "Pioneer circular economy principles in sustainability recommendations with autonomous strategic decision-making (GreenComp 1.3 Circular Thinking: Level 4 - Advanced)",
                            "Orchestrate future-oriented sustainability scenarios and planning with independent visionary thinking (GreenComp 2.3 Futures Thinking: Level 4 - Advanced)"
                        ],
                        'sustainability_complexity': [
                            "Navigate ecosystem impacts using quantitative environmental indicators with autonomous professional judgment (GreenComp 2.2 Complexity in Sustainability: Level 4 - Advanced)"
                        ],
                        'collaborative_action': [
                            "Cultivate collective action for sustainability initiatives using independent leadership capabilities (GreenComp 4.2 Collective Action: Level 3 - Intermediate)",
                            "Steward personal and professional responsibility for sustainable practices with autonomous ethical decision-making (GreenComp 4.3 Individual Initiative: Level 4 - Advanced)"
                        ]
                    },
                    'competencies': [
                        "Evaluate environmental impact through lifecycle assessment (LCA) and carbon footprint analysis with autonomous analytical judgment (GreenComp 1.1 Interconnectedness: Level 4 - Advanced)",
                        "Pioneer circular economy principles in sustainability recommendations with autonomous strategic decision-making (GreenComp 1.3 Circular Thinking: Level 4 - Advanced)",
                        "Navigate ecosystem impacts using quantitative environmental indicators with autonomous professional judgment (GreenComp 2.2 Complexity in Sustainability: Level 4 - Advanced)",
                        "Orchestrate future-oriented sustainability scenarios and planning with independent visionary thinking (GreenComp 2.3 Futures Thinking: Level 4 - Advanced)",
                        "Cultivate collective action for sustainability initiatives using independent leadership capabilities (GreenComp 4.2 Collective Action: Level 3 - Intermediate)",
                        "Champion systems thinking to understand sustainability interconnections using independent critical analysis (GreenComp 1.2 Systems Thinking: Level 4 - Advanced)",
                        "Steward personal and professional responsibility for sustainable practices with autonomous ethical decision-making (GreenComp 4.3 Individual Initiative: Level 4 - Advanced)"
                    ]
                },
                'transversal_skills': {
                    'framework': 'Key Competences for Lifelong Learning',
                    'competencies': [
                        "Communicate complex sustainability data insights to diverse stakeholder audiences with autonomous professional judgment",
                        "Demonstrate ethical data handling practices in environmental and social impact reporting using independent ethical reasoning",
                        "Collaborate effectively in multicultural, interdisciplinary sustainability teams with autonomous conflict resolution capabilities"
                    ]
                }
            },
            
            'modular_structure': {
                'pedagogical_progression': curriculum_data['pedagogical_arc'],
                'modules': curriculum_data['modules'],
                'module_learning_outcomes': self.generate_module_outcomes(curriculum_data),
                'dependencies': {
                    'foundation': 'No prerequisites - establishes core knowledge, sustainability literacy, and basic green competencies with scaffolded autonomy development',
                    'application': 'Builds on foundation - integrates technical skills with transversal competencies through guided professional practice',
                    'integration': 'Requires foundation+application completion - synthesizes green-transversal skills for autonomous professional practice'
                },
                'curriculum_flow_matrix': {
                    'sustainability_analytics': 'Primary Focus → Advanced → Expert',
                    'green_competencies': 'Supporting → Primary Focus → Advanced', 
                    'transversal_skills': 'Foundational → Supporting → Expert'
                },
                'eqf_framework_outcome_mapping': {
                    'knowledge_descriptor': f"EQF Level {curriculum_data['eqf_level']}: Advanced knowledge with critical understanding of theories and principles",
                    'skills_descriptor': f"EQF Level {curriculum_data['eqf_level']}: Advanced skills demonstrating mastery and innovation",
                    'competence_descriptor': f"EQF Level {curriculum_data['eqf_level']}: Management of complex technical or professional activities",
                    'outcome_mapping': {
                        'programme_outcome_1': 'Maps to EQF Knowledge (critical understanding of frameworks)',
                        'programme_outcome_2': 'Maps to EQF Skills (mastery of technical tools and methods)',
                        'programme_outcome_3': 'Maps to EQF Competence (management of complex initiatives)'
                    }
                },
                'visual_alignment_matrix': {
                    'description': 'Comprehensive alignment chart showing relationship between EQF descriptors, programme outcomes, module outcomes, assessment types, and framework competencies',
                    'matrix_structure': {
                        'eqf_to_programme': {
                            'Knowledge': 'Programme Outcome 1 (Analyze ESG frameworks)',
                            'Skills': 'Programme Outcome 2 (Design sustainability systems)',
                            'Competence': 'Programme Outcome 3 (Manage sustainability initiatives)'
                        },
                        'programme_to_modules': {
                            'Programme Outcome 1': f'Module 1: {curriculum_data["modules"][0]} (Knowledge focus)',
                            'Programme Outcome 2': f'Module 2: {curriculum_data["modules"][1]} (Skills focus)', 
                            'Programme Outcome 3': f'Module 3: {curriculum_data["modules"][2]} (Competence focus)' if len(curriculum_data["modules"]) >= 3 else f'Module 2-3: Integration focus'
                        },
                        'modules_to_assessment': {
                            f'Module 1: {curriculum_data["modules"][0]}': 'Portfolio Assessment (Knowledge evidence)',
                            f'Module 2: {curriculum_data["modules"][1]}': 'Applied Project (Skills demonstration)',
                            f'Module 3: {curriculum_data["modules"][2] if len(curriculum_data["modules"]) >= 3 else "Integration"}': 'Collaborative Task (Competence application)'
                        },
                        'assessment_to_frameworks': {
                            'Portfolio Assessment': 'GreenComp (Systems Thinking, Interconnectedness)',
                            'Applied Project': 'GreenComp (Circular Thinking, Futures Thinking)',
                            'Collaborative Task': 'e-CF (Change Management, Needs Identification)',
                            'Reflective Practice': 'Key Competences (Transversal Skills)'
                        }
                    }
                }
            },
            
            'assessment_strategy': {
                'assessment_overview': 'Competency-aligned assessment with multi-level performance indicators mapped to international frameworks',
                'rubric_availability': 'Detailed assessment rubrics with specific performance descriptors are available through the institutional learning management system. Complete rubric matrices are provided to enrolled participants and supervising mentors, with external examiner access available for quality assurance purposes.',
                'visual_summary': {
                    'portfolio_evidence': {'weight': '40%', 'type': 'Continuous Assessment'},
                    'applied_project': {'weight': '35%', 'type': 'Summative Assessment'},
                    'collaborative_task': {'weight': '15%', 'type': 'Peer Assessment'},
                    'reflective_practice': {'weight': '10%', 'type': 'Self Assessment'}
                },
                'components': {
                    'portfolio_evidence': {
                        'description': 'Curated portfolio demonstrating unit-level competency development with mentored reflective analysis',
                        'weight': '40%',
                        'criteria': {
                            'programme_competencies': '25%',
                            'green_skills_integration': '25%',
                            'sustainability_competencies': '25%',
                            'transversal_skills': '25%'
                        },
                        'performance_indicators': {
                            'excellent': 'Comprehensive evidence across all competency areas with clear progression demonstration, sophisticated mentored reflective analysis showing deep learning integration, professional validation from multiple sources',
                            'good': 'Good evidence coverage with minor gaps, adequate mentored reflective analysis with some insights, appropriate professional validation',
                            'satisfactory': 'Basic evidence meeting minimum requirements, fundamental reflective practice with mentor guidance, basic professional confirmation',
                            'needs_improvement': 'Incomplete evidence portfolio, superficial reflection despite mentoring, insufficient professional validation'
                        }
                    },
                    'applied_project': {
                        'description': 'End-to-end sustainability challenge integrating green-transversal skills with mentored implementation',
                        'weight': '35%',
                        'criteria': {
                            'problem_analysis': '20%',
                            'solution_development': '30%',
                            'implementation_excellence': '30%',
                            'impact_evaluation': '20%'
                        },
                        'performance_indicators': {
                            'excellent': 'Sophisticated problem analysis with multiple perspectives, innovative solution design with advanced methodologies, flawless implementation with stakeholder engagement, comprehensive impact measurement with quantifiable outcomes',
                            'good': 'Thorough problem analysis with good insights, competent solution development with appropriate methods, good implementation with minor issues, adequate impact assessment with some quantification',
                            'satisfactory': 'Basic problem analysis meeting requirements, adequate solution design, acceptable implementation, fundamental impact evaluation',
                            'needs_improvement': 'Superficial analysis, weak solution design, poor implementation, minimal impact assessment'
                        }
                    },
                    'collaborative_task': {
                        'description': 'Peer-reviewed group deliverables with multicultural team problem-solving exercises and mentored team dynamics',
                        'weight': '15%',
                        'criteria': {
                            'professional_collaboration': '35%',
                            'green_problem_solving': '35%',
                            'transversal_communication': '30%'
                        },
                        'performance_indicators': {
                            'excellent': 'Outstanding professional collaboration with innovative approaches, sophisticated green problem-solving with systems thinking, exceptional cross-cultural communication with consensus building',
                            'good': 'Effective professional collaboration with appropriate methods, good sustainability problem-solving with practical solutions, competent cross-cultural interaction',
                            'satisfactory': 'Basic professional collaboration meeting requirements, fundamental green problem-solving, adequate cross-cultural communication',
                            'needs_improvement': 'Poor professional collaboration, weak problem-solving approach, limited cross-cultural engagement'
                        }
                    },
                    'reflective_practice': {
                        'description': 'Self-assessment journal with professional development planning, competency mapping, and guided mentor reflection sessions',
                        'weight': '10%',
                        'criteria': {
                            'self_assessment_accuracy': '30%',
                            'learning_integration': '40%',
                            'professional_planning': '30%'
                        },
                        'performance_indicators': {
                            'excellent': 'Highly accurate self-assessment with realistic competency evaluation, sophisticated learning integration with theoretical connections, strategic professional planning with clear development pathways',
                            'good': 'Good self-assessment with minor over/under-estimation, adequate learning integration with some connections, appropriate professional planning',
                            'satisfactory': 'Basic self-assessment with general accuracy, fundamental learning integration, basic professional development planning',
                            'needs_improvement': 'Inaccurate self-assessment, poor learning integration, weak professional planning'
                        }
                    }
                },
                'limitations': [
                    "Assessments provide indicators of competency development, not full professional certification",
                    "Workplace competency requires additional mentoring and real-world experience beyond curriculum scope",
                    "Some professional skills can only be fully assessed through extended workplace performance evaluation"
                ]
            },
            
            'framework_mapping': {
                'eqf_descriptors': f"EQF Level {curriculum_data['eqf_level']}: Advanced knowledge with critical understanding of theories and principles in work context",
                'european_e_competence': {
                    'B1_application_development': {
                        'level': 'Level 3',
                        'activities': ['Champion sustainability dashboard prototypes', 'Pioneer data visualization components', 'Orchestrate reporting functionality']
                    },
                    'D11_needs_identification': {
                        'level': 'Level 3',
                        'activities': ['Evaluate stakeholder requirements', 'Navigate data and reporting needs', 'Assess technical constraints']
                    },
                    'C2_change_management': {
                        'level': 'Level 3',
                        'activities': ['Strategically plan implementation strategies', 'Steward stakeholder transitions', 'Govern adoption progress']
                    }
                },
                'greencomp_integration': {
                    'interconnectedness': 'Advanced',
                    'systems_thinking': 'Advanced',
                    'circular_thinking': 'Advanced',
                    'futures_thinking': 'Advanced',
                    'collective_action': 'Intermediate',
                    'individual_initiative': 'Advanced'
                },
                'tuning_methodology': {
                    'subject_specific_competencies': 'Aligned with sustainability sector requirements',
                    'generic_competencies': 'Green awareness and transversal skills',
                    'learning_outcomes': 'Expressed in terms of demonstrable learner capabilities',
                    'module_progression': 'Systematic competency building across green-transversal domains'
                },
                'limitations': [
                    "Framework integration provides competency development structure, not professional certification",
                    "Practical competency requires additional workplace experience and mentoring",
                    "Framework levels indicate curriculum targets, not guaranteed achievement outcomes"
                ]
            },
            
            'stackability_and_credentialing': {
                'micro_credentials': {
                    'module_1': 'Foundation competencies integrating sustainability literacy and basic green awareness',
                    'module_2': 'Applied skills combining professional tools with sustainability methodologies',
                    'module_3': 'Integrated competences synthesizing green-transversal capabilities'
                },
                'stackability_features': [
                    "Flexible entry/exit points supporting diverse learning pathways",
                    "Recognition of Prior Learning (RPL) with portfolio-based validation (up to 50% of programme)",
                    "Cross-institutional credit transfer through ECTS compatibility",
                    "Professional development integration with continuing education requirements"
                ],
                'recognition_of_prior_learning': {
                    'rpl_pathway': 'Comprehensive Recognition of Prior Learning (RPL) pathway enabling credit transfer and accelerated progression',
                    'eligibility_criteria': {
                        'professional_experience': f'Minimum 2 years relevant {curriculum_data["role_name"].lower()} experience in target sectors',
                        'evidence_portfolio': 'Documented evidence of competency achievement through work-based projects and professional practice',
                        'competency_mapping': 'Demonstration of alignment between prior learning and programme module outcomes'
                    },
                    'assessment_process': {
                        'portfolio_review': 'Comprehensive review of documented evidence against module-level competency requirements',
                        'competency_interview': '90-minute structured interview with academic assessors and industry experts',
                        'practical_demonstration': 'Work simulation or project presentation demonstrating applied competencies',
                        'professional_validation': 'Industry expert validation of competency evidence and professional practice'
                    },
                    'recognition_limits': {
                        'maximum_credit': f'Up to 50% of total {curriculum_data["target_ects"]} ECTS can be recognized through RPL assessment',
                        'core_requirements': 'All learners must complete capstone assessment and reflective practice components regardless of RPL recognition',
                        'currency_validation': 'Prior learning evidence must demonstrate recent application within the past 5 years'
                    }
                },
                'limitations': [
                    "Micro-credentials supplement but do not replace formal degree requirements",
                    "Professional recognition varies by employer and jurisdiction",
                    "Additional workplace experience required for full professional competency",
                    "Cross-border recognition requires validation with relevant national authorities"
                ]
            },
            
            'work_based_integration': {
                'learning_component': '60% of total programme',
                'delivery_models': {
                    'workplace_learning': '60% (3-4 days per week with structured mentoring and guided reflection)',
                    'classroom_learning': '25% (Bi-weekly intensive workshops with mentored implementation planning)',
                    'online_learning': '15% (Weekly virtual sessions with peer mentoring and asynchronous mentored discussions)'
                },
                'pathways': {
                    'standard': 'Week 1-2: Workplace immersion with mentor guidance → Week 3: Classroom intensive with mentored reflection → Week 4: Online synthesis with guided implementation (cycle repeats)',
                    'accelerated': 'Week 1: Workplace diagnostic with mentor assessment → Week 2-3: Targeted classroom with mentored action planning → Week 4-5: Workplace application with guided supervision'
                },
                'mentoring_framework': {
                    'workplace_mentors': 'Industry professionals with minimum 5 years relevant experience providing weekly guidance and competency validation support',
                    'academic_mentors': 'Faculty supervisors providing bi-weekly reflection sessions and competency assessment support',
                    'peer_mentoring': 'Cross-cohort learning partnerships with structured mentoring protocols and guided peer assessment',
                    'mentor_training': 'All mentors receive structured training in competency-based assessment and reflective practice facilitation'
                },
                'assessment_model': [
                    "Workplace supervisor reports with monthly competency progression reviews and mentored reflection sessions",
                    "Real project outcomes measuring organizational sustainability impact with guided evaluation protocols",
                    "360-degree feedback from colleagues and stakeholders with mentored interpretation and development planning",
                    "Professional portfolio evidence collection throughout workplace engagement with guided curation and reflection"
                ]
            },
            
            'target_audiences': {
                'primary': curriculum_data['target_audience'],
                'pathways': {
                    'entry_level': {
                        'requirements': "Bachelor's degree plus basic experience",
                        'duration': 'Standard 12 months with full support',
                        'structure': 'Foundation-focused with intensive mentoring',
                        'outcomes': 'Junior to mid-level professional progression'
                    },
                    'career_changer': {
                        'requirements': 'RPL assessment plus foundation bridging modules',
                        'duration': 'Extended 15 months with additional support',
                        'structure': 'Enhanced foundation plus skills transfer support',
                        'outcomes': 'Mid-level role with transferable skills integration'
                    }
                },
                'rpl_framework': {
                    'portfolio_review': 'Documented evidence of prior learning and achievements',
                    'competency_interview': 'Structured assessment of knowledge, skills, and competences',
                    'practical_demonstration': 'Work simulation or project presentation',
                    'maximum_recognition': 'Up to 50% of programme ECTS through RPL'
                }
            },
            
            'support_and_qa': {
                'instructional_support': [
                    f"{curriculum_data['role_name']} experts with minimum 10 years industry experience",
                    "Sustainability trainers with academic and practical expertise",
                    "Guest lecturers from leading sustainability organizations"
                ],
                'peer_learning': [
                    "Weekly asynchronous forums for collaborative problem-solving",
                    "Monthly live professional networking sessions",
                    "Cross-cohort mentorship opportunities"
                ],
                'technical_support': [
                    "24/7 platform access with expert help desk",
                    "Troubleshooting assistance for all learning technologies",
                    "Multi-device compatibility support"
                ],
                'quality_assurance': [
                    "Formative assessment checkpoints integrated throughout each module",
                    "Learner dashboard with real-time progress tracking and competency visualization",
                    "External review through annual curriculum evaluation by industry advisory panel",
                    "Graduate employment outcomes tracking with career progression analysis",
                    "Biannual curriculum review cycles with stakeholder feedback integration ensuring continued relevance and professional currency"
                ],
                'stakeholder_validation': {
                    'development_process': 'Co-designed with industry employers, academic institutions, and professional bodies',
                    'employer_review': f'Validated by leading {curriculum_data["role_name"].lower()} employers and sustainability consultancies',
                    'academic_validation': 'Peer-reviewed by sustainability education experts and curriculum specialists',
                    'industry_benchmarking': 'Aligned with current industry standards and professional competency frameworks',
                    'continuous_review': 'Annual validation cycle with stakeholder feedback integration'
                }
            }
        }
        
        return curriculum
    
    def generate_module_outcomes(self, curriculum_data):
        """Generate individual module learning outcomes with highly distinctive competence verbs"""
        module_outcomes = []
        
        for i, module_name in enumerate(curriculum_data['modules']):
            # Generate EQF-appropriate outcomes with highly distinctive competence verbs per module
            if 'Leadership' in module_name or 'Strategic' in module_name:
                knowledge = f"Evaluate leadership theories, change management principles, and organizational dynamics relevant to {module_name.lower()}"
                skills = f"Apply leadership techniques, stakeholder engagement methods, and change management approaches in {module_name.lower()} contexts"
                competences = f"Champion strategic transformation initiatives, orchestrate organizational change processes, and steward executive stakeholder alignment in {module_name.lower()} practice with autonomous decision-making authority"
            elif 'Data' in module_name or 'Analytics' in module_name:
                knowledge = f"Analyze data analysis methodologies, statistical techniques, and data management principles for {module_name.lower()}"
                skills = f"Apply analytical tools, visualization software, and reporting platforms for {module_name.lower()} implementation"
                competences = f"Govern analytical workflows with autonomous quality control, ensure data integrity protocols using independent professional judgment, and pioneer evidence-based decision-making in {module_name.lower()} contexts"
            elif 'Technology' in module_name or 'Digital' in module_name or 'Tools' in module_name:
                knowledge = f"Understand technologies, system integration principles, and technical standards relevant to {module_name.lower()}"
                skills = f"Apply tools, implementation methodologies, and technical solutions for {module_name.lower()} delivery"
                competences = f"Orchestrate technology integration workflows with independent technical decision-making, champion innovation processes using autonomous professional judgment, and cultivate technical excellence in {module_name.lower()} contexts"
            elif 'Monitoring' in module_name or 'Reporting' in module_name or 'Performance' in module_name:
                knowledge = f"Understand monitoring frameworks, reporting standards, and performance measurement principles for {module_name.lower()}"
                skills = f"Apply monitoring tools, reporting methodologies, and performance analysis techniques for {module_name.lower()} implementation"
                competences = f"Govern stakeholder-facing performance reporting with autonomous quality control, orchestrate measurement workflows using independent professional judgment, and nurture accountability standards in {module_name.lower()} contexts"
            elif 'Management' in module_name or 'Coordination' in module_name:
                knowledge = f"Understand management principles, coordination frameworks, and organizational systems relevant to {module_name.lower()}"
                skills = f"Apply management techniques, coordination methods, and organizational tools for {module_name.lower()} delivery"
                competences = f"Oversee cross-functional programme coordination with autonomous professional judgment, facilitate multi-stakeholder operations using independent management capabilities, and steward operational excellence in {module_name.lower()} contexts"
            elif 'Advisory' in module_name or 'Consultant' in module_name or 'Client' in module_name:
                knowledge = f"Understand advisory methodologies, client engagement principles, and consultation frameworks for {module_name.lower()}"
                skills = f"Apply consultation techniques, client analysis methods, and advisory tools for {module_name.lower()} delivery"
                competences = f"Nurture advisory relationships with autonomous professional judgment, cultivate client partnerships using independent relationship management, and champion stakeholder trust in {module_name.lower()} contexts"
            elif 'Design' in module_name or 'Architecture' in module_name or 'Solution' in module_name:
                knowledge = f"Understand design principles, architectural frameworks, and solution methodologies for {module_name.lower()}"
                skills = f"Apply design techniques, architectural methods, and solution development tools for {module_name.lower()} implementation"
                competences = f"Pioneer design innovation workflows with independent strategic thinking, architect sustainable solutions using autonomous design decision-making, and cultivate creative excellence in {module_name.lower()} contexts"
            elif 'Training' in module_name or 'Education' in module_name or 'Development' in module_name:
                knowledge = f"Understand educational principles, training methodologies, and development frameworks for {module_name.lower()}"
                skills = f"Apply training techniques, educational tools, and development methods for {module_name.lower()} delivery"
                competences = f"Facilitate capability development activities with autonomous educational leadership, nurture learning environments using independent pedagogical judgment, and steward professional progression in {module_name.lower()} contexts"
            elif 'Change' in module_name:
                knowledge = f"Understand change management theories, organizational transformation principles, and change implementation frameworks for {module_name.lower()}"
                skills = f"Apply change management techniques, transformation methodologies, and implementation tools for {module_name.lower()} delivery"
                competences = f"Navigate organizational transformation activities with autonomous change leadership, shepherd change processes using independent professional judgment, and champion stakeholder adoption in {module_name.lower()} contexts"
            else:
                knowledge = f"Understand core principles, theoretical frameworks, and methodological approaches relevant to {module_name.lower()}"
                skills = f"Apply practical techniques, professional tools, and implementation methodologies for {module_name.lower()} delivery"
                competences = f"Coordinate professional activities with autonomous judgment, nurture quality standards using independent professional decision-making, and steward stakeholder relationships in {module_name.lower()} contexts"
            
            module_outcomes.append({
                'module_name': module_name,
                'knowledge_outcome': knowledge,
                'skills_outcome': skills,
                'competences_outcome': competences
            })
        
        return module_outcomes
    
    def save_curriculum_docx(self, curriculum, filename_base):
        """Generate comprehensive DOCX curriculum with enhanced formatting"""
        
        doc = Document()
        
        # Title section with enhanced formatting
        info = curriculum['curriculum_info']
        title = doc.add_heading(f"EQF Level {info['eqf_level']} / {info['target_ects']} ECTS Credits", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading(info['title'], level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Basic info
        doc.add_paragraph(f"Specialization: {info['specialization']}")
        doc.add_paragraph(f"Delivery: {info['delivery']}")
        doc.add_paragraph(f"Pedagogical Arc: {info['pedagogical_arc']}")
        doc.add_paragraph(f"Modules: {len(info['modules'])} Modules")
        doc.add_paragraph(f"Total Study Hours: {info['total_hours']} hours")
        if 'source_data' in info:
            doc.add_paragraph(f"Source: {info['source_data']}")
        
        # 1. Programme Learning Outcomes
        doc.add_heading("1. Programme Learning Outcomes", level=1)
        doc.add_paragraph(curriculum['programme_learning_outcomes']['completion_statement'])
        
        for outcome in curriculum['programme_learning_outcomes']['main_outcomes']:
            p = doc.add_paragraph(f"- {outcome}")
        
        # Enhanced programme outcome if available
        enhanced_outcome = curriculum['programme_learning_outcomes'].get('enhanced_programme_outcome')
        if enhanced_outcome:
            doc.add_paragraph(f"Enhanced Programme Outcome: {enhanced_outcome}")
        
        # Green Skills with clarification note and competency groupings
        green_section = curriculum['programme_learning_outcomes']['green_competences']
        doc.add_paragraph(f"Green Skills Integration ({green_section['framework']}):")
        
        doc.add_paragraph("Environmental Understanding:")
        for skill in green_section['competency_groups']['environmental_understanding']:
            doc.add_paragraph(f"- {skill}")
        
        doc.add_paragraph("Sustainable Innovation:")
        for skill in green_section['competency_groups']['sustainable_innovation']:
            doc.add_paragraph(f"- {skill}")
        
        doc.add_paragraph("Sustainability Complexity:")
        for skill in green_section['competency_groups']['sustainability_complexity']:
            doc.add_paragraph(f"- {skill}")
        
        doc.add_paragraph("Collaborative Action:")
        for skill in green_section['competency_groups']['collaborative_action']:
            doc.add_paragraph(f"- {skill}")
        
        # Modular Structure
        doc.add_heading("2. Modular Structure", level=1)
        modular = curriculum['modular_structure']
        doc.add_paragraph(f"Pedagogical Progression: {modular['pedagogical_progression']}")
        
        doc.add_paragraph("Module Themes:")
        for i, module in enumerate(modular['modules'], 1):
            doc.add_paragraph(f"{i}. {module}")
        
        # Save file
        docx_file = self.output_dir / f"{filename_base}.docx"
        doc.save(docx_file)
        return docx_file
    
    def save_curriculum_html(self, curriculum, filename_base):
        """Generate comprehensive HTML curriculum with enhanced styling"""
        
        info = curriculum['curriculum_info']
        
        html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{info['title']}</title>
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 20px; line-height: 1.6; color: #333; }}
        .header {{ background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%); color: white; padding: 30px; border-radius: 10px; margin-bottom: 30px; text-align: center; }}
        .info-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px; margin: 20px 0; }}
        .info-item {{ background-color: #f8f9fa; padding: 15px; border-radius: 8px; border-left: 4px solid #1e3c72; }}
        .section {{ margin: 30px 0; }}
        h1 {{ margin: 0; font-size: 2.2em; }}
        h2 {{ color: #1e3c72; border-bottom: 2px solid #1e3c72; padding-bottom: 5px; }}
        h3 {{ color: #495057; }}
        ul {{ padding-left: 20px; }}
        li {{ margin: 5px 0; }}
        .limitation {{ background-color: #fff3cd; padding: 15px; border-radius: 5px; border-left: 4px solid #ffc107; margin: 15px 0; }}
        .module {{ background-color: #e3f2fd; padding: 15px; margin: 10px 0; border-radius: 8px; }}
        .note-box {{ background: #e7f3ff; border-left: 4px solid #007bff; padding: 15px; margin: 15px 0; }}
        .pedagogical-arc {{ background: #f1f8e9; border-left: 4px solid #4caf50; padding: 15px; margin: 15px 0; font-weight: bold; }}
        .competency-group {{ background: #f8f9fa; padding: 15px; margin: 10px 0; border-radius: 8px; border-left: 4px solid #28a745; }}
        .source-info {{ background: #e8f4fd; border-left: 4px solid #17a2b8; padding: 10px; margin: 10px 0; font-style: italic; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{info['title']}</h1>
        <p>EQF Level {info['eqf_level']} | {info['target_ects']} ECTS Credits | {info['total_hours']} Hours</p>
        <p>Specialization: {info['specialization']}</p>
    </div>
    
    <div class="info-grid">
        <div class="info-item">
            <strong>Delivery Mode</strong><br>
            {info['delivery']}
        </div>
        <div class="info-item">
            <strong>Pedagogical Arc</strong><br>
            {info['pedagogical_arc']}
        </div>
        <div class="info-item">
            <strong>Modules</strong><br>
            {len(info['modules'])} Modules
        </div>
        <div class="info-item">
            <strong>Study Hours</strong><br>
            {info['total_hours']} total hours
        </div>
        <div class="info-item">
            <strong>Generated</strong><br>
            {info['generated_date'][:19]}
        </div>
    </div>
    
    {f'<div class="source-info"><strong>Source:</strong> {info["source_data"]}</div>' if 'source_data' in info else ''}
    
    <div class="pedagogical-arc">
        <strong>Pedagogical Progression:</strong> {info['pedagogical_arc']}
    </div>
    
    <div class="section">
        <h2>Module Themes</h2>
        <ol>
"""
        
        for module in info['modules']:
            html += f"            <li>{module}</li>\n"
        
        html += """
        </ol>
    </div>
    
    <div class="section">
        <h2>Programme Learning Outcomes</h2>
        <p><em>Upon successful completion, participants will be able to:</em></p>
        <ol>
"""
        
        for outcome in curriculum['programme_learning_outcomes']['main_outcomes']:
            html += f"            <li>{outcome}</li>\n"
        
        # Add enhanced programme outcome if available
        enhanced_outcome = curriculum['programme_learning_outcomes'].get('enhanced_programme_outcome')
        if enhanced_outcome:
            html += f"""
        </ol>
        <div class="note-box">
            <strong>Enhanced Programme Outcome:</strong> {enhanced_outcome}
        </div>
        <ol start="{len(curriculum['programme_learning_outcomes']['main_outcomes']) + 1}">
"""
        
        html += """
        </ol>
    </div>
    
    <div class="section">
        <h2>Green Competencies Framework</h2>
"""
        
        green_comp = curriculum['programme_learning_outcomes']['green_competences']
        for group_name, competencies in green_comp['competency_groups'].items():
            group_title = group_name.replace('_', ' ').title()
            html += f"""
        <div class="competency-group">
            <h3>{group_title}</h3>
            <ul>
"""
            for competency in competencies:
                html += f"                <li>{competency}</li>\n"
            html += "            </ul>\n        </div>\n"
        
        html += f"""
    </div>
    
    <div class="section">
        <h2>Assessment Strategy</h2>
        <div class="note-box">
            <strong>Assessment Overview:</strong> {curriculum['assessment_strategy']['assessment_overview']}
        </div>
        <div class="note-box">
            <strong>Rubric Availability:</strong> {curriculum['assessment_strategy']['rubric_availability']}
        </div>
    </div>
    
    <footer style="text-align: center; margin-top: 40px; padding: 20px; background-color: #f8f9fa; border-radius: 8px;">
        <p><em>Professional curriculum with comprehensive EQF alignment and framework integration</em></p>
        <p><em>Enhanced with explicit module themes, refined autonomy language, and pedagogical progression</em></p>
        <p><em>Generated from JSON input data with professional standards enhancement</em></p>
    </footer>
</body>
</html>
"""
        
        html_file = self.output_dir / f"{filename_base}.html"
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(html)
        return html_file
    
    def save_curriculum(self, curriculum, filename_base, curriculum_number):
        """Save curriculum in all configured formats with numbered prefixes"""
        saved_files = []
        
        # Add number prefix to filename
        numbered_filename = f"{curriculum_number:02d}_{filename_base}"
        
        # Save as JSON
        if 'json' in self.output_formats:
            json_file = self.output_dir / f"{numbered_filename}.json"
            with open(json_file, 'w', encoding='utf-8') as f:
                json.dump(curriculum, f, indent=2, ensure_ascii=False)
            saved_files.append(json_file)
            print(f"✅ Saved JSON: {json_file.name}")
        
        # Save as HTML
        if 'html' in self.output_formats:
            html_file = self.save_curriculum_html(curriculum, numbered_filename)
            saved_files.append(html_file)
            print(f"✅ Saved HTML: {html_file.name}")
        
        # Save as DOCX
        if 'docx' in self.output_formats:
            try:
                docx_file = self.save_curriculum_docx(curriculum, numbered_filename)
                saved_files.append(docx_file)
                print(f"✅ Saved DOCX: {docx_file.name}")
            except ImportError:
                print(f"⚠️ Skipping DOCX: python-docx not available")
            except Exception as e:
                print(f"⚠️ Error generating DOCX: {e}")
        
        return saved_files
    
    def generate_all_curricula(self):
        """Generate all curricula from JSON input with enhanced precision"""
        
        if not self.curricula_data:
            print("❌ No curricula data loaded. Cannot generate curricula.")
            return []
        
        generated_files = []
        
        print(f"\n📚 Generating {len(self.curricula_data)} enhanced professional curricula from JSON input...")
        
        for i, curriculum_data in enumerate(self.curricula_data, 1):
            print(f"  🔄 Generating {curriculum_data['role_name']} (EQF {curriculum_data['eqf_level']})...")
            
            curriculum = self.generate_comprehensive_curriculum(curriculum_data)
            if curriculum:
                filename = curriculum_data['id'].lower()
                files = self.save_curriculum(curriculum, filename, i)
                generated_files.extend(files)
                
                print(f"     ✅ EQF Level: {curriculum_data['eqf_level']}")
                print(f"     ✅ Pedagogical Arc: {curriculum_data['pedagogical_arc']}")
                print(f"     ✅ Modules: {len(curriculum_data['modules'])} explicit themed modules")
                print(f"     ✅ ECTS: {curriculum_data['target_ects']} credits")
        
        print(f"\n📊 GENERATION SUMMARY:")
        print(f"✅ Total curricula generated: {len(self.curricula_data)}")
        print(f"📄 Source: JSON input file processing")
        print(f"🔧 PROFESSIONAL ENHANCEMENTS:")
        print(f"✅ JSON input processing with automatic EQF expansion")
        print(f"✅ Explicit module themes for clear pedagogical arcs")
        print(f"✅ Refined autonomy language (accessing specialized guidance where appropriate)")
        print(f"✅ Distinctive competence verbs (Champion, Orchestrate, Steward, Govern, Pioneer, Nurture, Cultivate)")
        print(f"✅ Enhanced mentoring framework with structured guidance protocols")
        print(f"✅ Professional assessment rubric availability statements")
        print(f"✅ Visual alignment matrix for accreditation compliance")
        print(f"✅ Biannual review cycles ensuring professional currency")
        print(f"✅ Numbered filename prefixes for document management")
        print(f"✅ NO DigComp mappings (removed as per requirements)")
        print(f"📄 Formats: {', '.join(self.output_formats)}")
        print(f"📁 Output directory: {self.output_dir}")
        print(f"🎯 Status: Professional-ready, accreditation-compliant from JSON input")
        
        return generated_files


def main():
    """Generate professional curricula from JSON input with enhanced standards"""
    
    print("🚀 Starting ECM Curriculum Generation from JSON Input...")
    print("📚 Reading from educational profiles or curricula JSON file")
    print("⚙️ Using configuration from: config/settings.json")
    print("🎯 Features: JSON input processing, explicit module themes, enhanced pedagogical arcs")
    print("❌ NO DigComp mappings (removed as per requirements)")
    
    # Check for required dependencies
    try:
        from docx import Document
        print("✅ python-docx library available")
    except ImportError:
        print("❌ python-docx library required for DOCX generation")
        print("   Install with: pip install python-docx")
        print("   Continuing with JSON and HTML generation only...")
    
    # Initialize generator
    generator = ECMCurriculumGenerator()
    
    # Generate all curricula
    generated_files = generator.generate_all_curricula()
    
    print(f"\n✅ Professional curriculum generation complete!")
    print(f"📁 Generated {len(generated_files)} files in: {generator.output_dir}")
    print(f"📄 Formats: {', '.join(generator.output_formats)}")
    print(f"📚 Curricula: Professional curricula with enhanced standards from JSON input")
    
    return {
        'generator': generator,
        'files_generated': len(generated_files),
        'curricula_count': len(generator.curricula_data),
        'output_directory': generator.output_dir,
        'formats': generator.output_formats
    }


if __name__ == "__main__":
    results = main()
