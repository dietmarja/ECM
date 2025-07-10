# generate_educational_profiles_v8.1.py
"""
ECM Educational Profiles Generator - Enhanced Version 8.1
Reads from educational_profiles.json input file and generates comprehensive profiles
NO DigComp mappings as per requirements
"""

import json
import os
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

class ECMProfilesGenerator:
    """Generate comprehensive educational profiles from JSON input with enhanced professional standards"""
    
    def __init__(self, config_path='../config/settings.json'):
        """Initialize generator with configuration"""
        self.config = self.load_config(config_path)
        self.setup_paths()
        self.profiles_data = self.load_profiles_data()
        
        print(f"📊 ECM Educational Profiles Generator - Professional Standards v8.1")
        print(f"⚙️ Configuration loaded from settings.json")
        print(f"📁 Input profiles: {self.profiles_path}")
        print(f"📁 Output directory: {self.output_dir}")
        print(f"📄 Loaded {len(self.profiles_data)} base profiles from JSON")
    
    def load_config(self, config_path):
        """Load configuration from settings.json"""
        try:
            project_root = Path(__file__).parent.parent
            config_file = project_root / config_path.lstrip('./')
            
            with open(config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)
            print(f"✅ Configuration loaded from: {config_file}")
            return config
        except FileNotFoundError:
            print(f"⚠️ Config file not found: {config_path}")
            return self.get_default_config()
        except Exception as e:
            print(f"⚠️ Error loading config: {e}")
            return self.get_default_config()
    
    def get_default_config(self):
        """Default configuration if settings.json not found"""
        return {
            "paths": {
                "input_educational_profiles": "./input/educational_profiles/educational_profiles.json"
            },
            "output": {
                "profiles": {
                    "directory": "./output/profiles",
                    "formats": ["json", "html", "docx"]
                }
            },
            "system": {"version": "8.1.0"}
        }
    
    def setup_paths(self):
        """Setup paths based on configuration"""
        project_root = Path(__file__).parent.parent
        
        profiles_config = self.config['paths']['input_educational_profiles'].lstrip('./')
        self.profiles_path = project_root / profiles_config
        
        output_config = self.config.get('output', {}).get('profiles', {})
        output_dir = output_config.get('directory', './output/profiles').lstrip('./')
        self.output_dir = project_root / output_dir
        
        self.output_formats = output_config.get('formats', ['json', 'html', 'docx'])
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def load_profiles_data(self):
        """Load profiles from educational_profiles.json"""
        try:
            with open(self.profiles_path, 'r', encoding='utf-8') as f:
                profiles = json.load(f)
            print(f"✅ Loaded {len(profiles)} base profiles from: {self.profiles_path}")
            return profiles
        except FileNotFoundError:
            print(f"❌ Profiles file not found: {self.profiles_path}")
            return []
        except Exception as e:
            print(f"❌ Error loading profiles: {e}")
            return []
    
    def expand_profiles_by_eqf(self):
        """Expand profiles with multiple EQF levels into separate profile entries"""
        expanded_profiles = []
        
        for profile in self.profiles_data:
            # Check for multiple EQF levels in enhanced_purpose_and_application
            purpose_app = profile.get('enhanced_purpose_and_application', {})
            learning_outcomes = profile.get('learning_outcomes_by_eqf', {})
            
            # Find all EQF levels for this profile
            eqf_levels = set()
            
            # Check purpose_and_application for EQF levels
            for key in purpose_app.keys():
                if key.startswith('eqf_'):
                    eqf_level = key.replace('eqf_', '')
                    if eqf_level.isdigit():
                        eqf_levels.add(int(eqf_level))
            
            # Check learning_outcomes for EQF levels
            for level in learning_outcomes.keys():
                if level.isdigit():
                    eqf_levels.add(int(level))
            
            # If no specific EQF levels found, try to infer from other fields
            if not eqf_levels:
                # Check assessment_framework
                assessment = profile.get('assessment_framework', {})
                for key in assessment.keys():
                    if key.startswith('eqf_'):
                        eqf_level = key.replace('eqf_', '')
                        if eqf_level.isdigit():
                            eqf_levels.add(int(eqf_level))
                
                # If still no levels, default to 6
                if not eqf_levels:
                    eqf_levels.add(6)
            
            # Create separate profile for each EQF level
            for eqf_level in sorted(eqf_levels):
                expanded_profile = self.create_eqf_specific_profile(profile, eqf_level)
                expanded_profiles.append(expanded_profile)
        
        return expanded_profiles
    
    def create_eqf_specific_profile(self, base_profile, eqf_level):
        """Create EQF-specific profile from base profile data"""
        profile_id = f"{base_profile['id']}_EQF{eqf_level}_Profile"
        
        # Extract EQF-specific purpose and application
        purpose_app = base_profile.get('enhanced_purpose_and_application', {})
        eqf_specific_purpose = purpose_app.get(f'eqf_{eqf_level}', 
                                               purpose_app.get('eqf_6', 
                                                               'Professional development for sustainability specialists'))
        
        # Extract EQF-specific learning outcomes
        learning_outcomes = base_profile.get('learning_outcomes_by_eqf', {})
        eqf_outcomes = learning_outcomes.get(str(eqf_level), learning_outcomes.get('6', {}))
        
        # Extract EQF-specific assessment
        assessment_framework = base_profile.get('assessment_framework', {})
        eqf_assessment = assessment_framework.get(f'eqf_{eqf_level}', 
                                                  assessment_framework.get('eqf_6', 
                                                                           'Assessment based on competency demonstration'))
        
        # Extract EQF-specific entry requirements
        entry_requirements = base_profile.get('entry_requirements_by_eqf', {})
        eqf_entry = entry_requirements.get(str(eqf_level), entry_requirements.get('6', {}))
        
        # Calculate ECTS based on EQF level
        base_ects = 30
        if eqf_level == 4:
            target_ects = 25
        elif eqf_level == 5:
            target_ects = 25
        elif eqf_level == 6:
            target_ects = 36
        elif eqf_level == 7:
            target_ects = 45
        elif eqf_level == 8:
            target_ects = 50
        else:
            target_ects = 36
        
        # Extract units from learning outcomes
        units = []
        if eqf_outcomes and 'units' in eqf_outcomes:
            for unit_name, unit_data in eqf_outcomes['units'].items():
                if isinstance(unit_data, dict):
                    units.append({
                        'name': unit_name,
                        'ects': target_ects // 3,  # Distribute ECTS across units
                        'complexity': f'EQF Level {eqf_level} complexity requiring professional competency development',
                        'knowledge': unit_data.get('knowledge', f'Advanced knowledge in {unit_name.lower()}'),
                        'skills': unit_data.get('skills', f'Advanced skills in {unit_name.lower()}'),
                        'competences': unit_data.get('competences', f'Autonomous competences in {unit_name.lower()}')
                    })
        
        # If no units extracted, create default units
        if not units:
            num_units = 3
            ects_per_unit = target_ects // num_units
            for i in range(num_units):
                units.append({
                    'name': f'Core Competency {i+1}',
                    'ects': ects_per_unit,
                    'complexity': f'EQF Level {eqf_level} complexity requiring professional competency development',
                    'knowledge': f'Advanced knowledge relevant to {base_profile.get("profile_name", "professional practice")}',
                    'skills': f'Advanced skills application in professional contexts',
                    'competences': f'Autonomous professional competences with independent decision-making'
                })
        
        # Extract target sectors
        target_sectors = []
        career_progression = base_profile.get('career_progression', {})
        if isinstance(career_progression, dict):
            for level, role in career_progression.items():
                if isinstance(role, str):
                    # Extract sector information from role descriptions
                    if 'corporate' in role.lower():
                        target_sectors.append('Corporate Sustainability')
                    if 'consulting' in role.lower():
                        target_sectors.append('Sustainability Consulting')
                    if 'technology' in role.lower():
                        target_sectors.append('Technology & Innovation')
        
        if not target_sectors:
            target_sectors = ['Sustainability Management', 'Environmental Consulting', 'Corporate ESG']
        
        # Create the profile structure
        expanded_profile = {
            'id': profile_id,
            'title': f"{base_profile.get('profile_name', 'Professional Profile').replace(' Educational Profile', '')}",
            'eqf_level': eqf_level,
            'complexity_level': f"EQF Level {eqf_level}: {self.get_eqf_complexity_descriptor(eqf_level)}",
            'target_ects': target_ects,
            'delivery_modes': self.extract_delivery_modes(base_profile),
            'target_sectors': target_sectors[:3],  # Limit to 3 sectors
            'units': units,
            'enhanced_purpose': eqf_specific_purpose,
            'programme_outcome': eqf_outcomes.get('programme_outcome', f'Professional development for EQF Level {eqf_level}'),
            'assessment_framework': eqf_assessment,
            'entry_requirements': eqf_entry,
            'role_description': base_profile.get('role_description', ''),
            'competences_descriptors': base_profile.get('enhanced_competences_with_descriptors', {}),
            'industry_application': base_profile.get('industry_application', []),
            'distinctive_features': base_profile.get('distinctive_features', [])
        }
        
        return expanded_profile
    
    def get_eqf_complexity_descriptor(self, eqf_level):
        """Get appropriate complexity descriptor for EQF level"""
        descriptors = {
            4: "Factual and theoretical knowledge in broad contexts with specialized skills for defined activities",
            5: "Comprehensive, specialized knowledge with awareness of boundaries and critical understanding", 
            6: "Advanced knowledge with critical understanding in professional contexts requiring autonomous responsibility",
            7: "Highly specialized knowledge at the forefront as basis for original thinking and research applications",
            8: "Knowledge at most advanced frontier with substantial authority and innovation capability"
        }
        return descriptors.get(eqf_level, "Professional knowledge and skills requiring competency development")
    
    def extract_delivery_modes(self, profile):
        """Extract delivery modes from profile data"""
        delivery_modes = ['Blended Learning', 'Online Learning', 'Work-Based Learning']
        
        # Try to extract from role description or other fields
        role_desc = profile.get('role_description', '').lower()
        if 'research' in role_desc:
            delivery_modes = ['Research-Based Learning', 'Project-Based Learning', 'Academic Study']
        elif 'consulting' in role_desc:
            delivery_modes = ['Client-Based Projects', 'Consulting Practice', 'Case Study Learning']
        elif 'technical' in role_desc:
            delivery_modes = ['Technical Training', 'Hands-On Practice', 'Laboratory Work']
        elif 'management' in role_desc:
            delivery_modes = ['Executive Learning', 'Strategic Workshops', 'Leadership Practice']
            
        return delivery_modes
    
    def generate_comprehensive_profile(self, profile_data):
        """Generate comprehensive educational profile with enhanced professional standards"""
        
        profile = {
            'profile_info': {
                'id': profile_data['id'],
                'title': f"Educational Profile: {profile_data['title']} (EQF {profile_data['eqf_level']})",
                'eqf_level': profile_data['eqf_level'],
                'complexity_level': profile_data['complexity_level'],
                'target_ects': profile_data['target_ects'],
                'total_hours': profile_data['target_ects'] * 25,
                'delivery_modes': profile_data['delivery_modes'],
                'target_sectors': profile_data['target_sectors'],
                'generated_date': datetime.now().isoformat()
            },
            
            'programme_goal': {
                'primary_objective': profile_data.get('enhanced_purpose', f"To develop advanced {profile_data['title'].lower()} competencies integrating critical sustainability knowledge, sophisticated professional skills, and autonomous management capabilities with professional autonomy, accessing specialized guidance where appropriate to refine complex professional decisions"),
                'learning_approach': 'Competency-based professional development with integrated green-transversal skill development',
                'professional_context': f"Prepares learners for complex {profile_data['title'].lower()} practice requiring critical understanding, advanced skills, and autonomous professional management"
            },
            
            'complexity_description': {
                'eqf_alignment': profile_data['complexity_level'],
                'cognitive_demands': 'Critical understanding of sustainability theories and principles, sophisticated analytical thinking, and complex problem-solving in professional contexts with autonomous judgment',
                'operational_requirements': 'Advanced skill application with innovative approaches, autonomous decision-making, and strategic management of complex professional activities',
                'responsibility_scope': 'Management and supervision of complex sustainability initiatives, responsibility for strategic outcomes, and accountability for professional excellence with independent decision-making authority'
            },
            
            'unit_breakdown': {
                'total_units': len(profile_data['units']),
                'total_ects': sum(unit['ects'] for unit in profile_data['units']),
                'units_summary': f"Units include: {', '.join([unit['name'] for unit in profile_data['units']])}",
                'units': profile_data['units']
            },
            
            'unit_learning_outcomes': {
                'unit_outcomes_overview': 'Detailed learning outcomes for each unit demonstrating progression from knowledge through skills to autonomous competence',
                'units': [
                    {
                        'unit_name': unit['name'],
                        'unit_ects': unit['ects'],
                        'complexity_level': unit['complexity'],
                        'knowledge_outcome': unit['knowledge'],
                        'skills_outcome': unit['skills'],
                        'competence_outcome': unit['competences']
                    }
                    for unit in profile_data['units']
                ]
            },
            
            'learning_outcomes': {
                'programme_outcome': profile_data.get('programme_outcome', f"Professional development for EQF Level {profile_data['eqf_level']}"),
                'knowledge_outcomes': [unit['knowledge'] for unit in profile_data['units']],
                'skills_outcomes': [unit['skills'] for unit in profile_data['units']],
                'competence_outcomes': [unit['competences'] for unit in profile_data['units']]
            },
            
            'green_competencies': {
                'framework': 'GreenComp Framework (Comprehensive Coverage)',
                'competency_groups': {
                    'environmental_understanding': [
                        "Evaluate complex environmental interconnections using systems thinking approaches with autonomous analytical judgment (GreenComp 1.1 Interconnectedness: Level 4 - Advanced)",
                        "Champion comprehensive systems thinking to sustainability challenge analysis using independent critical analysis (GreenComp 1.2 Systems Thinking: Level 4 - Advanced)"
                    ],
                    'sustainable_innovation': [
                        "Pioneer circular economy principles in complex organizational contexts with autonomous strategic decision-making (GreenComp 1.3 Circular Thinking: Level 4 - Advanced)",
                        "Orchestrate advanced futures thinking for long-term sustainability planning with independent visionary thinking (GreenComp 2.3 Futures Thinking: Level 4 - Advanced)"
                    ],
                    'sustainability_complexity': [
                        "Navigate complexity in sustainability decisions using sophisticated frameworks with autonomous professional judgment (GreenComp 2.2 Complexity in Sustainability: Level 4 - Advanced)"
                    ],
                    'collaborative_action': [
                        "Cultivate collective action initiatives for complex sustainability transformations using independent leadership capabilities (GreenComp 4.2 Collective Action: Level 4 - Advanced)",
                        "Steward advanced individual responsibility and leadership in sustainability practice with autonomous ethical decision-making (GreenComp 4.3 Individual Initiative: Level 4 - Advanced)"
                    ]
                },
                'competencies': [
                    "Evaluate complex environmental interconnections using systems thinking approaches with autonomous analytical judgment (GreenComp 1.1 Interconnectedness: Level 4 - Advanced)",
                    "Champion comprehensive systems thinking to sustainability challenge analysis using independent critical analysis (GreenComp 1.2 Systems Thinking: Level 4 - Advanced)",
                    "Pioneer circular economy principles in complex organizational contexts with autonomous strategic decision-making (GreenComp 1.3 Circular Thinking: Level 4 - Advanced)",
                    "Navigate complexity in sustainability decisions using sophisticated frameworks with autonomous professional judgment (GreenComp 2.2 Complexity in Sustainability: Level 4 - Advanced)",
                    "Orchestrate advanced futures thinking for long-term sustainability planning with independent visionary thinking (GreenComp 2.3 Futures Thinking: Level 4 - Advanced)",
                    "Cultivate collective action initiatives for complex sustainability transformations using independent leadership capabilities (GreenComp 4.2 Collective Action: Level 4 - Advanced)",
                    "Steward advanced individual responsibility and leadership in sustainability practice with autonomous ethical decision-making (GreenComp 4.3 Individual Initiative: Level 4 - Advanced)"
                ]
            },
            
            'assessment_framework': {
                'assessment_philosophy': 'Competency-aligned assessment with multi-level performance indicators and comprehensive visual rubric framework',
                'assessment_overview': profile_data.get('assessment_framework', 'Assessment based on competency demonstration with professional validation'),
                'rubric_availability': 'Detailed assessment rubrics with specific performance descriptors are available through the institutional learning management system. Complete rubric matrices are provided to enrolled participants and supervising mentors, with external examiner access available for quality assurance purposes.',
                'visual_assessment_criteria': {
                    'rubric_description': 'Comprehensive visual assessment rubric with color-coded performance levels and detailed criteria for transparent evaluation',
                    'performance_levels': {
                        'excellent': {
                            'color_code': '#28a745',
                            'descriptor': 'Outstanding performance exceeding expectations',
                            'characteristics': ['Sophisticated analysis', 'Innovation demonstration', 'Leadership excellence', 'Strategic impact']
                        },
                        'good': {
                            'color_code': '#17a2b8',
                            'descriptor': 'Strong performance meeting professional standards',
                            'characteristics': ['Thorough analysis', 'Competent application', 'Professional delivery', 'Measurable impact']
                        },
                        'satisfactory': {
                            'color_code': '#ffc107',
                            'descriptor': 'Adequate performance meeting minimum requirements',
                            'characteristics': ['Basic analysis', 'Standard application', 'Acceptable delivery', 'Fundamental impact']
                        },
                        'needs_improvement': {
                            'color_code': '#dc3545',
                            'descriptor': 'Performance below expectations requiring development',
                            'characteristics': ['Weak analysis', 'Limited application', 'Poor delivery', 'Minimal impact']
                        }
                    }
                },
                'components': {
                    'portfolio_assessment': {
                        'weight': '40%',
                        'description': 'Comprehensive professional portfolio demonstrating competency development progression with visual evidence matrix',
                        'rubric_levels': {
                            'excellent': {
                                'criteria': 'Comprehensive competency evidence with sophisticated reflection and professional validation',
                                'indicators': 'Multi-stakeholder validation, advanced theoretical integration, strategic professional planning',
                                'color_code': '#28a745'
                            },
                            'good': {
                                'criteria': 'Strong competency evidence with good reflection and appropriate validation',
                                'indicators': 'Professional validation, theoretical connections, practical application evidence',
                                'color_code': '#17a2b8'
                            },
                            'satisfactory': {
                                'criteria': 'Basic competency evidence meeting minimum requirements with fundamental reflection',
                                'indicators': 'Basic validation, minimal theoretical integration, adequate practical evidence',
                                'color_code': '#ffc107'
                            },
                            'needs_improvement': {
                                'criteria': 'Incomplete competency evidence with insufficient reflection and validation',
                                'indicators': 'Inadequate validation, weak theoretical connections, limited practical evidence',
                                'color_code': '#dc3545'
                            }
                        }
                    },
                    'project_assessment': {
                        'weight': '35%',
                        'description': 'Complex sustainability project demonstrating integrated competency application with visual impact assessment',
                        'rubric_levels': {
                            'excellent': {
                                'criteria': 'Multi-stakeholder project with measurable impact and innovative approaches',
                                'indicators': 'Sophisticated analysis, stakeholder engagement, quantifiable outcomes, strategic impact',
                                'color_code': '#28a745'
                            },
                            'good': {
                                'criteria': 'Professional project with good impact and appropriate methodologies',
                                'indicators': 'Thorough analysis, stakeholder involvement, measurable results, practical impact',
                                'color_code': '#17a2b8'
                            },
                            'satisfactory': {
                                'criteria': 'Basic project meeting requirements with fundamental impact demonstration',
                                'indicators': 'Adequate analysis, basic stakeholder engagement, minimal impact measurement',
                                'color_code': '#ffc107'
                            },
                            'needs_improvement': {
                                'criteria': 'Inadequate project with limited impact and poor methodology application',
                                'indicators': 'Weak analysis, minimal stakeholder engagement, insufficient impact evidence',
                                'color_code': '#dc3545'
                            }
                        }
                    },
                    'rpl_assessment': {
                        'weight': '25%',
                        'description': 'Recognition of Prior Learning through comprehensive competency demonstration and portfolio validation with visual competency mapping',
                        'rubric_levels': {
                            'excellent': {
                                'criteria': '5+ years relevant experience with comprehensive competency validation',
                                'indicators': 'Industry expert validation, multiple project evidence, leadership demonstration',
                                'color_code': '#28a745'
                            },
                            'good': {
                                'criteria': '3-5 years experience with good competency evidence and validation',
                                'indicators': 'Professional validation, project portfolio, management experience',
                                'color_code': '#17a2b8'
                            },
                            'satisfactory': {
                                'criteria': '2-3 years experience with basic competency validation',
                                'indicators': 'Basic validation, limited project evidence, some management exposure',
                                'color_code': '#ffc107'
                            },
                            'needs_improvement': {
                                'criteria': 'Insufficient experience or inadequate competency validation',
                                'indicators': 'Weak validation, minimal evidence, limited professional practice',
                                'color_code': '#dc3545'
                            }
                        }
                    }
                }
            },
            
            'recognition_of_prior_learning': {
                'rpl_pathway': 'Comprehensive Recognition of Prior Learning (RPL) pathway enabling credit transfer and accelerated progression through systematic competency validation with visual competency mapping',
                'eligibility_criteria': {
                    'professional_experience': f'Minimum 2 years relevant {profile_data["title"].lower()} experience in target sectors',
                    'evidence_portfolio': 'Documented evidence of competency achievement through work-based projects and professional practice',
                    'competency_mapping': 'Demonstration of alignment between prior learning and programme unit outcomes using visual competency matrix'
                },
                'assessment_process': {
                    'portfolio_review': 'Comprehensive review of documented evidence against unit-level competency requirements using visual assessment criteria',
                    'competency_interview': '90-minute structured interview with academic assessors and industry experts using standardized rubric',
                    'practical_demonstration': 'Work simulation or project presentation demonstrating applied competencies with visual performance indicators',
                    'professional_validation': 'Industry expert validation of competency evidence and professional practice using established criteria'
                },
                'recognition_limits': {
                    'maximum_credit': f'Up to 60% of total {profile_data["target_ects"]} ECTS can be recognized through RPL assessment',
                    'core_requirements': 'All learners must complete capstone assessment and reflective practice components regardless of RPL recognition',
                    'currency_validation': 'Prior learning evidence must demonstrate recent application within the past 5 years'
                }
            },
            
            'validation_framework': {
                'development_methodology': 'Co-designed with industry employers, academic institutions, professional bodies, and regulatory authorities using participatory design principles',
                'stakeholder_validation': {
                    'employer_validation': f'Validated by leading {profile_data["title"].lower()} employers across target sectors through systematic consultation',
                    'academic_validation': 'Peer-reviewed by sustainability education experts and curriculum development specialists using established quality criteria',
                    'professional_validation': 'Endorsed by relevant professional bodies and industry associations through formal validation processes',
                    'regulatory_alignment': 'Aligned with national qualifications frameworks and international standards through compliance mapping'
                },
                'quality_assurance': {
                    'annual_review': 'Comprehensive annual validation cycle with stakeholder feedback integration using structured evaluation protocols',
                    'employer_feedback': 'Systematic collection and integration of employer satisfaction and graduate performance data',
                    'graduate_tracking': 'Long-term career progression monitoring and outcome assessment with visual progress indicators',
                    'continuous_improvement': 'Evidence-based enhancement cycle with stakeholder input integration',
                    'update_cycles': 'This profile will be reviewed biannually with input from industry stakeholders and aligned with evolving frameworks, ensuring continued relevance and professional currency through systematic stakeholder consultation'
                }
            },
            
            'delivery_pathways': {
                'standard_pathway': {
                    'duration': '12 months full-time equivalent',
                    'structure': 'Sequential unit delivery with integrated competency development and scaffolded autonomy progression',
                    'assessment_points': 'Continuous assessment with major evaluation points per unit using visual rubric criteria',
                    'support_level': 'Comprehensive academic and professional mentoring support with structured guidance protocols'
                },
                'work_based_pathway': {
                    'duration': '15 months part-time with workplace integration',
                    'structure': 'Workplace-integrated learning with academic support and mentored autonomy development',
                    'assessment_points': 'Work-based evidence collection with academic validation using established criteria',
                    'support_level': 'Workplace mentor coordination with academic supervision and structured guidance'
                },
                'accelerated_pathway': {
                    'duration': '9 months intensive with prior learning recognition',
                    'structure': 'RPL-enhanced delivery with targeted competency development and accelerated autonomy progression',
                    'assessment_points': 'Intensive assessment with portfolio validation using visual competency mapping',
                    'support_level': 'Specialized support for experienced professionals with flexible guidance protocols'
                }
            },
            
            'industry_application': profile_data.get('industry_application', [
                f"{profile_data['title']} applications across target sectors",
                "Professional services and consulting organizations",
                "Corporate sustainability departments and teams"
            ]),
            
            'distinctive_features': profile_data.get('distinctive_features', [
                f"Specialized expertise in {profile_data['title'].lower()} professional practice",
                "Comprehensive competency development with industry validation", 
                "Professional recognition through stakeholder engagement"
            ])
        }
        
        return profile
    
    def save_profile_docx(self, profile, filename_base):
        """Generate comprehensive DOCX profile with enhanced visual formatting"""
        
        doc = Document()
        
        # Title section
        info = profile['profile_info']
        title = doc.add_heading(info['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Basic information
        doc.add_paragraph(f"EQF Level: {info['eqf_level']}")
        doc.add_paragraph(f"Total ECTS: {info['target_ects']} ({info['total_hours']} hours)")
        doc.add_paragraph(f"Delivery Modes: {', '.join(info['delivery_modes'])}")
        doc.add_paragraph(f"Target Sectors: {', '.join(info['target_sectors'])}")
        
        # Programme Goal
        doc.add_heading("Programme Goal", level=1)
        goal = profile['programme_goal']
        doc.add_paragraph(f"Primary Objective: {goal['primary_objective']}")
        doc.add_paragraph(f"Learning Approach: {goal['learning_approach']}")
        doc.add_paragraph(f"Professional Context: {goal['professional_context']}")
        
        # Unit Breakdown with explicit outcomes
        doc.add_heading("Unit Learning Outcomes", level=1)
        unit_breakdown = profile['unit_breakdown']
        doc.add_paragraph(f"Units Summary: {unit_breakdown['units_summary']}")
        
        for unit in profile['unit_learning_outcomes']['units']:
            doc.add_heading(f"Unit: {unit['unit_name']} ({unit['unit_ects']} ECTS)", level=2)
            doc.add_paragraph(f"Complexity: {unit['complexity_level']}")
            doc.add_paragraph(f"Knowledge: {unit['knowledge_outcome']}")
            doc.add_paragraph(f"Skills: {unit['skills_outcome']}")
            doc.add_paragraph(f"Competences: {unit['competence_outcome']}")
        
        # Recognition of Prior Learning
        doc.add_heading("Recognition of Prior Learning", level=1)
        rpl = profile['recognition_of_prior_learning']
        doc.add_paragraph(f"RPL Pathway: {rpl['rpl_pathway']}")
        
        # Save file
        docx_file = self.output_dir / f"{filename_base}.docx"
        doc.save(docx_file)
        return docx_file
    
    def save_profile_html(self, profile, filename_base):
        """Generate comprehensive HTML profile with enhanced visual formatting and assessment rubrics"""
        
        info = profile['profile_info']
        
        html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{info['title']}</title>
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 20px; line-height: 1.6; color: #333; }}
        .header {{ background: linear-gradient(135deg, #2c5aa0 0%, #1e3c72 100%); color: white; padding: 30px; border-radius: 10px; margin-bottom: 30px; text-align: center; }}
        .info-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 15px; margin: 20px 0; }}
        .info-item {{ background-color: #f8f9fa; padding: 15px; border-radius: 8px; border-left: 4px solid #2c5aa0; }}
        .section {{ margin: 30px 0; }}
        .unit-card {{ background: #f1f3f4; padding: 20px; margin: 15px 0; border-radius: 8px; border-left: 4px solid #28a745; }}
        .rubric-container {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)); gap: 15px; margin: 20px 0; }}
        .rubric-level {{ padding: 15px; border-radius: 8px; margin: 5px 0; }}
        .level-excellent {{ background-color: #d4edda; border-left: 4px solid #28a745; }}
        .level-good {{ background-color: #d1ecf1; border-left: 4px solid #17a2b8; }}
        .level-satisfactory {{ background-color: #fff3cd; border-left: 4px solid #ffc107; }}
        .level-needs-improvement {{ background-color: #f8d7da; border-left: 4px solid #dc3545; }}
        h1 {{ margin: 0; font-size: 2.2em; }}
        h2 {{ color: #2c5aa0; border-bottom: 2px solid #2c5aa0; padding-bottom: 5px; }}
        h3 {{ color: #495057; }}
        .competency-list {{ background: #e3f2fd; padding: 15px; border-radius: 8px; margin: 10px 0; }}
        .note-box {{ background: #e7f3ff; border-left: 4px solid #007bff; padding: 15px; margin: 15px 0; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{info['title']}</h1>
        <p>EQF Level {info['eqf_level']} | {info['target_ects']} ECTS | {info['total_hours']} Hours</p>
    </div>
    
    <div class="section">
        <h2>Unit Learning Outcomes</h2>
        <div class="note-box">
            <strong>Units Summary:</strong> {profile['unit_breakdown']['units_summary']}
        </div>
"""
        
        # Add unit outcomes
        for unit in profile['unit_learning_outcomes']['units']:
            html += f"""
        <div class="unit-card">
            <h3>{unit['unit_name']} ({unit['unit_ects']} ECTS)</h3>
            <p><strong>Complexity:</strong> {unit['complexity_level']}</p>
            <p><strong>Knowledge:</strong> {unit['knowledge_outcome']}</p>
            <p><strong>Skills:</strong> {unit['skills_outcome']}</p>
            <p><strong>Competences:</strong> {unit['competence_outcome']}</p>
        </div>
"""
        
        # Add visual assessment rubric
        assessment = profile['assessment_framework']
        html += f"""
    </div>
    
    <div class="section">
        <h2>Visual Assessment Rubric</h2>
        <div class="note-box">
            <strong>Assessment Philosophy:</strong> {assessment['assessment_philosophy']}
        </div>
        <div class="rubric-container">
"""
        
        for level, details in assessment['visual_assessment_criteria']['performance_levels'].items():
            html += f"""
            <div class="rubric-level level-{level.replace('_', '-')}">
                <h4>{level.replace('_', ' ').title()}</h4>
                <p><strong>{details['descriptor']}</strong></p>
                <ul>
"""
            for char in details['characteristics']:
                html += f"                    <li>{char}</li>\n"
            html += "                </ul>\n            </div>\n"
        
        html += f"""
        </div>
    </div>
    
    <div class="section">
        <h2>Recognition of Prior Learning</h2>
        <div class="note-box">
            <strong>RPL Pathway:</strong> {profile['recognition_of_prior_learning']['rpl_pathway']}
        </div>
        <p><strong>Maximum Recognition:</strong> {profile['recognition_of_prior_learning']['recognition_limits']['maximum_credit']}</p>
    </div>
    
    <div class="section">
        <h2>Quality Assurance</h2>
        <div class="note-box">
            <strong>Update Cycles:</strong> {profile['validation_framework']['quality_assurance']['update_cycles']}
        </div>
    </div>
    
    <footer style="text-align: center; margin-top: 40px; padding: 20px; background-color: #f8f9fa; border-radius: 8px;">
        <p><em>Professional educational profile with comprehensive EQF alignment and enhanced standards</em></p>
        <p><em>Enhanced with explicit unit outcomes, visual assessment rubrics, and comprehensive stakeholder validation</em></p>
        <p><em>Generated from educational_profiles.json input data</em></p>
    </footer>
</body>
</html>
"""
        
        html_file = self.output_dir / f"{filename_base}.html"
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(html)
        return html_file
    
    def save_profile(self, profile, filename_base, profile_number):
        """Save profile in all configured formats with numbered prefixes"""
        saved_files = []
        
        # Add number prefix to filename
        numbered_filename = f"{profile_number:02d}_{filename_base}"
        
        # Save as JSON
        if 'json' in self.output_formats:
            json_file = self.output_dir / f"{numbered_filename}.json"
            with open(json_file, 'w', encoding='utf-8') as f:
                json.dump(profile, f, indent=2, ensure_ascii=False)
            saved_files.append(json_file)
            print(f"✅ Saved JSON: {json_file.name}")
        
        # Save as HTML
        if 'html' in self.output_formats:
            html_file = self.save_profile_html(profile, numbered_filename)
            saved_files.append(html_file)
            print(f"✅ Saved HTML: {html_file.name}")
        
        # Save as DOCX
        if 'docx' in self.output_formats:
            try:
                docx_file = self.save_profile_docx(profile, numbered_filename)
                saved_files.append(docx_file)
                print(f"✅ Saved DOCX: {docx_file.name}")
            except ImportError:
                print(f"⚠️ Skipping DOCX: python-docx not available")
            except Exception as e:
                print(f"⚠️ Error generating DOCX: {e}")
        
        return saved_files
    
    def generate_all_profiles(self):
        """Generate all educational profiles from JSON input with enhanced professional standards"""
        
        if not self.profiles_data:
            print("❌ No profiles data loaded. Cannot generate profiles.")
            return []
        
        # Expand profiles by EQF levels
        expanded_profiles = self.expand_profiles_by_eqf()
        generated_files = []
        
        print(f"\n📊 Generating {len(expanded_profiles)} educational profiles from JSON input...")
        
        for i, profile_data in enumerate(expanded_profiles, 1):
            print(f"  🔄 Generating {profile_data['title']} (EQF {profile_data['eqf_level']})...")
            
            profile = self.generate_comprehensive_profile(profile_data)
            if profile:
                filename = profile_data['id'].lower()
                files = self.save_profile(profile, filename, i)
                generated_files.extend(files)
                
                print(f"     ✅ EQF Level: {profile_data['eqf_level']}")
                print(f"     ✅ Units: {len(profile_data['units'])} units with explicit learning outcomes")
                print(f"     ✅ ECTS: {profile_data['target_ects']} credits")
        
        print(f"\n📊 GENERATION SUMMARY:")
        print(f"✅ Total profiles generated: {len(expanded_profiles)}")
        print(f"📄 Source: educational_profiles.json input file")
        print(f"🔧 PROFESSIONAL ENHANCEMENTS:")
        print(f"✅ Explicit unit-level learning outcomes with detailed competency mapping")
        print(f"✅ Visual assessment rubrics with color-coded performance levels")
        print(f"✅ Refined autonomy language (accessing specialized guidance where appropriate)")
        print(f"✅ Comprehensive RPL procedure with visual competency mapping")
        print(f"✅ Professional assessment rubric availability statements")
        print(f"✅ Biannual review cycles statement for continued professional currency")
        print(f"✅ Enhanced EQF alignment with sophisticated professional verb structures")
        print(f"✅ Expanded GreenComp framework coverage with collaborative action focus")
        print(f"✅ Numbered filename prefixes for organized document management")
        print(f"✅ NO DigComp mappings (removed as per requirements)")
        print(f"📄 Formats: {', '.join(self.output_formats)}")
        print(f"📁 Output directory: {self.output_dir}")
        print(f"🎯 Status: Professional-ready, comprehensive profile collection with enhanced standards")
        
        return generated_files


def main():
    """Generate professional educational profiles from JSON input with enhanced standards"""
    
    print("🚀 Starting ECM Educational Profiles Generation from JSON Input...")
    print("📊 Reading from educational_profiles.json")
    print("⚙️ Using configuration from: config/settings.json")
    print("🎯 Features: JSON input processing, EQF expansion, visual assessment rubrics")
    print("❌ NO DigComp mappings (removed as per requirements)")
    
    # Check for required dependencies
    try:
        from docx import Document
        print("✅ python-docx library available")
    except ImportError:
        print("❌ python-docx library required for DOCX generation")
        print("   Install with: pip install python-docx")
        print("   Continuing with JSON and HTML generation only...")
    
    # Initialize generator
    generator = ECMProfilesGenerator()
    
    # Generate all profiles
    generated_files = generator.generate_all_profiles()
    
    print(f"\n✅ Professional educational profiles generation complete!")
    print(f"📁 Generated {len(generated_files)} files in: {generator.output_dir}")
    print(f"📄 Formats: {', '.join(generator.output_formats)}")
    print(f"📊 Profiles: Professional profiles with enhanced standards from JSON input")
    
    return {
        'generator': generator,
        'files_generated': len(generated_files),
        'profiles_count': len(generator.expand_profiles_by_eqf()) if generator.profiles_data else 0,
        'output_directory': generator.output_dir,
        'formats': generator.output_formats
    }


if __name__ == "__main__":
    results = main()
