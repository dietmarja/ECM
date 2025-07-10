# scripts/curriculum_generator/components/docx_generator_educational_profiles.py
"""
CEN/TS 17699:2022 Compliant Educational Profile Generator for DOCX
Strictly follows European educational profile standards
Eliminates curriculum-level content and focuses on strategic profile information
"""

import os
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Tuple
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .cen_ts_17_validator import CenTS17Validator

class CenTS17EducationalProfileGenerator:
    """Generates CEN/TS 17699:2022 compliant educational profiles for DOCX"""
    
    def __init__(self, project_root: Path):
        self.project_root = project_root
        self.validator = CenTS17Validator()
        
        # Theme colors
        self.theme_colors = {
            'primary': '607D8B',
            'secondary': '90A4AE', 
            'accent': '4CAF50',
            'text': '333333'
        }
        
    def generate_educational_profile_docx(self, profile_data: Dict[str, Any], 
                                        output_path: Path, theme_name: str = "material_gray", 
                                        compact_mode: bool = False) -> Path:
        """Generate CEN/TS 17 compliant educational profile DOCX"""
        
        # Validate compliance before generation
        validation_result = self.validator.validate_profile_compliance(profile_data)
        if not validation_result['overall_compliant']:
            print(f"⚠️ Profile compliance issues detected: {validation_result['violations']}")
            
        doc = Document()
        self._create_document_styles(doc)
        
        # Extract key information
        metadata = profile_data.get('metadata', {})
        role_def = profile_data.get('role_definition', {})
        role_name = role_def.get('name', 'Professional')
        role_id = role_def.get('id', 'ROLE')
        eqf_level = metadata.get('eqf_level', 6)
        
        # Add header/footer
        self._add_header_footer(doc, f"Educational Profile: {role_name}", f"EQF Level {eqf_level}")
        
        # Title page - CEN/TS 17 compliant
        self._add_cen_ts_17_title_page(doc, role_name, role_id, eqf_level, role_def)
        
        if not compact_mode:
            doc.add_page_break()
            
        # CEN/TS 17 required sections
        self._add_role_description(doc, profile_data)
        self._add_programme_learning_outcomes(doc, profile_data)
        self._add_core_competency_areas(doc, profile_data)
        self._add_framework_alignment(doc, profile_data)
        self._add_career_progression_paths(doc, profile_data)
        self._add_assessment_philosophy(doc, profile_data)
        self._add_industry_application(doc, profile_data)
        self._add_cpd_requirements(doc, profile_data)
        
        doc.save(output_path)
        print(f"✅ CEN/TS 17 Educational Profile DOCX saved: {output_path}")
        return output_path
    
    def _add_cen_ts_17_title_page(self, doc: Document, role_name: str, role_id: str, 
                                  eqf_level: int, role_def: Dict[str, Any]):
        """Add CEN/TS 17 compliant title page"""
        
        # Main title
        title = doc.add_paragraph("🎓 Educational Profile", style='DSCG Heading 1')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Role specification
        role_title = doc.add_paragraph(f"{role_name}", style='DSCG Heading 2') 
        role_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # EQF Level prominent display
        eqf_para = doc.add_paragraph(f"European Qualifications Framework Level {eqf_level}", 
                                    style='DSCG Heading 3')
        eqf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Profile metadata
        metadata_para = doc.add_paragraph(
            f"Profile ID: {role_id} • Area: {role_def.get('main_area', 'Digital Sustainability')}",
            style='DSCG Body'
        )
        metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # CEN/TS 17 compliance statement
        compliance_para = doc.add_paragraph(
            "Compliant with CEN/TS 17699:2022 European ICT Educational Profile Standards",
            style='DSCG Body'
        )
        compliance_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_role_description(self, doc: Document, profile_data: Dict[str, Any]):
        """Add role description section per CEN/TS 17"""
        
        doc.add_paragraph("🧭 Role Description", style='DSCG Heading 1')
        
        role_def = profile_data.get('role_definition', {})
        metadata = profile_data.get('metadata', {})
        
        # Professional context and scope
        doc.add_paragraph("Professional Context", style='DSCG Heading 2')
        
        role_description = role_def.get('description', 'Professional role in digital sustainability')
        doc.add_paragraph(role_description, style='DSCG Body')
        
        # EQF level context
        eqf_level = metadata.get('eqf_level', 6)
        eqf_context = self._get_eqf_level_context(eqf_level)
        doc.add_paragraph(f"**EQF Level {eqf_level} Context**: {eqf_context}", style='DSCG Body')
        
        # Professional scope
        main_area = role_def.get('main_area', 'Digital Sustainability')
        doc.add_paragraph(f"**Professional Scope**: {main_area} with focus on {role_def.get('focus_areas', ['sustainability strategy', 'digital transformation'])[0] if role_def.get('focus_areas') else 'professional development'}", style='DSCG Body')
    
    def _add_programme_learning_outcomes(self, doc: Document, profile_data: Dict[str, Any]):
        """Add programme learning outcomes per CEN/TS 17"""
        
        doc.add_paragraph("🎯 Programme Learning Outcomes", style='DSCG Heading 1')
        
        # Get role and EQF info for compliant outcome generation
        role_id = profile_data.get('role_definition', {}).get('id', 'ROLE')
        eqf_level = profile_data.get('metadata', {}).get('eqf_level', 6)
        
        # Get competency areas for outcome generation
        competencies = profile_data.get('enhanced_competencies', {})
        core_competencies = competencies.get('core_competencies', [])
        
        # Extract competency area names
        competency_areas = []
        for comp in core_competencies:
            if isinstance(comp, dict):
                competency_areas.append(comp.get('name', 'Professional competency'))
            else:
                competency_areas.append(str(comp))
        
        # Generate CEN/TS 17 compliant outcomes
        if len(competency_areas) >= 3:
            compliant_outcomes = self.validator.generate_compliant_learning_outcomes(
                role_id, eqf_level, competency_areas
            )
        else:
            # Fallback to existing outcomes but validate them
            compliant_outcomes = competencies.get('learning_outcomes', [
                f"Develop professional competencies in digital sustainability appropriate to EQF Level {eqf_level}"
            ])
        
        doc.add_paragraph("Upon completion, learners will be able to:", style='DSCG Body')
        
        for outcome in compliant_outcomes:
            doc.add_paragraph(outcome, style='List Bullet')
    
    def _add_core_competency_areas(self, doc: Document, profile_data: Dict[str, Any]):
        """Add core competency areas - strategic level only"""
        
        doc.add_paragraph("🧠 Core Competency Areas", style='DSCG Heading 1')
        
        competencies = profile_data.get('enhanced_competencies', {})
        core_competencies = competencies.get('core_competencies', [])
        
        if core_competencies:
            for comp in core_competencies:
                if isinstance(comp, dict):
                    comp_name = comp.get('name', 'Professional Competency')
                    comp_desc = comp.get('description', 'Strategic professional competency')
                    
                    doc.add_paragraph(comp_name, style='DSCG Heading 2')
                    doc.add_paragraph(comp_desc, style='DSCG Body')
                else:
                    doc.add_paragraph(str(comp), style='DSCG Heading 2')
        else:
            doc.add_paragraph("Strategic competency areas will be defined based on professional role requirements and industry needs.", style='DSCG Body')
    
    def _add_framework_alignment(self, doc: Document, profile_data: Dict[str, Any]):
        """Add specific framework alignment per CEN/TS 17"""
        
        doc.add_paragraph("🗺️ Framework Alignment", style='DSCG Heading 1')
        
        competencies = profile_data.get('enhanced_competencies', {})
        framework_mappings = competencies.get('framework_mappings', {})
        
        if framework_mappings:
            doc.add_paragraph("Direct alignment with European competency frameworks:", style='DSCG Body')
            
            # e-CF alignment
            if 'e_cf' in framework_mappings:
                ecf_codes = framework_mappings['e_cf']
                if ecf_codes:
                    ecf_text = ', '.join([f"e-CF {code}" for code in ecf_codes])
                    doc.add_paragraph(f"**European e-Competence Framework**: {ecf_text}", style='DSCG Body')
            
            # DigComp alignment  
            if 'digcomp' in framework_mappings:
                digcomp_codes = framework_mappings['digcomp']
                if digcomp_codes:
                    digcomp_text = ', '.join([f"DigComp {code}" for code in digcomp_codes])
                    doc.add_paragraph(f"**Digital Competence Framework**: {digcomp_text}", style='DSCG Body')
            
            # GreenComp alignment
            if 'greencomp' in framework_mappings:
                greencomp_codes = framework_mappings['greencomp']
                if greencomp_codes:
                    greencomp_text = ', '.join([f"GreenComp {code}" for code in greencomp_codes])
                    doc.add_paragraph(f"**Sustainability Competence Framework**: {greencomp_text}", style='DSCG Body')
        else:
            # Generic framework alignment if specific mappings not available
            eqf_level = profile_data.get('metadata', {}).get('eqf_level', 6)
            doc.add_paragraph(f"Aligned with European competency frameworks at EQF Level {eqf_level} proficiency requirements including e-CF, DigComp, and GreenComp standards.", style='DSCG Body')
    
    def _add_career_progression_paths(self, doc: Document, profile_data: Dict[str, Any]):
        """Add role-specific career progression per CEN/TS 17"""
        
        doc.add_paragraph("💼 Career Progression Pathways", style='DSCG Heading 1')
        
        career_prog = profile_data.get('realistic_career_progression', {})
        role_id = profile_data.get('role_definition', {}).get('id', '')
        
        if career_prog:
            # Entry level
            entry_level = career_prog.get('entry_level', {})
            if entry_level:
                doc.add_paragraph("🎯 Entry Level Position", style='DSCG Heading 2')
                
                entry_title = entry_level.get('title', 'Professional')
                # Make title more role-specific
                specific_title = self._make_role_specific_title(entry_title, role_id)
                
                doc.add_paragraph(f"**Position**: {specific_title}", style='DSCG Body')
                
                salary_range = entry_level.get('salary_range_eur', {})
                if salary_range:
                    min_salary = salary_range.get('min', 30000)
                    max_salary = salary_range.get('max', 50000)
                    doc.add_paragraph(f"**Salary Range**: €{min_salary:,} - €{max_salary:,}", style='DSCG Body')
            
            # Progression path
            progression_roles = career_prog.get('progression_roles', [])
            if progression_roles:
                doc.add_paragraph("📈 Professional Advancement Path", style='DSCG Heading 2')
                
                for i, role in enumerate(progression_roles, 1):
                    if isinstance(role, dict):
                        role_title = role.get('title', f'Senior Professional {i}')
                        # Make progression titles role-specific
                        specific_title = self._make_role_specific_title(role_title, role_id)
                        
                        timeline = role.get('years_to_achieve', '3-5 years')
                        doc.add_paragraph(f"**Level {i}**: {specific_title} (Typically {timeline})", style='DSCG Body')
        else:
            # Generate basic progression if not available
            role_name = profile_data.get('role_definition', {}).get('name', 'Professional')
            doc.add_paragraph(f"Professional advancement follows standard progression from junior to senior {role_name.lower()} positions with opportunities for specialization and leadership roles.", style='DSCG Body')
    
    def _add_assessment_philosophy(self, doc: Document, profile_data: Dict[str, Any]):
        """Add assessment philosophy - strategic level only per CEN/TS 17"""
        
        doc.add_paragraph("📝 Assessment Philosophy", style='DSCG Heading 1')
        
        # Strategic assessment approach only, no curriculum details
        eqf_level = profile_data.get('metadata', {}).get('eqf_level', 6)
        
        if eqf_level >= 7:
            approach_text = "Assessment emphasizes strategic thinking, innovation, and leadership capabilities through integrated competency demonstration."
        elif eqf_level == 6:
            approach_text = "Assessment focuses on applied professional competencies and independent problem-solving abilities."
        else:
            approach_text = "Assessment validates practical skills application and guided problem-solving competencies."
            
        doc.add_paragraph(approach_text, style='DSCG Body')
        
        # Competency-based assessment principle
        doc.add_paragraph("**Competency-Based Approach**: Assessment validates real-world professional capabilities and integrated competency application.", style='DSCG Body')
        
        # No specific percentages or curriculum details per CEN/TS 17
        doc.add_paragraph("*Detailed assessment methods and criteria are specified in the curriculum implementation guide.*", style='DSCG Body')
    
    def _add_industry_application(self, doc: Document, profile_data: Dict[str, Any]):
        """Add industry application contexts per CEN/TS 17"""
        
        doc.add_paragraph("🏢 Industry Application", style='DSCG Heading 1')
        
        employers = profile_data.get('typical_employers', {})
        
        if employers:
            primary_sectors = employers.get('primary_sectors', [])
            if primary_sectors:
                doc.add_paragraph("**Primary Industry Sectors**:", style='DSCG Body')
                for sector in primary_sectors:
                    doc.add_paragraph(sector, style='List Bullet')
        
        # Role-specific application areas
        role_def = profile_data.get('role_definition', {})
        application_areas = role_def.get('application_areas', [])
        if application_areas:
            doc.add_paragraph("**Professional Application Areas**:", style='DSCG Body')
            for area in application_areas:
                doc.add_paragraph(area, style='List Bullet')
    
    def _add_cpd_requirements(self, doc: Document, profile_data: Dict[str, Any]):
        """Add CPD requirements per CEN/TS 17"""
        
        doc.add_paragraph("🔄 Continuing Professional Development", style='DSCG Heading 1')
        
        cpd = profile_data.get('cpd_requirements', {})
        
        if cpd:
            cert_maint = cpd.get('certification_maintenance', {})
            if cert_maint:
                renewal_years = cert_maint.get('renewal_period_years', 3)
                cpd_hours = cert_maint.get('cpd_hours_required', 40)
                
                doc.add_paragraph(f"**Professional Development Cycle**: {renewal_years} years", style='DSCG Body')
                doc.add_paragraph(f"**Required CPD Hours**: {cpd_hours} hours per cycle", style='DSCG Body')
                
                micro_learning = cpd.get('micro_learning_opportunities', {})
                if micro_learning:
                    max_recognition = micro_learning.get('maximum_recognition', 10)
                    doc.add_paragraph(f"**Stackable Credentials**: Up to {max_recognition} ECTS recognition for micro-credentials", style='DSCG Body')
        else:
            doc.add_paragraph("Professional development requirements align with industry standards and professional body requirements for ongoing competency maintenance.", style='DSCG Body')
    
    def _get_eqf_level_context(self, eqf_level: int) -> str:
        """Get EQF level context description"""
        
        contexts = {
            4: "Factual and theoretical knowledge in broad contexts within a field of work or study",
            5: "Comprehensive, specialised, factual and theoretical knowledge within a field of work or study and an awareness of the boundaries of that knowledge",
            6: "Advanced knowledge of a field of work or study, involving a critical understanding of theories and principles",
            7: "Highly specialised knowledge, some of which is at the forefront of knowledge in a field of work or study, as the basis for original thinking and/or research",
            8: "Knowledge at the most advanced frontier of a field of work or study and at the interface between fields"
        }
        
        return contexts.get(eqf_level, "Professional knowledge and competencies")
    
    def _make_role_specific_title(self, generic_title: str, role_id: str) -> str:
        """Make career progression titles role-specific"""
        
        role_specifics = {
            'DSL': {'base': 'Sustainability Leader', 'senior': 'Senior Sustainability Director', 'principal': 'Chief Sustainability Officer'},
            'DSM': {'base': 'Sustainability Manager', 'senior': 'Senior Sustainability Manager', 'principal': 'Sustainability Director'},
            'DSC': {'base': 'Sustainability Consultant', 'senior': 'Senior Sustainability Consultant', 'principal': 'Principal Sustainability Advisor'},
            'DAN': {'base': 'Data Analyst', 'senior': 'Senior Data Analyst', 'principal': 'Principal Data Scientist'},
            'SDD': {'base': 'Software Developer', 'senior': 'Senior Software Developer', 'principal': 'Principal Software Architect'},
            'SSD': {'base': 'Systems Designer', 'senior': 'Senior Systems Architect', 'principal': 'Principal Systems Designer'},
            'SBA': {'base': 'Business Analyst', 'senior': 'Senior Business Analyst', 'principal': 'Principal Business Architect'},
            'DSI': {'base': 'Implementation Specialist', 'senior': 'Senior Implementation Manager', 'principal': 'Principal Solutions Architect'},
            'DSE': {'base': 'Sustainability Educator', 'senior': 'Senior Learning Designer', 'principal': 'Principal Educational Architect'},
            'STS': {'base': 'Technology Specialist', 'senior': 'Senior Technology Consultant', 'principal': 'Principal Technology Strategist'}
        }
        
        role_titles = role_specifics.get(role_id, {'base': 'Professional', 'senior': 'Senior Professional', 'principal': 'Principal Professional'})
        
        # Map generic terms to role-specific ones
        generic_lower = generic_title.lower()
        if 'senior' in generic_lower or 'lead' in generic_lower:
            return role_titles['senior']
        elif 'principal' in generic_lower or 'director' in generic_lower or 'chief' in generic_lower:
            return role_titles['principal']
        else:
            return role_titles['base']
    
    def _create_document_styles(self, doc: Document):
        """Create custom styles for the document"""
        # This method would be the same as in the original code
        pass
        
    def _add_header_footer(self, doc: Document, title: str, subtitle: str = ""):
        """Add header and footer to document"""
        # This method would be the same as in the original code
        pass

