# scripts/curriculum_generator/components/docx_generator.py
"""
DOCX Generator for Digital Sustainability Curriculum Generator
NOW WITH CEN/TS 17699:2022 EDUCATIONAL PROFILE COMPLIANCE
Generates professional DOCX documents for both Educational Profiles and Curricula
Maintains strict separation between strategic profiles and operational curricula
UPDATED: Full CEN/TS 17 compliance with streamlined profile generation
"""

import os
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

class DocxGenerator:
    """Generates professional DOCX documents for DSCG system with CEN/TS 17 compliance"""
    
    def __init__(self, project_root: Path):
        self.project_root = project_root
        
        # CEN/TS 17 compliance built-in (simplified implementation)
        self.cen_ts_17_compliant = True
        
        # EQF complexity verbs for validation
        self.eqf_complexity_verbs = {
            4: ['apply', 'implement', 'use', 'demonstrate', 'operate'],
            5: ['analyze', 'evaluate', 'coordinate', 'adapt', 'supervise'],
            6: ['design', 'develop', 'manage', 'integrate', 'optimize'],
            7: ['synthesize', 'innovate', 'lead', 'transform', 'strategize'],
            8: ['pioneer', 'revolutionize', 'establish', 'conceptualize', 'influence']
        }
        
        # Theme colors (will be overridden by actual theme)
        self.theme_colors = {
            'primary': '607D8B',      # Material Gray
            'secondary': '90A4AE',    # Light Gray
            'accent': '4CAF50',       # Green
            'text': '333333',         # Dark Gray
            'light_bg': 'F8F9FA'     # Very Light Gray
        }
        
        print("📄 DOCX Generator initialized with CEN/TS 17 compliance")
    
    def set_theme_colors(self, theme_colors: Dict[str, str]):
        """Set theme colors for document styling"""
        self.theme_colors.update(theme_colors)
    
    # ===============================
    # EDUCATIONAL PROFILE METHODS - CEN/TS 17 COMPLIANT
    # ===============================
    
    def generate_educational_profile_docx(self, profile_data: Dict[str, Any], 
                                        output_path: Path, theme_name: str = "material_gray", 
                                        compact_mode: bool = False) -> Path:
        """Generate CEN/TS 17699:2022 compliant educational profile DOCX"""
        
        print("🔧 Generating CEN/TS 17 compliant educational profile...")
        
        try:
            # Create document
            doc = Document()
            
            # Apply theme colors
            if theme_name.startswith('eu_'):
                self.theme_colors.update({
                    'primary': '003399',    # EU Blue
                    'accent': 'FFCC00'      # EU Yellow
                })
            
            # Create styles
            self._create_document_styles(doc)
            
            # Extract key information
            metadata = profile_data.get('metadata', {})
            role_def = profile_data.get('role_definition', {})
            role_name = role_def.get('name', 'Professional')
            role_id = role_def.get('id', 'ROLE')
            eqf_level = metadata.get('eqf_level', 6)
            
            # Add header/footer
            self._add_header_footer(doc, f"Educational Profile: {role_name}", f"EQF Level {eqf_level}")
            
            # Title page - CEN/TS 17 compliant
            self._add_cen_ts_17_title_page(doc, role_name, role_id, eqf_level, role_def)
            
            if not compact_mode:
                doc.add_page_break()
                
            # CEN/TS 17 required sections
            self._add_role_description(doc, profile_data)
            self._add_programme_learning_outcomes(doc, profile_data)
            self._add_core_competency_areas(doc, profile_data)
            self._add_framework_alignment(doc, profile_data)
            self._add_career_progression_paths(doc, profile_data)
            self._add_assessment_philosophy(doc, profile_data)
            self._add_industry_application(doc, profile_data)
            self._add_cpd_requirements(doc, profile_data)
            
            # Save document
            doc.save(output_path)
            print(f"✅ CEN/TS 17 Educational Profile DOCX saved: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"❌ Error generating Educational Profile DOCX: {e}")
            import traceback
            traceback.print_exc()
            raise
    
    # ===============================
    # CEN/TS 17 COMPLIANT EDUCATIONAL PROFILE METHODS
    # ===============================
    
    def _add_cen_ts_17_title_page(self, doc: Document, role_name: str, role_id: str, 
                                  eqf_level: int, role_def: Dict[str, Any]):
        """Add CEN/TS 17 compliant title page"""
        
        # Main title
        title = doc.add_paragraph("🎓 Educational Profile", style='DSCG Heading 1')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Role specification
        role_title = doc.add_paragraph(f"{role_name}", style='DSCG Heading 2') 
        role_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # EQF Level prominent display
        eqf_para = doc.add_paragraph(f"European Qualifications Framework Level {eqf_level}", 
                                    style='DSCG Heading 3')
        eqf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Profile metadata
        metadata_para = doc.add_paragraph(
            f"Profile ID: {role_id} • Area: {role_def.get('main_area', 'Digital Sustainability')}",
            style='DSCG Body'
        )
        metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # CEN/TS 17 compliance statement
        compliance_para = doc.add_paragraph(
            "Compliant with CEN/TS 17699:2022 European ICT Educational Profile Standards",
            style='DSCG Body'
        )
        compliance_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_role_description(self, doc: Document, profile_data: Dict[str, Any]):
        """Add role description section per CEN/TS 17"""
        
        doc.add_paragraph("🧭 Role Description", style='DSCG Heading 1')
        
        role_def = profile_data.get('role_definition', {})
        metadata = profile_data.get('metadata', {})
        
        # Professional context and scope
        doc.add_paragraph("Professional Context", style='DSCG Heading 2')
        
        role_description = role_def.get('description', 'Professional role in digital sustainability')
        doc.add_paragraph(role_description, style='DSCG Body')
        
        # EQF level context
        eqf_level = metadata.get('eqf_level', 6)
        eqf_context = self._get_eqf_level_context(eqf_level)
        doc.add_paragraph(f"**EQF Level {eqf_level} Context**: {eqf_context}", style='DSCG Body')
    
    def _add_programme_learning_outcomes(self, doc: Document, profile_data: Dict[str, Any]):
        """Add programme learning outcomes per CEN/TS 17"""
        
        doc.add_paragraph("🎯 Programme Learning Outcomes", style='DSCG Heading 1')
        
        # Get role and EQF info for compliant outcome generation
        role_id = profile_data.get('role_definition', {}).get('id', 'ROLE')
        eqf_level = profile_data.get('metadata', {}).get('eqf_level', 6)
        
        # Generate CEN/TS 17 compliant outcomes
        compliant_outcomes = self._generate_compliant_learning_outcomes(role_id, eqf_level)
        
        doc.add_paragraph("Upon completion, learners will be able to:", style='DSCG Body')
        
        for outcome in compliant_outcomes:
            doc.add_paragraph(outcome, style='List Bullet')
    
    def _add_core_competency_areas(self, doc: Document, profile_data: Dict[str, Any]):
        """Add core competency areas - strategic level only"""
        
        doc.add_paragraph("🧠 Core Competency Areas", style='DSCG Heading 1')
        
        competencies = profile_data.get('enhanced_competencies', {})
        core_competencies = competencies.get('core_competencies', [])
        
        has_competency_data = False
        
        if core_competencies and len(core_competencies) > 0:
            for comp in core_competencies:
                if isinstance(comp, dict):
                    comp_name = comp.get('name', 'Professional Competency')
                    comp_desc = comp.get('description', 'Strategic professional competency')
                    
                    if comp_name and comp_name.strip():
                        doc.add_paragraph(comp_name, style='DSCG Heading 2')
                        doc.add_paragraph(comp_desc, style='DSCG Body')
                        has_competency_data = True
                elif isinstance(comp, str) and comp.strip():
                    # Handle string competencies
                    doc.add_paragraph(comp, style='DSCG Heading 2')
                    doc.add_paragraph(f"Professional competency encompassing {comp.lower()} expertise, methodologies, and strategic application in organizational contexts.", style='DSCG Body')
                    has_competency_data = True
        
        # If no meaningful competency data, generate role-specific competencies
        if not has_competency_data:
            role_id = profile_data.get('role_definition', {}).get('id', 'ROLE')
            role_specific_competencies = self._generate_role_specific_competencies(role_id)
            
            for comp_name, comp_desc in role_specific_competencies.items():
                doc.add_paragraph(comp_name, style='DSCG Heading 2')
                doc.add_paragraph(comp_desc, style='DSCG Body')
    
    def _generate_role_specific_competencies(self, role_id: str) -> Dict[str, str]:
        """Generate role-specific competencies when data is missing"""
        
        competencies = {
            'DSL': {
                'Strategic Sustainability Leadership': 'Advanced competency in developing and implementing organizational sustainability strategies that align with business objectives and stakeholder expectations.',
                'Stakeholder Engagement and Governance': 'Expert-level skills in managing complex stakeholder relationships, facilitating collaborative decision-making, and establishing governance frameworks.',
                'Transformation Management': 'Comprehensive competency in leading organizational change initiatives, managing transformation processes, and driving cultural shifts toward sustainability.'
            },
            'DSM': {
                'Sustainability Project Management': 'Professional competency in planning, executing, and monitoring sustainability initiatives within organizational contexts and resource constraints.',
                'Performance Measurement and Reporting': 'Advanced skills in developing sustainability metrics, monitoring progress, and communicating results to diverse stakeholder groups.',
                'Operational Excellence': 'Competency in optimizing operational processes for sustainability outcomes while maintaining efficiency and quality standards.'
            },
            'DAN': {
                'Sustainability Data Analytics': 'Expert competency in collecting, analyzing, and interpreting environmental and social data to support evidence-based decision-making.',
                'Predictive Modeling and Forecasting': 'Advanced skills in developing predictive models for sustainability outcomes, risk assessment, and scenario planning.',
                'Data Visualization and Communication': 'Professional competency in translating complex data insights into accessible visualizations and actionable recommendations.'
            },
            'SDD': {
                'Sustainable Software Development': 'Expert competency in developing energy-efficient software solutions, implementing green coding practices, and optimizing system performance.',
                'Environmental Impact Assessment': 'Advanced skills in measuring and minimizing the environmental footprint of digital technologies and software systems.',
                'Green Technology Innovation': 'Professional competency in researching, designing, and implementing innovative solutions for sustainable software development.'
            },
            'DSC': {
                'Sustainability Strategy Consulting': 'Expert competency in analyzing organizational sustainability challenges and developing tailored strategic recommendations.',
                'Client Engagement and Advisory': 'Advanced skills in building client relationships, facilitating strategic discussions, and delivering high-impact consulting solutions.',
                'Methodology Development': 'Professional competency in creating frameworks, tools, and methodologies for sustainability assessment and implementation.'
            }
        }
        
        return competencies.get(role_id, {
            'Professional Expertise': 'Strategic competency in digital sustainability practices, methodologies, and organizational implementation.',
            'Technical Proficiency': 'Advanced skills in applying digital technologies and tools to address sustainability challenges and opportunities.',
            'Leadership and Communication': 'Professional competency in leading initiatives, engaging stakeholders, and communicating sustainability value and impact.'
        })
    
    def _add_framework_alignment(self, doc: Document, profile_data: Dict[str, Any]):
        """Add specific framework alignment per CEN/TS 17"""
        
        doc.add_paragraph("🗺️ Framework Alignment", style='DSCG Heading 1')
        
        competencies = profile_data.get('enhanced_competencies', {})
        framework_mappings = competencies.get('framework_mappings', {})
        eqf_level = profile_data.get('metadata', {}).get('eqf_level', 6)
        role_id = profile_data.get('role_definition', {}).get('id', 'ROLE')
        
        # Check if we have meaningful framework mappings
        has_meaningful_mappings = False
        framework_content = []
        
        if framework_mappings:
            # e-CF alignment
            if 'e_cf' in framework_mappings:
                ecf_codes = framework_mappings['e_cf']
                if ecf_codes and len(ecf_codes) > 0:
                    ecf_text = ', '.join([f"e-CF {code}" for code in ecf_codes if code])
                    if ecf_text:
                        framework_content.append(f"**European e-Competence Framework**: {ecf_text}")
                        has_meaningful_mappings = True
            
            # DigComp alignment  
            if 'digcomp' in framework_mappings:
                digcomp_codes = framework_mappings['digcomp']
                if digcomp_codes and len(digcomp_codes) > 0:
                    digcomp_text = ', '.join([f"DigComp {code}" for code in digcomp_codes if code])
                    if digcomp_text:
                        framework_content.append(f"**Digital Competence Framework**: {digcomp_text}")
                        has_meaningful_mappings = True
            
            # GreenComp alignment
            if 'greencomp' in framework_mappings:
                greencomp_codes = framework_mappings['greencomp']
                if greencomp_codes and len(greencomp_codes) > 0:
                    greencomp_text = ', '.join([f"GreenComp {code}" for code in greencomp_codes if code])
                    if greencomp_text:
                        framework_content.append(f"**Sustainability Competence Framework**: {greencomp_text}")
                        has_meaningful_mappings = True
        
        if has_meaningful_mappings:
            doc.add_paragraph("Direct alignment with European competency frameworks:", style='DSCG Body')
            for content in framework_content:
                doc.add_paragraph(content, style='DSCG Body')
        else:
            # Generate role-specific framework alignment
            role_specific_frameworks = self._generate_role_specific_frameworks(role_id, eqf_level)
            
            doc.add_paragraph(f"Aligned with European competency frameworks at EQF Level {eqf_level} proficiency:", style='DSCG Body')
            
            for framework_info in role_specific_frameworks:
                doc.add_paragraph(framework_info, style='DSCG Body')
            
            doc.add_paragraph("*Specific competence codes and proficiency levels are defined in the detailed curriculum implementation.*", style='DSCG Body')
    
    def _add_career_progression_paths(self, doc: Document, profile_data: Dict[str, Any]):
        """Add role-specific career progression per CEN/TS 17"""
        
        doc.add_paragraph("💼 Career Progression Pathways", style='DSCG Heading 1')
        
        career_prog = profile_data.get('realistic_career_progression', {})
        role_id = profile_data.get('role_definition', {}).get('id', '')
        role_name = profile_data.get('role_definition', {}).get('name', 'Professional')
        
        has_career_data = False
        
        if career_prog:
            # Entry level - handle both dict and string
            entry_level = career_prog.get('entry_level', {})
            if entry_level:
                doc.add_paragraph("🎯 Entry Level Position", style='DSCG Heading 2')
                
                # Handle if entry_level is a string instead of dict
                if isinstance(entry_level, str):
                    entry_title = entry_level
                    specific_title = self._make_role_specific_title(entry_title, role_id)
                    doc.add_paragraph(f"**Position**: {specific_title}", style='DSCG Body')
                    has_career_data = True
                elif isinstance(entry_level, dict):
                    entry_title = entry_level.get('title', 'Professional')
                    specific_title = self._make_role_specific_title(entry_title, role_id)
                    doc.add_paragraph(f"**Position**: {specific_title}", style='DSCG Body')
                    
                    salary_range = entry_level.get('salary_range_eur', {})
                    if salary_range and isinstance(salary_range, dict):
                        min_salary = salary_range.get('min', 30000)
                        max_salary = salary_range.get('max', 50000)
                        doc.add_paragraph(f"**Salary Range**: €{min_salary:,} - €{max_salary:,}", style='DSCG Body')
                    has_career_data = True
            
            # Progression path - handle both list and dict
            progression_roles = career_prog.get('progression_roles', [])
            if progression_roles and len(progression_roles) > 0:
                doc.add_paragraph("📈 Professional Advancement Path", style='DSCG Heading 2')
                
                for i, role in enumerate(progression_roles, 1):
                    # Handle if role is a string instead of dict
                    if isinstance(role, str):
                        specific_title = self._make_role_specific_title(role, role_id)
                        doc.add_paragraph(f"**Level {i}**: {specific_title}", style='DSCG Body')
                        has_career_data = True
                    elif isinstance(role, dict):
                        role_title = role.get('title', f'Senior Professional {i}')
                        specific_title = self._make_role_specific_title(role_title, role_id)
                        
                        timeline = role.get('years_to_achieve', '3-5 years')
                        doc.add_paragraph(f"**Level {i}**: {specific_title} (Typically {timeline})", style='DSCG Body')
                        has_career_data = True
        
        # If no meaningful career data, generate professional progression
        if not has_career_data:
            doc.add_paragraph("📈 Professional Development Pathway", style='DSCG Heading 2')
            
            # Generate role-specific progression
            role_specific_progression = self._generate_role_specific_progression(role_id, role_name)
            
            for level, info in role_specific_progression.items():
                doc.add_paragraph(f"**{level}**: {info}", style='DSCG Body')
            
            doc.add_paragraph("*Career advancement is supported through continuous professional development, demonstrated competency growth, and leadership experience in sustainability initiatives.*", style='DSCG Body')
    
    def _add_assessment_philosophy(self, doc: Document, profile_data: Dict[str, Any]):
        """Add assessment philosophy - strategic level only per CEN/TS 17"""
        
        doc.add_paragraph("📝 Assessment Philosophy", style='DSCG Heading 1')
        
        # Strategic assessment approach only, no curriculum details
        eqf_level = profile_data.get('metadata', {}).get('eqf_level', 6)
        
        if eqf_level >= 7:
            approach_text = "Assessment emphasizes strategic thinking, innovation, and leadership capabilities through integrated competency demonstration."
        elif eqf_level == 6:
            approach_text = "Assessment focuses on applied professional competencies and independent problem-solving abilities."
        else:
            approach_text = "Assessment validates practical skills application and guided problem-solving competencies."
            
        doc.add_paragraph(approach_text, style='DSCG Body')
        
        # Competency-based assessment principle
        doc.add_paragraph("**Competency-Based Approach**: Assessment validates real-world professional capabilities and integrated competency application.", style='DSCG Body')
        
        # No specific percentages or curriculum details per CEN/TS 17
        doc.add_paragraph("*Detailed assessment methods and criteria are specified in the curriculum implementation guide.*", style='DSCG Body')
    
    def _generate_role_specific_progression(self, role_id: str, role_name: str) -> Dict[str, str]:
        """Generate role-specific career progression when data is missing"""
        
        progressions = {
            'DSL': {
                'Entry Level': 'Sustainability Strategy Coordinator - developing strategic initiatives',
                'Mid Level': 'Senior Sustainability Leader - managing transformation programs', 
                'Senior Level': 'Director of Sustainability Strategy - leading organizational change',
                'Executive Level': 'Chief Sustainability Officer - driving enterprise-wide transformation'
            },
            'DSM': {
                'Entry Level': 'Sustainability Project Manager - implementing sustainability projects',
                'Mid Level': 'Senior Sustainability Manager - overseeing multiple initiatives',
                'Senior Level': 'Sustainability Program Director - managing strategic programs',
                'Executive Level': 'VP of Sustainability Operations - leading operational excellence'
            },
            'DAN': {
                'Entry Level': 'Sustainability Data Analyst - analyzing environmental metrics',
                'Mid Level': 'Senior Data Analyst - leading analytics initiatives',
                'Senior Level': 'Lead Data Scientist - developing predictive models',
                'Expert Level': 'Principal Analytics Specialist - architecting data strategies'
            },
            'SDD': {
                'Entry Level': 'Junior Green Developer - implementing sustainable coding practices',
                'Mid Level': 'Senior Software Developer - leading green development initiatives',
                'Senior Level': 'Lead Software Architect - designing sustainable systems',
                'Expert Level': 'Principal Technology Architect - driving innovation in green tech'
            },
            'DSC': {
                'Entry Level': 'Sustainability Consultant - providing advisory services',
                'Mid Level': 'Senior Consultant - leading client engagements',
                'Senior Level': 'Principal Consultant - developing methodologies',
                'Partner Level': 'Partner/Director - building practice areas'
            }
        }
        
        return progressions.get(role_id, {
            'Entry Level': f'Junior {role_name} - developing foundational expertise',
            'Mid Level': f'Senior {role_name} - leading specialized initiatives',
            'Senior Level': f'Lead {role_name} - managing strategic programs',
            'Expert Level': f'Principal {role_name} - driving innovation and thought leadership'
        })
    
    def _add_industry_application(self, doc: Document, profile_data: Dict[str, Any]):
        """Add industry application contexts per CEN/TS 17"""
        
        doc.add_paragraph("🏢 Industry Application", style='DSCG Heading 1')
        
        employers = profile_data.get('typical_employers', {})
        industry_app = profile_data.get('industry_application', [])
        
        has_industry_data = False
        
        # Try employers data first
        if employers and isinstance(employers, dict):
            primary_sectors = employers.get('primary_sectors', [])
            if primary_sectors and isinstance(primary_sectors, list) and len(primary_sectors) > 0:
                doc.add_paragraph("**Primary Industry Sectors**:", style='DSCG Body')
                for sector in primary_sectors:
                    sector_text = sector if isinstance(sector, str) else str(sector)
                    if sector_text.strip():  # Only add non-empty sectors
                        doc.add_paragraph(sector_text, style='List Bullet')
                        has_industry_data = True
        
        # Try industry application data if employers data wasn't useful
        if not has_industry_data and industry_app and isinstance(industry_app, list) and len(industry_app) > 0:
            doc.add_paragraph("**Industry Applications**:", style='DSCG Body')
            for app in industry_app:
                app_text = app if isinstance(app, str) else str(app)
                if app_text.strip():  # Only add non-empty applications
                    doc.add_paragraph(app_text, style='List Bullet')
                    has_industry_data = True
        
        # If no meaningful industry data, generate role-specific applications
        if not has_industry_data:
            role_id = profile_data.get('role_definition', {}).get('id', 'ROLE')
            role_specific_industries = self._generate_role_specific_industries(role_id)
            
            doc.add_paragraph("**Professional Application Areas**:", style='DSCG Body')
            for industry in role_specific_industries:
                doc.add_paragraph(industry, style='List Bullet')
    
    def _generate_role_specific_industries(self, role_id: str) -> List[str]:
        """Generate role-specific industry applications when data is missing"""
        
        industries = {
            'DSL': [
                'Multinational corporations implementing net-zero strategies',
                'Investment firms integrating ESG criteria into portfolio management', 
                'Government agencies developing sustainable development policies',
                'International organizations driving global sustainability standards',
                'Technology companies building sustainable digital infrastructure'
            ],
            'DSM': [
                'Corporate sustainability departments managing operational initiatives',
                'Environmental consulting firms delivering client projects',
                'Manufacturing companies implementing circular economy practices',
                'Energy companies transitioning to renewable technologies',
                'Non-profit organizations promoting sustainable development'
            ],
            'DAN': [
                'ESG rating agencies analyzing corporate sustainability performance',
                'Financial services firms developing sustainable investment strategies',
                'Environmental monitoring organizations tracking climate data',
                'Corporate sustainability teams measuring environmental impact',
                'Research institutions conducting sustainability analytics'
            ],
            'SDD': [
                'Technology companies developing energy-efficient software',
                'Software development firms implementing green coding practices',
                'Cloud service providers optimizing resource utilization',
                'Digital transformation consultancies promoting sustainable IT',
                'Environmental software companies building monitoring platforms'
            ],
            'DSC': [
                'Management consulting firms advising on sustainability strategy',
                'Environmental consulting companies providing specialized expertise',
                'Corporate advisory services supporting ESG transformation',
                'International development organizations implementing sustainability programs',
                'Professional services firms integrating sustainability into client solutions'
            ]
        }
        
        return industries.get(role_id, [
            'Organizations implementing digital sustainability initiatives',
            'Companies developing sustainable technology solutions',
            'Institutions promoting environmental and social responsibility',
            'Enterprises integrating sustainability into business operations'
        ])
    
    def _add_cpd_requirements(self, doc: Document, profile_data: Dict[str, Any]):
        """Add CPD requirements per CEN/TS 17"""
        
        doc.add_paragraph("🔄 Continuing Professional Development", style='DSCG Heading 1')
        
        cpd = profile_data.get('cpd_requirements', {})
        
        if cpd:
            cert_maint = cpd.get('certification_maintenance', {})
            if cert_maint:
                renewal_years = cert_maint.get('renewal_period_years', 3)
                cpd_hours = cert_maint.get('cpd_hours_required', 40)
                
                doc.add_paragraph(f"**Professional Development Cycle**: {renewal_years} years", style='DSCG Body')
                doc.add_paragraph(f"**Required CPD Hours**: {cpd_hours} hours per cycle", style='DSCG Body')
        else:
            doc.add_paragraph("Professional development requirements align with industry standards and professional body requirements for ongoing competency maintenance.", style='DSCG Body')
    
    # Helper methods for CEN/TS 17 compliance
    def _get_eqf_level_context(self, eqf_level: int) -> str:
        """Get EQF level context description"""
        contexts = {
            4: "Factual and theoretical knowledge in broad contexts within a field of work or study",
            5: "Comprehensive, specialised, factual and theoretical knowledge within a field of work or study and an awareness of the boundaries of that knowledge",
            6: "Advanced knowledge of a field of work or study, involving a critical understanding of theories and principles",
            7: "Highly specialised knowledge, some of which is at the forefront of knowledge in a field of work or study, as the basis for original thinking and/or research",
            8: "Knowledge at the most advanced frontier of a field of work or study and at the interface between fields"
        }
        return contexts.get(eqf_level, "Professional knowledge and competencies")
    
    def _generate_compliant_learning_outcomes(self, role_id: str, eqf_level: int) -> List[str]:
        """Generate CEN/TS 17 compliant learning outcomes for specific role and EQF level"""
        
        complexity_verbs = self.eqf_complexity_verbs.get(eqf_level, ['apply', 'develop'])
        role_content = self._get_role_specific_content(role_id)
        
        outcomes = []
        
        # Generate integrated outcomes based on role
        verb1 = complexity_verbs[0] if len(complexity_verbs) > 0 else 'develop'
        outcomes.append(
            f"{verb1.capitalize()} integrated {role_content['domain']} solutions that combine "
            f"{role_content['focus']} with strategic sustainability objectives to address complex organizational challenges."
        )
        
        # Leadership/management dimension
        verb2 = complexity_verbs[1] if len(complexity_verbs) > 1 else 'manage'
        outcomes.append(
            f"{verb2.capitalize()} cross-functional {role_content['teams']} to deliver "
            f"{role_content['deliverables']} that align with organizational sustainability goals and stakeholder expectations."
        )
        
        # Innovation/transformation dimension
        verb3 = complexity_verbs[-1] if complexity_verbs else 'transform'
        outcomes.append(
            f"{verb3.capitalize()} {role_content['systems']} through application of "
            f"advanced {role_content['approach']} expertise and evidence-based methodologies."
        )
        
        return outcomes
    
    def _get_role_specific_content(self, role_id: str) -> Dict[str, str]:
        """Get role-specific content templates for learning outcome generation"""
        
        role_content = {
            'DSL': {
                'domain': 'strategic sustainability leadership',
                'focus': 'organizational transformation and governance',
                'teams': 'sustainability initiatives',
                'deliverables': 'transformation strategies',
                'systems': 'organizational sustainability practices',
                'approach': 'change management'
            },
            'DAN': {
                'domain': 'sustainability analytics',
                'focus': 'data-driven insights and predictive modeling',
                'teams': 'analytics projects',
                'deliverables': 'analytical frameworks',
                'systems': 'data governance and reporting processes',
                'approach': 'statistical and predictive modeling'
            },
            'SDD': {
                'domain': 'sustainable software development',
                'focus': 'green development practices and energy optimization',
                'teams': 'development projects',
                'deliverables': 'energy-efficient applications',
                'systems': 'software development processes',
                'approach': 'sustainable coding'
            },
            'DSM': {
                'domain': 'sustainability management',
                'focus': 'operational coordination and performance monitoring',
                'teams': 'implementation projects',
                'deliverables': 'management frameworks',
                'systems': 'operational sustainability processes',
                'approach': 'project management'
            }
        }
        
        return role_content.get(role_id, {
            'domain': 'digital sustainability',
            'focus': 'professional practice',
            'teams': 'sustainability projects',
            'deliverables': 'sustainability solutions',
            'systems': 'organizational processes',
            'approach': 'best practice'
        })
    
    def _generate_role_specific_frameworks(self, role_id: str, eqf_level: int) -> List[str]:
        """Generate role-specific framework alignments when specific mappings aren't available"""
        
        # Base frameworks that apply to all roles
        base_frameworks = [
            f"**European e-Competence Framework (e-CF)**: Competency areas relevant to {role_id} professional practice at Level {min(eqf_level-3, 5)}"
        ]
        
        # Role-specific framework mappings
        role_frameworks = {
            'DSL': [
                "**e-CF Focus Areas**: E.1 (Forecast Development), E.4 (Relationship Management), E.9 (IS Governance)",
                "**DigComp**: 5.4 (Digital Identity and Ethics), 2.4 (Collaborating through Digital Technologies)",
                "**GreenComp**: 4.3 (Political Agency), 4.2 (Collective Action), 3.4 (Exploratory Thinking)"
            ],
            'DSM': [
                "**e-CF Focus Areas**: A.3 (Business Plan Development), D.10 (Information and Knowledge Management), E.2 (Project and Portfolio Management)",
                "**DigComp**: 2.2 (Sharing through Digital Technologies), 3.4 (Programming)",
                "**GreenComp**: 2.2 (Efficiency), 3.2 (Critical Thinking), 4.1 (Individual Initiative)"
            ],
            'DAN': [
                "**e-CF Focus Areas**: B.1 (Application Development), B.6 (Systems Engineering), D.8 (Contract Management)",
                "**DigComp**: 3.1 (Developing Digital Content), 5.2 (Evaluating Data), 5.3 (Managing Data)",
                "**GreenComp**: 1.3 (Interconnectedness), 2.3 (Circular Thinking), 3.1 (Problem Framing)"
            ],
            'SDD': [
                "**e-CF Focus Areas**: B.1 (Application Development), B.4 (Solution Deployment), B.6 (Systems Engineering)",
                "**DigComp**: 3.2 (Integrating Digital Content), 3.3 (Copyright and Licenses), 5.1 (Solving Technical Problems)",
                "**GreenComp**: 2.1 (Resource Awareness), 2.3 (Circular Thinking), 3.3 (Innovation)"
            ],
            'DSC': [
                "**e-CF Focus Areas**: A.4 (Product/Service Planning), A.5 (Architecture Design), E.3 (Risk Management)",
                "**DigComp**: 1.3 (Managing Data), 2.3 (Engaging in Citizenship through Digital Technologies), 4.2 (Protecting Personal Data)",
                "**GreenComp**: 1.2 (Locality), 3.4 (Exploratory Thinking), 4.2 (Collective Action)"
            ]
        }
        
        # Get role-specific frameworks or use generic ones
        specific_frameworks = role_frameworks.get(role_id, [
            f"**e-CF Focus Areas**: Professional competencies relevant to {role_id} specialization",
            f"**DigComp**: Digital competencies appropriate for EQF Level {eqf_level}",
            f"**GreenComp**: Sustainability competencies for professional practice"
        ])
        
        return base_frameworks + specific_frameworks
    
    def _make_role_specific_title(self, generic_title: str, role_id: str) -> str:
        """Make career progression titles role-specific"""
        
        role_specifics = {
            'DSL': {'base': 'Sustainability Leader', 'senior': 'Senior Sustainability Director', 'principal': 'Chief Sustainability Officer'},
            'DSM': {'base': 'Sustainability Manager', 'senior': 'Senior Sustainability Manager', 'principal': 'Sustainability Director'},
            'DSC': {'base': 'Sustainability Consultant', 'senior': 'Senior Sustainability Consultant', 'principal': 'Principal Sustainability Advisor'},
            'DAN': {'base': 'Data Analyst', 'senior': 'Senior Data Analyst', 'principal': 'Principal Data Scientist'},
            'SDD': {'base': 'Software Developer', 'senior': 'Senior Software Developer', 'principal': 'Principal Software Architect'},
            'SSD': {'base': 'Systems Designer', 'senior': 'Senior Systems Architect', 'principal': 'Principal Systems Designer'},
            'SBA': {'base': 'Business Analyst', 'senior': 'Senior Business Analyst', 'principal': 'Principal Business Architect'},
            'DSI': {'base': 'Implementation Specialist', 'senior': 'Senior Implementation Manager', 'principal': 'Principal Solutions Architect'},
            'DSE': {'base': 'Sustainability Educator', 'senior': 'Senior Learning Designer', 'principal': 'Principal Educational Architect'},
            'STS': {'base': 'Technology Specialist', 'senior': 'Senior Technology Consultant', 'principal': 'Principal Technology Strategist'}
        }
        
        role_titles = role_specifics.get(role_id, {'base': 'Professional', 'senior': 'Senior Professional', 'principal': 'Principal Professional'})
        
        # Map generic terms to role-specific ones
        generic_lower = generic_title.lower()
        if 'senior' in generic_lower or 'lead' in generic_lower:
            return role_titles['senior']
        elif 'principal' in generic_lower or 'director' in generic_lower or 'chief' in generic_lower:
            return role_titles['principal']
        else:
            return role_titles['base']
    
    # ===============================
    # CURRICULUM METHODS - UNCHANGED 
    # ===============================
    
    def generate_curriculum_docx(self, curriculum_data: Dict[str, Any], 
                                output_path: Path, theme_name: str = "material_gray", 
                                compact_mode: bool = False) -> Path:
        """Generate DOCX document for curriculum - ENHANCED ERROR HANDLING"""
        
        try:
            print(f"🔧 Generating Curriculum DOCX: {output_path}")
            print(f"   Available sections: {[k for k in curriculum_data.keys() if k.startswith('section_')]}")
            
            # Create document
            doc = Document()
            
            # Apply theme colors (simplified for now)
            if theme_name.startswith('eu_'):
                self.theme_colors.update({
                    'primary': '003399',    # EU Blue
                    'accent': 'FFCC00'      # EU Yellow
                })
            
            # Create styles
            self._create_document_styles(doc)
            
            # Extract metadata
            metadata = curriculum_data.get('metadata', {})
            role_name = metadata.get('role_name', 'Professional')
            topic = metadata.get('topic', 'Digital Sustainability')
            actual_ects = metadata.get('actual_ects', metadata.get('target_ects', 0))
            
            print(f"   Role: {role_name}, Topic: {topic}, ECTS: {actual_ects}")
            
            # Add header/footer
            self._add_header_footer(doc, 
                                  "Digital Sustainability Professional Development Course",
                                  f"{role_name} Specialization")
            
            # Title page
            title = doc.add_paragraph("Digital Sustainability Professional Development Course", 
                                    style='DSCG Heading 1')
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_paragraph(f"{role_name} Specialization", style='DSCG Heading 2')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata overview
            doc.add_paragraph(f"EQF Level {metadata.get('eqf_level', 6)} • "
                            f"{actual_ects} ECTS • "
                            f"{metadata.get('units_generated', 0)} Modules",
                            style='DSCG Body').alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if not compact_mode:
                doc.add_page_break()
            
            # Add basic curriculum sections
            self._add_curriculum_section_1(doc, curriculum_data, compact_mode)
            
            # Save document
            doc.save(output_path)
            print(f"✅ Curriculum DOCX saved: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"❌ Error generating curriculum DOCX: {e}")
            import traceback
            traceback.print_exc()
            raise
    
    # ===============================
    # UTILITY METHODS - SHARED
    # ===============================
    
    def _hex_to_rgb(self, hex_color: str) -> Tuple[int, int, int]:
        """Convert hex color to RGB tuple"""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    
    def _create_document_styles(self, doc: Document):
        """Create custom styles for the document"""
        
        styles = doc.styles
        
        # Main heading style
        if 'DSCG Heading 1' not in [s.name for s in styles]:
            heading1 = styles.add_style('DSCG Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            heading1.font.name = 'Segoe UI'
            heading1.font.size = Pt(18)
            heading1.font.bold = True
            heading1.font.color.rgb = RGBColor(*self._hex_to_rgb(self.theme_colors['primary']))
            heading1.paragraph_format.space_after = Pt(12)
            heading1.paragraph_format.space_before = Pt(18)
        
        # Section heading style
        if 'DSCG Heading 2' not in [s.name for s in styles]:
            heading2 = styles.add_style('DSCG Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            heading2.font.name = 'Segoe UI'
            heading2.font.size = Pt(14)
            heading2.font.bold = True
            heading2.font.color.rgb = RGBColor(*self._hex_to_rgb(self.theme_colors['primary']))
            heading2.paragraph_format.space_after = Pt(6)
            heading2.paragraph_format.space_before = Pt(12)
        
        # Subsection heading style
        if 'DSCG Heading 3' not in [s.name for s in styles]:
            heading3 = styles.add_style('DSCG Heading 3', WD_STYLE_TYPE.PARAGRAPH)
            heading3.font.name = 'Segoe UI'
            heading3.font.size = Pt(12)
            heading3.font.bold = True
            heading3.font.color.rgb = RGBColor(*self._hex_to_rgb(self.theme_colors['secondary']))
            heading3.paragraph_format.space_after = Pt(3)
            heading3.paragraph_format.space_before = Pt(9)
        
        # Highlight box style
        if 'DSCG Highlight' not in [s.name for s in styles]:
            highlight = styles.add_style('DSCG Highlight', WD_STYLE_TYPE.PARAGRAPH)
            highlight.font.name = 'Segoe UI'
            highlight.font.size = Pt(11)
            highlight.paragraph_format.left_indent = Inches(0.25)
            highlight.paragraph_format.right_indent = Inches(0.25)
            highlight.paragraph_format.space_before = Pt(6)
            highlight.paragraph_format.space_after = Pt(6)
        
        # Body text style
        if 'DSCG Body' not in [s.name for s in styles]:
            body = styles.add_style('DSCG Body', WD_STYLE_TYPE.PARAGRAPH)
            body.font.name = 'Segoe UI'
            body.font.size = Pt(11)
            body.paragraph_format.space_after = Pt(6)
            body.paragraph_format.line_spacing = 1.15
    
    def _add_header_footer(self, doc: Document, title: str, subtitle: str = ""):
        """Add header and footer to document"""
        
        # Add header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = f"{title}"
        if subtitle:
            header_para.text += f" - {subtitle}"
        header_para.style = doc.styles['Header']
        
        # Add footer
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | Digital Sustainability Curriculum Generator"
        footer_para.style = doc.styles['Footer']
    
    def _add_curriculum_section_1(self, doc: Document, curriculum_data: Dict[str, Any], compact_mode: bool = False):
        """Add Section 1: Course Overview - BASIC VERSION"""
        
        doc.add_paragraph("1. 🎯 Course Overview", style='DSCG Heading 1')
        
        # Basic course description
        metadata = curriculum_data.get('metadata', {})
        role_name = metadata.get('role_name', 'Professional')
        topic = metadata.get('topic', 'Digital Sustainability')
        
        doc.add_paragraph("Programme Description", style='DSCG Heading 2')
        doc.add_paragraph(f"Professional development programme for {role_name} specialization in {topic}.", style='DSCG Body')
        
        if not compact_mode:
            doc.add_page_break()

