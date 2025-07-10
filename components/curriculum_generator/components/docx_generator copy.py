# scripts/curriculum_generator/components/docx_generator.py
"""
DOCX Generator for Digital Sustainability Curriculum Generator
Generates professional DOCX documents for both Educational Profiles and Curricula
Maintains consistency with JSON/HTML outputs and theme system
"""

import os
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

class DocxGenerator:
    """Generates professional DOCX documents for DSCG system"""
    
    def __init__(self, project_root: Path):
        self.project_root = project_root
        
        # Theme colors (will be overridden by actual theme)
        self.theme_colors = {
            'primary': '607D8B',      # Material Gray
            'secondary': '90A4AE',    # Light Gray
            'accent': '4CAF50',       # Green
            'text': '333333',         # Dark Gray
            'light_bg': 'F8F9FA'     # Very Light Gray
        }
        
        print("📄 DOCX Generator initialized")
    
    def set_theme_colors(self, theme_colors: Dict[str, str]):
        """Set theme colors for document styling"""
        self.theme_colors.update(theme_colors)
    
    def _hex_to_rgb(self, hex_color: str) -> Tuple[int, int, int]:
        """Convert hex color to RGB tuple"""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    
    def _create_document_styles(self, doc: Document):
        """Create custom styles for the document"""
        
        styles = doc.styles
        
        # Main heading style
        if 'DSCG Heading 1' not in [s.name for s in styles]:
            heading1 = styles.add_style('DSCG Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            heading1.font.name = 'Segoe UI'
            heading1.font.size = Pt(18)
            heading1.font.bold = True
            heading1.font.color.rgb = RGBColor(*self._hex_to_rgb(self.theme_colors['primary']))
            heading1.paragraph_format.space_after = Pt(12)
            heading1.paragraph_format.space_before = Pt(18)
        
        # Section heading style
        if 'DSCG Heading 2' not in [s.name for s in styles]:
            heading2 = styles.add_style('DSCG Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            heading2.font.name = 'Segoe UI'
            heading2.font.size = Pt(14)
            heading2.font.bold = True
            heading2.font.color.rgb = RGBColor(*self._hex_to_rgb(self.theme_colors['primary']))
            heading2.paragraph_format.space_after = Pt(6)
            heading2.paragraph_format.space_before = Pt(12)
        
        # Subsection heading style
        if 'DSCG Heading 3' not in [s.name for s in styles]:
            heading3 = styles.add_style('DSCG Heading 3', WD_STYLE_TYPE.PARAGRAPH)
            heading3.font.name = 'Segoe UI'
            heading3.font.size = Pt(12)
            heading3.font.bold = True
            heading3.font.color.rgb = RGBColor(*self._hex_to_rgb(self.theme_colors['secondary']))
            heading3.paragraph_format.space_after = Pt(3)
            heading3.paragraph_format.space_before = Pt(9)
        
        # Highlight box style
        if 'DSCG Highlight' not in [s.name for s in styles]:
            highlight = styles.add_style('DSCG Highlight', WD_STYLE_TYPE.PARAGRAPH)
            highlight.font.name = 'Segoe UI'
            highlight.font.size = Pt(11)
            highlight.paragraph_format.left_indent = Inches(0.25)
            highlight.paragraph_format.right_indent = Inches(0.25)
            highlight.paragraph_format.space_before = Pt(6)
            highlight.paragraph_format.space_after = Pt(6)
        
        # Body text style
        if 'DSCG Body' not in [s.name for s in styles]:
            body = styles.add_style('DSCG Body', WD_STYLE_TYPE.PARAGRAPH)
            body.font.name = 'Segoe UI'
            body.font.size = Pt(11)
            body.paragraph_format.space_after = Pt(6)
            body.paragraph_format.line_spacing = 1.15
    
    def _add_header_footer(self, doc: Document, title: str, subtitle: str = ""):
        """Add header and footer to document"""
        
        # Add header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = f"{title}"
        if subtitle:
            header_para.text += f" - {subtitle}"
        header_para.style = doc.styles['Header']
        
        # Add footer
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | Digital Sustainability Curriculum Generator"
        footer_para.style = doc.styles['Footer']
    
    def _add_table_of_contents(self, doc: Document, sections: List[Dict[str, str]]):
        """Add table of contents to document"""
        
        # Add TOC heading
        toc_heading = doc.add_paragraph("📋 Table of Contents", style='DSCG Heading 2')
        
        # Add TOC entries
        for i, section in enumerate(sections, 1):
            toc_entry = doc.add_paragraph(style='List Number')
            toc_entry.text = f"{section['title']}"
            toc_entry.paragraph_format.left_indent = Inches(0.25)
        
        # Add page break after TOC
        doc.add_page_break()
    
    def _add_metadata_table(self, doc: Document, metadata: Dict[str, Any]):
        """Add metadata overview table"""
        
        # Create metadata table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        table.columns[0].width = Inches(2)
        table.columns[1].width = Inches(4)
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Attribute'
        header_cells[1].text = 'Value'
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add metadata rows
        metadata_items = [
            ('Role', metadata.get('role_name', 'Professional')),
            ('EQF Level', str(metadata.get('eqf_level', 6))),
            ('ECTS Credits', f"{metadata.get('actual_ects', metadata.get('target_ects', 0))} ECTS"),
            ('Learning Units', str(metadata.get('units_generated', metadata.get('units_requested', 0)))),
            ('Generation Date', metadata.get('generation_date', datetime.now().isoformat())),
            ('System Version', metadata.get('system_version', 'DSCG v4.3'))
        ]
        
        for key, value in metadata_items:
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
        
        doc.add_paragraph()  # Add space after table
    
    def generate_curriculum_docx(self, curriculum_data: Dict[str, Any], 
                                output_path: Path, theme_name: str = "material_gray") -> Path:
        """Generate DOCX document for curriculum"""
        
        try:
            # Create document
            doc = Document()
            
            # Apply theme colors (simplified for now)
            if theme_name.startswith('eu_'):
                self.theme_colors.update({
                    'primary': '003399',    # EU Blue
                    'accent': 'FFCC00'      # EU Yellow
                })
            
            # Create styles
            self._create_document_styles(doc)
            
            # Extract metadata
            metadata = curriculum_data.get('metadata', {})
            role_name = metadata.get('role_name', 'Professional')
            topic = metadata.get('topic', 'Digital Sustainability')
            actual_ects = metadata.get('actual_ects', metadata.get('target_ects', 0))
            
            # Add header/footer
            self._add_header_footer(doc, 
                                  "Digital Sustainability Professional Development Course",
                                  f"{role_name} Specialization")
            
            # Title page
            title = doc.add_paragraph("Digital Sustainability Professional Development Course", 
                                    style='DSCG Heading 1')
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_paragraph(f"{role_name} Specialization", style='DSCG Heading 2')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata overview
            doc.add_paragraph(f"EQF Level {metadata.get('eqf_level', 6)} • "
                            f"{actual_ects} ECTS • "
                            f"{metadata.get('units_generated', 0)} Modules",
                            style='DSCG Body').alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_page_break()
            
            # Table of Contents
            sections = [
                {'title': '1. Course Overview'},
                {'title': '2. Learning Outcomes'},
                {'title': '3. Delivery Options'},
                {'title': '4. Course Organization'},
                {'title': '5. Entry Requirements'},
                {'title': '6. Qualification & Recognition'},
                {'title': '7. Assessment Methods'},
                {'title': '8. Framework Alignment'},
                {'title': '9. Key Benefits'},
                {'title': '10. Cross-Border Compatibility'}
            ]
            self._add_table_of_contents(doc, sections)
            
            # Section 1: Course Overview
            self._add_curriculum_section_1(doc, curriculum_data)
            
            # Section 2: Learning Outcomes
            self._add_curriculum_section_2(doc, curriculum_data)
            
            # Section 3: Delivery Options
            self._add_curriculum_section_3(doc, curriculum_data)
            
            # Section 4: Course Organization
            self._add_curriculum_section_4(doc, curriculum_data)
            
            # Section 5: Entry Requirements
            self._add_curriculum_section_5(doc, curriculum_data)
            
            # Section 6: Qualification & Recognition
            self._add_curriculum_section_6(doc, curriculum_data)
            
            # Section 7: Assessment Methods
            self._add_curriculum_section_7(doc, curriculum_data)
            
            # Section 8: Framework Alignment
            self._add_curriculum_section_8(doc, curriculum_data)
            
            # Section 9: Key Benefits
            self._add_curriculum_section_9(doc, curriculum_data)
            
            # Section 10: Cross-Border Compatibility
            self._add_curriculum_section_10(doc, curriculum_data)
            
            # Save document
            doc.save(output_path)
            print(f"✅ Curriculum DOCX saved: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"❌ Error generating curriculum DOCX: {e}")
            raise
    
    def _add_curriculum_section_1(self, doc: Document, curriculum_data: Dict[str, Any]):
        """Add Section 1: Course Overview"""
        
        doc.add_paragraph("1. 🎯 Course Overview", style='DSCG Heading 1')
        
        section_1 = curriculum_data.get('section_1_what_this_delivers', {})
        
        # Programme Description
        doc.add_paragraph("Programme Description", style='DSCG Heading 3')
        desc_para = doc.add_paragraph(section_1.get('programme_description', ''), style='DSCG Highlight')
        
        # Programme Overview
        doc.add_paragraph("Programme Metadata", style='DSCG Heading 3')
        overview = section_1.get('programme_overview', {})
        
        # Create overview table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light List Accent 1'
        
        overview_items = [
            ('Course Focus', overview.get('role_focus', '')),
            ('Learning Duration', overview.get('duration_weeks', '')),
            ('Total Study Time', overview.get('total_hours', '')),
            ('ECTS Credits', overview.get('ects_credits', '')),
            ('Learning Modules', overview.get('learning_units', '')),
            ('Average Module Size', overview.get('average_unit_size', ''))
        ]
        
        for key, value in overview_items:
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
        
        # Key Features
        doc.add_paragraph("Key Features", style='DSCG Heading 3')
        features = section_1.get('key_features', [])
        for feature in features:
            doc.add_paragraph(feature, style='List Bullet')
        
        doc.add_page_break()
    
    def _add_curriculum_section_2(self, doc: Document, curriculum_data: Dict[str, Any]):
        """Add Section 2: Learning Outcomes"""
        
        doc.add_paragraph("2. 🎯 Learning Outcomes", style='DSCG Heading 1')
        
        section_2 = curriculum_data.get('section_2_learning_outcomes', {})
        
        # Overview
        doc.add_paragraph(section_2.get('overview', ''), style='DSCG Body')
        
        # Technical Competencies
        doc.add_paragraph("Technical Competencies", style='DSCG Heading 3')
        technical = section_2.get('technical_competencies', [])
        for competency in technical:
            doc.add_paragraph(competency, style='List Bullet')
        
        # Transversal Skills
        doc.add_paragraph("Transversal Skills", style='DSCG Heading 3')
        transversal = section_2.get('transversal_competencies', [])
        for skill in transversal:
            doc.add_paragraph(skill, style='List Bullet')
        
        # Framework Alignment
        doc.add_paragraph("Framework Alignment", style='DSCG Heading 3')
        framework = section_2.get('framework_alignment', {})
        for fw_name, fw_desc in framework.items():
            doc.add_paragraph(f"**{fw_name.upper().replace('_', '-')}**: {fw_desc}", style='DSCG Body')
        
        doc.add_page_break()
    
    def _add_curriculum_section_3(self, doc: Document, curriculum_data: Dict[str, Any]):
        """Add Section 3: Delivery Options"""
        
        doc.add_paragraph("3. 📱 Delivery Options and Learning Formats", style='DSCG Heading 1')
        
        section_3 = curriculum_data.get('section_3_delivery_options', {})
        
        # Overview
        doc.add_paragraph(section_3.get('overview', ''), style='DSCG Body')
        
        # Delivery Channels
        doc.add_paragraph("Delivery Channels", style='DSCG Heading 3')
        channels = section_3.get('delivery_channels', {})
        
        for channel_id, channel_info in channels.items():
            # Channel title
            doc.add_paragraph(channel_info.get('title', channel_id.title()), style='DSCG Heading 3')
            
            # Channel description
            doc.add_paragraph(channel_info.get('description', ''), style='DSCG Body')
            
            # Channel details
            details = []
            if 'duration' in channel_info:
                details.append(f"**Duration**: {channel_info['duration']}")
            if 'best_for' in channel_info:
                details.append(f"**Best for**: {channel_info['best_for']}")
            
            for detail in details:
                doc.add_paragraph(detail, style='DSCG Body')
        
        doc.add_page_break()
    
    def _add_curriculum_section_4(self, doc: Document, curriculum_data: Dict[str, Any]):
        """Add Section 4: Course Organization"""
        
        doc.add_paragraph("4. 🧱 Course Organization", style='DSCG Heading 1')
        
        section_4 = curriculum_data.get('section_4_course_organisation', {})
        
        # UOL Explanation
        uol_explanation = section_4.get('uol_explanation', {})
        doc.add_paragraph("Understanding the Module Structure", style='DSCG Heading 3')
        doc.add_paragraph(uol_explanation.get('concept', ''), style='DSCG Highlight')
        
        # Benefits
        benefits = uol_explanation.get('benefits', [])
        if benefits:
            doc.add_paragraph("Benefits", style='DSCG Heading 3')
            for benefit in benefits:
                doc.add_paragraph(benefit, style='List Bullet')
        
        # Learning Units
        doc.add_paragraph("Learning Modules", style='DSCG Heading 3')
        learning_units = section_4.get('learning_units', [])
        
        for i, unit in enumerate(learning_units, 1):
            # Unit header
            unit_title = unit.get('unit_title', f'Module {i}')
            doc.add_paragraph(f"Module {i}: {unit_title}", style='DSCG Heading 3')
            
            # Unit details
            unit_details = []
            unit_details.append(f"**Level**: {unit.get('progression_level', 'Foundation')}")
            unit_details.append(f"**ECTS**: {unit.get('ects', 0)}")
            unit_details.append(f"**Duration**: {unit.get('estimated_hours', 'TBD')}")
            
            for detail in unit_details:
                doc.add_paragraph(detail, style='DSCG Body')
            
            # Learning Outcomes
            outcomes = unit.get('specific_learning_outcomes', [])
            if outcomes:
                doc.add_paragraph("Learning Outcomes:", style='DSCG Body')
                for outcome in outcomes[:3]:  # Limit to first 3
                    doc.add_paragraph(outcome, style='List Bullet')
        
        doc.add_page_break()
    
    def _add_curriculum_section_5(self, doc: Document, curriculum_data: Dict[str, Any]):
        """Add Section 5: Entry Requirements"""
        
        doc.add_paragraph("5. 🎓 Entry Requirements", style='DSCG Heading 1')
        
        section_5 = curriculum_data.get('section_5_entry_requirements', {})
        
        # Standard Requirements
        requirements = section_5.get('standard_requirements', {})
        if requirements:
            doc.add_paragraph("Standard Requirements", style='DSCG Heading 3')
            
            req_table = doc.add_table(rows=1, cols=2)
            req_table.style = 'Light List Accent 1'
            
            for req_type, req_value in requirements.items():
                row_cells = req_table.add_row().cells
                row_cells[0].text = req_type.replace('_', ' ').title()
                row_cells[1].text = str(req_value)
        
        # Alternative Pathways
        pathways = section_5.get('alternative_pathways', [])
        if pathways:
            doc.add_paragraph("Alternative Pathways", style='DSCG Heading 3')
            for pathway in pathways:
                doc.add_paragraph(pathway, style='List Bullet')
        
        doc.add_page_break()
    
    def _add_curriculum_section_6(self, doc: Document, curriculum_data: Dict[str, Any]):
        """Add Section 6: Qualification & Recognition"""
        
        doc.add_paragraph("6. 🏆 Qualification & Recognition", style='DSCG Heading 1')
        
        section_6 = curriculum_data.get('section_6_qualification_recognition', {})
        
        # Primary Qualification
        primary_qual = section_6.get('primary_qualification', '')
        if primary_qual:
            doc.add_paragraph(primary_qual, style='DSCG Highlight')
        
        # What You Receive
        receive = section_6.get('what_you_receive', [])
        if receive:
            doc.add_paragraph("What You Receive", style='DSCG Heading 3')
            for item in receive:
                doc.add_paragraph(item, style='List Bullet')
        
        doc.add_page_break()
    
    def _add_curriculum_section_7(self, doc: Document, curriculum_data: Dict[str, Any]):
        """Add Section 7: Assessment Methods"""
        
        doc.add_paragraph("7. 📊 Assessment Methods", style='DSCG Heading 1')
        
        section_7 = curriculum_data.get('section_7_assessment_methods', {})
        
        # Overview
        overview = section_7.get('overview', '')
        if overview:
            doc.add_paragraph(overview, style='DSCG Body')
        
        # Assessment Breakdown
        breakdown = section_7.get('assessment_breakdown', {})
        if breakdown:
            doc.add_paragraph("Assessment Breakdown", style='DSCG Heading 3')
            
            assess_table = doc.add_table(rows=1, cols=3)
            assess_table.style = 'Light Grid Accent 1'
            
            # Header
            header_cells = assess_table.rows[0].cells
            header_cells[0].text = 'Assessment Type'
            header_cells[1].text = 'Weight'
            header_cells[2].text = 'Description'
            
            for assess_type, assess_info in breakdown.items():
                row_cells = assess_table.add_row().cells
                row_cells[0].text = assess_type.replace('_', ' ').title()
                row_cells[1].text = f"{assess_info.get('percentage', 0)}%"
                row_cells[2].text = assess_info.get('description', '')
        
        doc.add_page_break()
    
    def _add_curriculum_section_8(self, doc: Document, curriculum_data: Dict[str, Any]):
        """Add Section 8: Framework Alignment"""
        
        doc.add_paragraph("8. 🗺️ Framework Alignment", style='DSCG Heading 1')
        
        # Extract framework mappings from learning units
        learning_units = curriculum_data.get('learning_units', [])
        
        if learning_units:
            doc.add_paragraph("Competency Framework Mapping", style='DSCG Heading 3')
            
            # Create framework table
            framework_table = doc.add_table(rows=1, cols=3)
            framework_table.style = 'Light Grid Accent 1'
            
            # Header
            header_cells = framework_table.rows[0].cells
            header_cells[0].text = 'Competency'
            header_cells[1].text = 'Framework'
            header_cells[2].text = 'Reference'
            
            # Sample mappings from first few units
            for unit in learning_units[:3]:
                mappings = unit.get('framework_mappings', {})
                unit_title = unit.get('unit_title', 'Professional Competency')
                
                for framework, codes in mappings.items():
                    if codes:
                        row_cells = framework_table.add_row().cells
                        row_cells[0].text = unit_title[:30] + "..." if len(unit_title) > 30 else unit_title
                        row_cells[1].text = framework.upper().replace('_', '-')
                        row_cells[2].text = ', '.join(codes) if isinstance(codes, list) else str(codes)
        
        doc.add_page_break()
    
    def _add_curriculum_section_9(self, doc: Document, curriculum_data: Dict[str, Any]):
        """Add Section 9: Key Benefits"""
        
        doc.add_paragraph("9. 💡 Key Benefits Summary", style='DSCG Heading 1')
        
        section_8 = curriculum_data.get('section_8_key_benefits_recap', {})
        
        # Benefits
        benefits = section_8.get('benefits', [])
        if benefits:
            for benefit in benefits:
                doc.add_paragraph(benefit, style='List Bullet')
        
        # Value Proposition
        value_prop = section_8.get('value_proposition', '')
        if value_prop:
            doc.add_paragraph("Value Proposition", style='DSCG Heading 3')
            doc.add_paragraph(value_prop, style='DSCG Highlight')
        
        doc.add_page_break()
    
    def _add_curriculum_section_10(self, doc: Document, curriculum_data: Dict[str, Any]):
        """Add Section 10: Cross-Border Compatibility"""
        
        doc.add_paragraph("10. 🌍 Cross-Border Compatibility & Recognition", style='DSCG Heading 1')
        
        section_9 = curriculum_data.get('section_9_cross_border_compatibility', {})
        
        # EU Recognition
        eu_recognition = section_9.get('eu_recognition', '')
        if eu_recognition:
            doc.add_paragraph(eu_recognition, style='DSCG Body')
        
        # Recognition Mechanisms
        mechanisms = section_9.get('recognition_mechanisms', [])
        if mechanisms:
            doc.add_paragraph("Recognition Mechanisms", style='DSCG Heading 3')
            for mechanism in mechanisms:
                doc.add_paragraph(mechanism, style='List Bullet')
        
        # Final spacing
        doc.add_paragraph()
    
    def generate_educational_profile_docx(self, profile_data: Dict[str, Any], 
                                        output_path: Path, theme_name: str = "material_gray") -> Path:
        """Generate DOCX document for educational profile"""
        
        try:
            # Create document
            doc = Document()
            
            # Apply theme colors
            if theme_name.startswith('eu_'):
                self.theme_colors.update({
                    'primary': '003399',    # EU Blue
                    'accent': 'FFCC00'      # EU Yellow
                })
            
            # Create styles
            self._create_document_styles(doc)
            
            # Extract key information
            metadata = profile_data.get('metadata', {})
            role_def = profile_data.get('role_definition', {})
            role_name = role_def.get('name', 'Professional')
            eqf_level = metadata.get('eqf_level', 6)
            
            # Add header/footer
            self._add_header_footer(doc, f"Educational Profile: {role_name}", f"EQF Level {eqf_level}")
            
            # Title page
            title = doc.add_paragraph("🎓 Educational Profile", style='DSCG Heading 1')
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_paragraph(f"{role_name} - EQF Level {eqf_level}", style='DSCG Heading 2')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            metadata_para = doc.add_paragraph(
                f"Role: {role_def.get('id', 'ROLE')} • "
                f"Area: {role_def.get('main_area', 'Digital Sustainability')}",
                style='DSCG Body'
            )
            metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_page_break()
            
            # Add sections
            self._add_profile_career_progression(doc, profile_data)
            self._add_profile_employers(doc, profile_data)
            self._add_profile_competencies(doc, profile_data)
            self._add_profile_program_structure(doc, profile_data)
            self._add_profile_assessment(doc, profile_data)
            self._add_profile_requirements(doc, profile_data)
            self._add_profile_cpd(doc, profile_data)
            
            # Save document
            doc.save(output_path)
            print(f"✅ Educational Profile DOCX saved: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"❌ Error generating Educational Profile DOCX: {e}")
            raise
    
    def _add_profile_career_progression(self, doc: Document, profile_data: Dict[str, Any]):
        """Add career progression section to educational profile"""
        
        doc.add_paragraph("💼 Career Progression & Opportunities", style='DSCG Heading 1')
        
        career_prog = profile_data.get('realistic_career_progression', {})
        
        # Entry Level
        entry_level = career_prog.get('entry_level', {})
        if entry_level:
            doc.add_paragraph("🎯 Entry Level Position", style='DSCG Heading 2')
            
            entry_table = doc.add_table(rows=1, cols=2)
            entry_table.style = 'Light List Accent 1'
            
            entry_items = [
                ('Position', entry_level.get('title', 'Professional')),
                ('EQF Level', str(entry_level.get('eqf_level', 6))),
                ('Salary Range', f"€{entry_level.get('salary_range_eur', {}).get('min', 30000):,} - €{entry_level.get('salary_range_eur', {}).get('max', 50000):,}"),
                ('Experience Required', entry_level.get('experience_required', 'Entry level'))
            ]
            
            for key, value in entry_items:
                row_cells = entry_table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = str(value)
        
        # Progression Roles
        progression_roles = career_prog.get('progression_roles', [])
        if progression_roles:
            doc.add_paragraph("📈 Career Progression Path", style='DSCG Heading 2')
            
            for role in progression_roles:
                if isinstance(role, dict):
                    doc.add_paragraph(role.get('title', 'Senior Role'), style='DSCG Heading 3')
                    doc.add_paragraph(f"**Timeline**: {role.get('years_to_achieve', '3-5')} years", style='DSCG Body')
                    doc.add_paragraph(f"**Salary Increase**: {role.get('salary_increase_percent', '30-50')}%", style='DSCG Body')
                    
                    skills = role.get('additional_skills_needed', [])
                    if skills:
                        skills_text = ', '.join(skills) if isinstance(skills, list) else str(skills)
                        doc.add_paragraph(f"**Skills Needed**: {skills_text}", style='DSCG Body')
        
        doc.add_page_break()
    
    def _add_profile_employers(self, doc: Document, profile_data: Dict[str, Any]):
        """Add typical employers section"""
        
        doc.add_paragraph("🏢 Typical Employers", style='DSCG Heading 1')
        
        employers = profile_data.get('typical_employers', {})
        
        # Primary Sectors
        primary = employers.get('primary_sectors', [])
        if primary:
            doc.add_paragraph("🏢 Primary Employers", style='DSCG Heading 2')
            for sector in primary:
                doc.add_paragraph(sector, style='List Bullet')
        
        # Secondary Sectors
        secondary = employers.get('secondary_sectors', [])
        if secondary:
            doc.add_paragraph("🏭 Secondary Sectors", style='DSCG Heading 2')
            for sector in secondary:
                doc.add_paragraph(sector, style='List Bullet')
        
        # Emerging Opportunities
        emerging = employers.get('emerging_opportunities', [])
        if emerging:
            doc.add_paragraph("🚀 Emerging Opportunities", style='DSCG Heading 2')
            for opp in emerging:
                doc.add_paragraph(opp, style='List Bullet')
        
        doc.add_page_break()
    
    def _add_profile_competencies(self, doc: Document, profile_data: Dict[str, Any]):
        """Add competencies section"""
        
        doc.add_paragraph("🎯 Competences", style='DSCG Heading 1')
        
        competencies = profile_data.get('enhanced_competencies', {})
        
        # Learning Outcomes
        learning_outcomes = competencies.get('learning_outcomes', [])
        if learning_outcomes:
            doc.add_paragraph("🎯 Target Learning Outcomes", style='DSCG Heading 2')
            for outcome in learning_outcomes:
                doc.add_paragraph(outcome, style='List Bullet')
        
        # Framework Mappings
        framework_mappings = competencies.get('framework_mappings', {})
        if framework_mappings:
            doc.add_paragraph("🗺️ Framework Alignment", style='DSCG Heading 2')
            
            for framework, codes in framework_mappings.items():
                framework_name = framework.upper().replace('_', '-')
                if isinstance(codes, list):
                    codes_text = ', '.join(str(code) for code in codes[:4])
                else:
                    codes_text = str(codes)
                doc.add_paragraph(f"**{framework_name}**: {codes_text}", style='DSCG Body')
        
        # Core Competencies
        core_competencies = competencies.get('core_competencies', [])
        if core_competencies:
            doc.add_paragraph("💪 Core Competencies", style='DSCG Heading 2')
            
            for comp in core_competencies:
                if isinstance(comp, dict):
                    doc.add_paragraph(comp.get('name', 'Core Competency'), style='DSCG Heading 3')
                    doc.add_paragraph(comp.get('description', 'Professional competency'), style='DSCG Body')
                    doc.add_paragraph(f"**Proficiency Level**: {comp.get('proficiency_level', 'Professional')}", style='DSCG Body')
                else:
                    doc.add_paragraph(str(comp), style='List Bullet')
        
        doc.add_page_break()
    
    def _add_profile_program_structure(self, doc: Document, profile_data: Dict[str, Any]):
        """Add program structure section"""
        
        doc.add_paragraph("📚 Program Structure", style='DSCG Heading 1')
        
        modular = profile_data.get('modular_structure', {})
        
        # Program Overview
        total_ects = modular.get('total_ects', 60)
        duration = modular.get('duration_semesters', 2)
        
        doc.add_paragraph("📚 Program Overview", style='DSCG Heading 2')
        doc.add_paragraph(f"**Total**: {total_ects} ECTS | **Duration**: {duration} semester(s)", style='DSCG Highlight')
        
        # Modules
        modules = modular.get('modules', [])
        if modules:
            doc.add_paragraph("📖 Modules", style='DSCG Heading 2')
            
            for module in modules:
                if isinstance(module, dict):
                    module_name = module.get('name', 'Module')
                    module_ects = module.get('ects', 7.5)
                    module_semester = module.get('semester', 1)
                    delivery_mode = module.get('delivery_mode', 'Blended')
                    
                    doc.add_paragraph(f"{module_name} ({module_ects} ECTS)", style='DSCG Heading 3')
                    doc.add_paragraph(f"**Semester**: {module_semester} | **Delivery**: {delivery_mode}", style='DSCG Body')
        
        doc.add_page_break()
    
    def _add_profile_assessment(self, doc: Document, profile_data: Dict[str, Any]):
        """Add assessment section"""
        
        doc.add_paragraph("📝 Assessment", style='DSCG Heading 1')
        
        assessment = profile_data.get('assessment_methods', {})
        
        # Primary Methods
        primary_methods = assessment.get('primary_methods', [])
        if primary_methods:
            doc.add_paragraph("📝 How You'll Be Assessed", style='DSCG Heading 2')
            for method in primary_methods:
                doc.add_paragraph(method, style='List Bullet')
        
        # Practical Components
        practical_comp = assessment.get('practical_components', {})
        if practical_comp:
            practical_percent = practical_comp.get('percentage', 50)
            doc.add_paragraph(f"**Practical Components**: {practical_percent}% of assessment", style='DSCG Body')
        
        # Final Assessment
        final_assessment = assessment.get('final_assessment', '')
        if final_assessment:
            doc.add_paragraph(f"**Final Assessment**: {final_assessment}", style='DSCG Body')
        
        doc.add_page_break()
    
    def _add_profile_requirements(self, doc: Document, profile_data: Dict[str, Any]):
        """Add entry requirements section"""
        
        doc.add_paragraph("🎓 Entry Requirements", style='DSCG Heading 1')
        
        entry_req = profile_data.get('entry_requirements', {})
        
        if entry_req:
            req_table = doc.add_table(rows=1, cols=2)
            req_table.style = 'Light List Accent 1'
            
            req_items = [
                ('Education', entry_req.get('formal_education', 'Secondary education')),
                ('Experience', entry_req.get('professional_experience', 'No experience required')),
                ('Digital Skills', entry_req.get('digital_competencies', 'Basic digital literacy')),
                ('Language', entry_req.get('language_requirements', 'English proficiency'))
            ]
            
            for key, value in req_items:
                row_cells = req_table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = str(value)
        
        doc.add_page_break()
    
    def _add_profile_cpd(self, doc: Document, profile_data: Dict[str, Any]):
        """Add CPD requirements section"""
        
        doc.add_paragraph("🔄 Professional Development", style='DSCG Heading 1')
        
        cpd = profile_data.get('cpd_requirements', {})
        
        # Certification Maintenance
        cert_maint = cpd.get('certification_maintenance', {})
        if cert_maint:
            doc.add_paragraph("🔄 Continuing Professional Development", style='DSCG Heading 2')
            
            renewal_years = cert_maint.get('renewal_period_years', 3)
            cpd_hours = cert_maint.get('cpd_hours_required', 40)
            
            doc.add_paragraph(f"**Renewal Period**: {renewal_years} years", style='DSCG Body')
            doc.add_paragraph(f"**CPD Hours Required**: {cpd_hours} hours", style='DSCG Body')
        
        # Micro Learning
        micro_learning = cpd.get('micro_learning_opportunities', {})
        if micro_learning:
            max_recognition = micro_learning.get('maximum_recognition', 10)
            doc.add_paragraph(f"**Stackable Credits**: Up to {max_recognition} ECTS per renewal", style='DSCG Body')
