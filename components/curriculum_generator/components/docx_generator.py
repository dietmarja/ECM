#!/usr/bin/env python3
"""
Enhanced DOCX Generator - Supports both Educational Profiles (CEN/TS 17699:2022) and Curricula
IMPORTANT: This preserves all existing curriculum generation functionality while adding
the three-table format for educational profiles ONLY.
"""

import os
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    print("⚠️  python-docx not available. Install with: pip install python-docx")
    DOCX_AVAILABLE = False

class DocxGenerator:
    """Enhanced DOCX generator supporting both curricula and CEN/TS 17699:2022 educational profiles"""
    
    def __init__(self, project_root: Path):
        self.project_root = project_root
        
        # Check if DOCX is available
        if not DOCX_AVAILABLE:
            print("❌ python-docx library not available")
            return
        
        # Theme colors for document styling
        self.theme_colors = {
            'primary': '607D8B',      # Material Gray
            'secondary': '90A4AE',    # Light Gray
            'accent': '4CAF50',       # Green
            'text': '333333',         # Dark Gray
            'light_bg': 'F8F9FA'     # Very Light Gray
        }
        
        print("📄 Enhanced DOCX Generator initialized (Curricula + Educational Profiles)")
    
    def set_theme_colors(self, theme_colors: Dict[str, str]):
        """Set theme colors for document styling"""
        self.theme_colors.update(theme_colors)
    
    # ===============================
    # CURRICULUM GENERATION METHODS (PRESERVED - NO CHANGES)
    # ===============================
    
    def generate_compact_curriculum_docx(self, curriculum_text: str, role_id: str, 
                                       eqf_level: int, ects: float, output_path: Path) -> Path:
        """Generate curriculum DOCX - PRESERVED EXISTING FUNCTIONALITY"""
        
        if not DOCX_AVAILABLE:
            print(f"❌ Cannot generate DOCX: python-docx not available")
            return output_path
        
        print(f"🔧 Generating Curriculum DOCX: {output_path}")
        
        try:
            # Create document
            doc = Document()
            
            # Create styles for curriculum (different from educational profiles)
            self._create_curriculum_styles(doc)
            
            # Add curriculum content
            lines = curriculum_text.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Handle different content types
                if line.startswith('Curriculum of'):
                    para = doc.add_heading(line, level=0)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif line.startswith('EQF Level'):
                    para = doc.add_paragraph(line, style='DSCG Body')
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Add horizontal line
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.runs[0] if para.runs else para.add_run()
                    run.add_break()
                elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.')):
                    doc.add_heading(line, level=1)
                elif line.endswith(':') and len(line) < 50:
                    doc.add_heading(line, level=2)
                elif '\t' in line and len(line.split('\t')) >= 3:
                    # Handle table content - preserve existing table generation
                    self._add_curriculum_table_row(doc, line)
                else:
                    doc.add_paragraph(line, style='DSCG Body')
            
            # Save document
            doc.save(output_path)
            print(f"✅ Curriculum DOCX saved: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"❌ Error generating Curriculum DOCX: {e}")
            raise
    
    def _add_curriculum_table_row(self, doc: Document, line: str):
        """Add table row for curriculum - PRESERVED HELPER METHOD"""
        parts = line.split('\t')
        # Simple table handling for curriculum content
        if len(parts) >= 2:
            para = doc.add_paragraph()
            para.add_run(parts[0]).bold = True
            para.add_run(f": {parts[1]}")
    
    def _create_curriculum_styles(self, doc: Document):
        """Create styles specifically for curriculum documents - PRESERVED METHOD"""
        if not DOCX_AVAILABLE:
            return
            
        styles = doc.styles
        
        # Main heading style for curricula
        if 'DSCG Heading 1' not in [s.name for s in styles]:
            heading1 = styles.add_style('DSCG Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            heading1.font.name = 'Segoe UI'
            heading1.font.size = Pt(18)
            heading1.font.bold = True
            heading1.font.color.rgb = RGBColor(*self._hex_to_rgb(self.theme_colors['primary']))
            heading1.paragraph_format.space_after = Pt(12)
            heading1.paragraph_format.space_before = Pt(18)
            heading1.paragraph_format.keep_with_next = True
        
        # Section heading style for curricula
        if 'DSCG Heading 2' not in [s.name for s in styles]:
            heading2 = styles.add_style('DSCG Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            heading2.font.name = 'Segoe UI'
            heading2.font.size = Pt(14)
            heading2.font.bold = True
            heading2.font.color.rgb = RGBColor(*self._hex_to_rgb(self.theme_colors['primary']))
            heading2.paragraph_format.space_after = Pt(6)
            heading2.paragraph_format.space_before = Pt(12)
            heading2.paragraph_format.keep_with_next = True
        
        # Body text style for curricula
        if 'DSCG Body' not in [s.name for s in styles]:
            body = styles.add_style('DSCG Body', WD_STYLE_TYPE.PARAGRAPH)
            body.font.name = 'Segoe UI'
            body.font.size = Pt(11)
            body.paragraph_format.space_after = Pt(6)
            body.paragraph_format.line_spacing = 1.15
            body.paragraph_format.widow_control = True
    
    # ===============================
    # EDUCATIONAL PROFILE GENERATION METHODS (THREE-TABLE FORMAT ONLY)
    # ===============================
    
    def generate_educational_profile_docx(self, profile_data: Dict[str, Any], 
                                        output_path: Path, theme_name: str = "material_gray", 
                                        compact_mode: bool = False) -> Path:
        """Generate CEN/TS 17699:2022 Annex E compliant educational profile DOCX with ONLY three tables"""
        
        if not DOCX_AVAILABLE:
            print(f"❌ Cannot generate Educational Profile DOCX: python-docx not available")
            return output_path
        
        print(f"🔧 Generating CEN/TS 17699:2022 Educational Profile: {output_path}")
        
        try:
            # Create document
            doc = Document()
            
            # Create styles for educational profiles (different from curricula)
            self._create_educational_profile_styles(doc)
            
            # Extract key information
            metadata = profile_data.get('metadata', {})
            role_name = metadata.get('role_name', 'Professional')
            eqf_level = metadata.get('eqf_level', 6)
            
            # Add header/footer
            self._add_educational_profile_header_footer(doc, f"Educational Profile: {role_name}", f"EQF Level {eqf_level}")
            
            # CEN/TS 17699:2022 Compliance Statement ONLY (no sections)
            compliance_para = doc.add_paragraph("*Compliant with CEN/TS 17699:2022 Annex E*", style='EP Compliance')
            compliance_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_paragraph("")  # Add spacing
            
            # Generate ONLY the three CEN/TS 17699:2022 tables (NO SECTIONS)
            self._add_table_1_educational_profile_description(doc, profile_data)
            self._add_table_2_learning_outcome_structure(doc, profile_data)
            self._add_table_3_assessment(doc, profile_data)
            
            # Save document
            doc.save(output_path)
            print(f"✅ CEN/TS 17699:2022 Educational Profile DOCX saved: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"❌ Error generating Educational Profile DOCX: {e}")
            import traceback
            traceback.print_exc()
            raise
    
    # ===============================
    # CEN/TS 17699:2022 THREE-TABLE METHODS (NO SECTIONS)
    # ===============================
    
    def _add_table_1_educational_profile_description(self, doc: Document, profile_data: Dict[str, Any]):
        """Add Table 1 — Educational Profile Description per CEN/TS 17699:2022"""
        
        doc.add_heading("Table 1 — Educational Profile Description", level=2, style='EP Heading 2')
        
        # Create table with 9 rows for all fields from E1-E3.docx specification
        table = doc.add_table(rows=9, cols=2)
        table.style = 'Table Grid'
        
        # Set column widths
        table.columns[0].width = Inches(2.8)
        table.columns[1].width = Inches(3.7)
        
        # Get data from profile_data structure
        if 'table_1' in profile_data:
            t1_data = profile_data['table_1']
            rows_data = [
                ("Title", t1_data.get('title', '')),
                ("Description", t1_data.get('description', '')),
                ("Goal", t1_data.get('goal', '')),
                ("Scope", t1_data.get('scope', '')),
                ("Competences", t1_data.get('competences', '')),
                ("Complexity", t1_data.get('complexity', '')),
                ("Deliverables", t1_data.get('deliverables', '')),
                ("Perspective (Educational)", t1_data.get('perspective_educational', '')),
                ("Perspective (Professional)", t1_data.get('perspective_professional', ''))
            ]
        else:
            # Fallback data structure
            rows_data = [
                ("Title", "Educational Profile"),
                ("Description", "Professional educational profile"),
                ("Goal", "Develop professional competencies"),
                ("Scope", "Professional practice contexts"),
                ("Competences", "Professional competency framework"),
                ("Complexity", "Professional level competency"),
                ("Deliverables", "Professional outputs"),
                ("Perspective (Educational)", "Educational progression pathways"),
                ("Perspective (Professional)", "Professional career pathways")
            ]
        
        for i, (field, content) in enumerate(rows_data):
            table.cell(i, 0).text = field
            table.cell(i, 1).text = content
        
        # Format table headers (left column) - make them bold and colored
        for i in range(9):
            for paragraph in table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.theme_colors['primary']))
        
        doc.add_paragraph("")  # Add spacing
    
    def _add_table_2_learning_outcome_structure(self, doc: Document, profile_data: Dict[str, Any]):
        """Add Table 2 — Learning Outcome Structure and Contents per CEN/TS 17699:2022"""
        
        doc.add_heading("Table 2 — Learning Outcome Structure and Contents", level=2, style='EP Heading 2')
        
        # Create table with nested structure
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Set column widths
        table.columns[0].width = Inches(2.8)
        table.columns[1].width = Inches(3.7)
        
        # Add headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Level"
        hdr_cells[1].text = "Content"
        
        # Format headers
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.theme_colors['primary']))
        
        # Get data from profile_data structure
        if 'table_2' in profile_data:
            t2_data = profile_data['table_2']
            
            # Add Programme Learning Outcomes
            prog_row = table.add_row()
            prog_row.cells[0].text = "Programme Learning Outcomes"
            prog_row.cells[1].text = t2_data.get('programme_learning_outcomes', '')
            
            # Format Programme Learning Outcomes cell
            for paragraph in prog_row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            
            # Add Unit Learning Outcomes
            unit_outcomes = t2_data.get('unit_learning_outcomes', {})
            for unit_name, unit_content in unit_outcomes.items():
                unit_row = table.add_row()
                unit_row.cells[0].text = f"Unit Learning Outcomes\n({unit_name})"
                unit_row.cells[1].text = unit_content
                
                # Format unit title
                for paragraph in unit_row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.theme_colors['secondary']))
        else:
            # Fallback if data not provided
            prog_row = table.add_row()
            prog_row.cells[0].text = "Programme Learning Outcomes"
            prog_row.cells[1].text = "Comprehensive professional competencies for educational profile."
        
        doc.add_paragraph("")  # Add spacing
    
    def _add_table_3_assessment(self, doc: Document, profile_data: Dict[str, Any]):
        """Add Table 3 — Assessment per CEN/TS 17699:2022"""
        
        doc.add_heading("Table 3 — Assessment", level=2, style='EP Heading 2')
        
        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Set column widths
        table.columns[0].width = Inches(2.2)
        table.columns[1].width = Inches(2.2)
        table.columns[2].width = Inches(2.1)
        
        # Add headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Unit Learning Outcome"
        hdr_cells[1].text = "Assessment Type"
        hdr_cells[2].text = "Validation of Prior Learning"
        
        # Format headers
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.theme_colors['primary']))
        
        # Get data from profile_data structure
        if 'table_3' in profile_data:
            assessment_data = profile_data['table_3']
            for assessment_item in assessment_data:
                row = table.add_row()
                row.cells[0].text = assessment_item.get('unit_learning_outcome', '')
                row.cells[1].text = assessment_item.get('assessment_type', '')
                row.cells[2].text = assessment_item.get('validation_prior_learning', '')
        else:
            # Fallback if data not provided
            fallback_assessments = [
                ("Unit 1: Foundation", "Portfolio Assessment", "Yes - through portfolio review"),
                ("Unit 2: Application", "Applied Project", "Partial - practical demonstration required"),
                ("Unit 3: Integration", "Integrated Assessment", "Limited - comprehensive assessment required")
            ]
            
            for unit_outcome, assessment_type, vpl_info in fallback_assessments:
                row = table.add_row()
                row.cells[0].text = unit_outcome
                row.cells[1].text = assessment_type
                row.cells[2].text = vpl_info
        
        doc.add_paragraph("")  # Add spacing
    
    # ===============================
    # EDUCATIONAL PROFILE STYLING METHODS
    # ===============================
    
    def _create_educational_profile_styles(self, doc: Document):
        """Create styles specifically for educational profiles (separate from curricula)"""
        
        if not DOCX_AVAILABLE:
            return
            
        styles = doc.styles
        
        # Compliance statement style
        if 'EP Compliance' not in [s.name for s in styles]:
            compliance = styles.add_style('EP Compliance', WD_STYLE_TYPE.PARAGRAPH)
            compliance.font.name = 'Segoe UI'
            compliance.font.size = Pt(11)
            compliance.font.italic = True
            compliance.font.color.rgb = RGBColor(*self._hex_to_rgb(self.theme_colors['secondary']))
            compliance.paragraph_format.space_after = Pt(12)
        
        # Table heading style for educational profiles
        if 'EP Heading 2' not in [s.name for s in styles]:
            heading2 = styles.add_style('EP Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            heading2.font.name = 'Segoe UI'
            heading2.font.size = Pt(14)
            heading2.font.bold = True
            heading2.font.color.rgb = RGBColor(*self._hex_to_rgb(self.theme_colors['primary']))
            heading2.paragraph_format.space_after = Pt(8)
            heading2.paragraph_format.space_before = Pt(16)
            heading2.paragraph_format.keep_with_next = True
    
    def _add_educational_profile_header_footer(self, doc: Document, title: str, subtitle: str = ""):
        """Add header and footer specifically for educational profiles"""
        
        if not DOCX_AVAILABLE:
            return
        
        # Add header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = f"{title}"
        if subtitle:
            header_para.text += f" - {subtitle}"
        header_para.style = doc.styles['Header']
        
        # Add footer
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | CEN/TS 17699:2022 Educational Profile"
        footer_para.style = doc.styles['Footer']
    
    # ===============================
    # UTILITY METHODS (SHARED)
    # ===============================
    
    def _hex_to_rgb(self, hex_color: str) -> Tuple[int, int, int]:
        """Convert hex color to RGB tuple"""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))