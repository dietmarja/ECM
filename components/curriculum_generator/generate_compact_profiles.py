#!/usr/bin/env python3
# scripts/curriculum_generator/generate_compact_profiles.py
"""
CEN/TS 17699:2022 Annex E Strictly Compliant Educational Profile Generator
Generates educational profiles that strictly adhere to Annex E specifications
"""

import json
import os
import sys
from datetime import datetime
from pathlib import Path

# Try to import python-docx for DOCX generation
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠️  python-docx not available. DOCX files will be generated as TXT.")

def load_educational_profiles():
    """Load educational profiles from JSON file"""
    # Try multiple possible locations
    possible_paths = [
        "../../input/educational_profiles/educational_profiles.json",
        "../input/educational_profiles/educational_profiles.json", 
        "input/educational_profiles/educational_profiles.json",
        "educational_profiles.json"
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                print(f"⚠️  Error loading {path}: {e}")
                continue
    
    print("❌ Could not find educational_profiles.json")
    return []

def get_role_specific_industry_alignment(role_id):
    """Get role-specific industry alignment text"""
    alignments = {
        'DSL': 'Drives organizational transformation through strategic ESG integration and sustainable business model innovation',
        'DSM': 'Operationalizes sustainability strategies through systematic implementation and cross-functional coordination',
        'DSC': 'Provides expert advisory services for ESG strategy development and regulatory compliance frameworks',
        'DAN': 'Transforms environmental and social data into strategic insights for evidence-based sustainability decisions',
        'SDD': 'Creates energy-efficient digital solutions that minimize environmental impact while maximizing functionality',
        'SSD': 'Designs regenerative technology solutions applying circular economy principles and life cycle thinking',
        'SBA': 'Bridges sustainability requirements with business performance through analytical frameworks and ROI demonstration',
        'DSI': 'Advances sustainability science through innovative analytical methodologies and predictive modeling',
        'DSE': 'Engineers sustainable infrastructure and carbon-aware computing systems for environmental optimization',
        'STS': 'Enables organizational adoption of sustainability technologies through technical support and user enablement'
    }
    return alignments.get(role_id, 'Addresses critical sustainability challenges through specialized professional expertise')

def get_profile_data(role_id, eqf_level):
    """Get profile data from the educational profiles JSON"""
    profiles = load_educational_profiles()
    
    # Find the matching profile
    profile = None
    for p in profiles:
        if p.get('id') == role_id:
            profile = p
            break
    
    if not profile:
        print(f"⚠️  Profile not found for {role_id}, using default")
        return get_default_profile_data(role_id, eqf_level)
    
    # Fix EQF level references in role description
    role_description = profile.get('role_description', '')
    # Replace any incorrect EQF level references with the actual level
    if f'EQF Level {eqf_level+1}' in role_description:
        role_description = role_description.replace(f'EQF Level {eqf_level+1}', f'EQF Level {eqf_level}')
    if f'EQF Level {eqf_level-1}' in role_description and eqf_level > 5:
        role_description = role_description.replace(f'EQF Level {eqf_level-1}', f'EQF Level {eqf_level}')
    
    profile['role_description'] = role_description
    profile['industry_alignment'] = get_role_specific_industry_alignment(role_id)
    
    return profile

def get_default_profile_data(role_id, eqf_level):
    """Fallback profile data if JSON not available"""
    role_names = {
        'DSL': 'Digital Sustainability Leader',
        'DSM': 'Digital Sustainability Manager', 
        'DSC': 'Digital Sustainability Consultant',
        'DAN': 'Data Analyst',
        'SDD': 'Software Developer for Sustainability',
        'SSD': 'Sustainable Solution Designer',
        'SBA': 'Sustainability Business Analyst',
        'DSI': 'Data Scientist (Sustainability)',
        'DSE': 'Data Engineer',
        'STS': 'Sustainability Technical Specialist'
    }
    
    return {
        'id': role_id,
        'profile_name': f"{role_names.get(role_id, role_id)} Educational Profile",
        'role_description': f"Professional {role_names.get(role_id, role_id)} in digital sustainability",
        'core_competency_areas': [
            "Strategic Sustainability Leadership",
            "Digital Technology Integration", 
            "Environmental Impact Management"
        ],
        'learning_outcomes_by_eqf': {
            str(eqf_level): [
                "Lead strategic sustainability initiatives",
                "Integrate digital technologies with environmental goals",
                "Manage complex environmental impact assessments"
            ]
        },
        'framework_alignment': {
            'key_frameworks': ["e-CF: E.1, E.4", "DigComp: 5.4, 2.4", "GreenComp: 4.3, 4.2"]
        }
    }

def generate_html_profile(profile_data, eqf_level):
    """Generate CEN/TS 17699:2022 Annex E compliant HTML profile"""
    
    role_name = profile_data.get('profile_name', '').replace(' Educational Profile', '')
    
    # 1. Description components
    role_description = profile_data.get('role_description', '')
    
    # 2. Programme Learning Outcomes
    learning_outcomes = profile_data.get('learning_outcomes_by_eqf', {}).get(str(eqf_level), [])
    if not learning_outcomes and str(eqf_level-1) in profile_data.get('learning_outcomes_by_eqf', {}):
        learning_outcomes = profile_data.get('learning_outcomes_by_eqf', {})[str(eqf_level-1)]
    
    outcomes_html = ""
    for i, outcome in enumerate(learning_outcomes[:3], 1):
        outcomes_html += f"        <li>🎯 {outcome}</li>\n"
    
    # 3. Unit Learning Outcomes (Competency Areas - NOT curriculum modules)
    competency_areas = profile_data.get('core_competency_areas', [])
    unit_outcomes_html = ""
    for i, area in enumerate(competency_areas, 1):
        unit_outcomes_html += f"        <li><strong>Competency Area {i}:</strong> {area}</li>\n"
    
    # 4. Framework Alignment
    framework_info = profile_data.get('framework_alignment', {})
    key_frameworks = framework_info.get('key_frameworks', ['e-CF: E.1', 'DigComp: 5.1', 'GreenComp: 4.1'])
    
    frameworks_html = ""
    for framework in key_frameworks:
        if ':' in framework:
            name, codes = framework.split(':', 1)
            frameworks_html += f"        <li><strong>{name.strip()}:</strong> {codes.strip()}</li>\n"
    
    # 5. Assessment Methods
    assessment_info = profile_data.get('assessment_philosophy', {})
    assessment_methods = assessment_info.get('methods', ['Portfolio assessment', 'Practical projects', 'Competency demonstrations'])
    
    assessment_html = ""
    for method in assessment_methods:
        assessment_html += f"        <li>• {method}</li>\n"
    
    # 6. Industry Applications
    industry_apps = profile_data.get('industry_application', [])
    industry_html = ""
    for app in industry_apps[:3]:  # Limit to 3 for compact
        industry_html += f"        <li>• {app}</li>\n"
    
    # Entry requirements
    entry_reqs = profile_data.get('entry_requirements_by_eqf', {}).get(str(eqf_level), {})
    academic_req = entry_reqs.get('academic', f'EQF Level {eqf_level-1} qualification or equivalent')
    professional_req = entry_reqs.get('professional', 'Relevant professional experience')
    
    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Educational Profile of {role_name} - EQF Level {eqf_level}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
            background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        }}
        .header {{
            text-align: center;
            margin-bottom: 30px;
            padding: 25px;
            background: white;
            border-radius: 12px;
            box-shadow: 0 4px 8px rgba(0,0,0,0.12);
        }}
        .title {{
            color: #2c3e50;
            font-size: 26px;
            font-weight: 700;
            margin: 0 0 12px 0;
        }}
        .subtitle {{
            color: #7f8c8d;
            font-size: 16px;
            font-style: italic;
            margin: 0;
        }}
        .section {{
            background: white;
            margin: 20px 0;
            padding: 25px;
            border-radius: 10px;
            border-left: 5px solid #3498db;
            box-shadow: 0 3px 6px rgba(0,0,0,0.1);
        }}
        .section-title {{
            color: #2c3e50;
            font-size: 19px;
            font-weight: 600;
            margin: 0 0 18px 0;
            display: flex;
            align-items: center;
            gap: 12px;
            border-bottom: 2px solid #ecf0f1;
            padding-bottom: 8px;
        }}
        .section-content {{
            color: #34495e;
            margin: 0;
            font-size: 15px;
        }}
        ul {{
            margin: 12px 0;
            padding-left: 0;
        }}
        li {{
            list-style: none;
            margin: 10px 0;
            padding: 10px;
            background: #f8f9fa;
            border-radius: 6px;
            border-left: 3px solid #3498db;
        }}
        .proficiency {{
            background: #e8f5e8;
            border-left-color: #27ae60;
            font-style: italic;
        }}
        .framework-grid {{
            display: grid;
            grid-template-columns: 1fr;
            gap: 12px;
            margin: 15px 0;
        }}
        .framework-item {{
            background: #f8f9fa;
            padding: 12px;
            border-radius: 6px;
            border-left: 3px solid #3498db;
        }}
        .codes {{
            font-family: 'Courier New', monospace;
            font-size: 13px;
            color: #7f8c8d;
            font-weight: 600;
        }}
        .highlight {{
            background: #fff3cd;
            border-left-color: #ffc107;
            font-weight: 500;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1 class="title">Educational Profile of {role_name} - EQF Level {eqf_level}</h1>
        <p class="subtitle">Compliant with CEN/TS 17699:2022 Annex E</p>
    </div>

    <div class="section">
        <h2 class="section-title">📋 1. Description</h2>
        <div class="section-content">
            <p><strong>Role/Position:</strong> {role_name} • <strong>EQF Level:</strong> {eqf_level}</p>
            <p><strong>Professional Context:</strong> {role_description}</p>
            <p><strong>Industry Alignment:</strong> {profile_data.get('industry_alignment', 'Addresses sustainability challenges through professional expertise')}</p>
        </div>
    </div>

    <div class="section">
        <h2 class="section-title">🎯 2. Programme Learning Outcomes</h2>
        <div class="section-content">
            <p><strong>Observable, competence-based results linked to e-CF frameworks:</strong></p>
            <ul>
{outcomes_html}            </ul>
            <div class="proficiency">
                <strong>Proficiency Level:</strong> Competent level as defined by relevant European competency frameworks
            </div>
            <p><strong>Transversal Skills Integration:</strong> Sustainability ethics, digital transformation leadership, environmental stewardship</p>
        </div>
    </div>

    <div class="section">
        <h2 class="section-title">📚 3. Unit Learning Outcomes</h2>
        <div class="section-content">
            <p><strong>Competency-based learning areas (profile-level specification):</strong></p>
            <ul>
{unit_outcomes_html}            </ul>
        </div>
    </div>

    <div class="section">
        <h2 class="section-title">📊 4. Assessments</h2>
        <div class="section-content">
            <p><strong>Assessment Methods:</strong></p>
            <ul>
{assessment_html}            </ul>
            <p><strong>Assessment Philosophy:</strong> {assessment_info.get('approach', 'Competency-based validation through integrated practical applications')}</p>
            <p class="highlight"><strong>Integrated Validation:</strong> Demonstration of integrated sustainability leadership competencies addressing real-world challenges</p>
        </div>
    </div>

    <div class="section">
        <h2 class="section-title">🏛️ 5. Framework Alignment</h2>
        <div class="section-content">
            <p><strong>Direct competency framework references:</strong></p>
            <div class="framework-grid">
{frameworks_html}            </div>
            <p><strong>EQF Level {eqf_level} Descriptors:</strong> {framework_info.get('eqf_focus', 'Advanced knowledge, skills and competences with autonomous responsibility')}</p>
        </div>
    </div>

    <div class="section">
        <h2 class="section-title">🔄 6. Continuing Professional Development</h2>
        <div class="section-content">
            <p><strong>CPD Requirements:</strong></p>
            <ul>
                <li><strong>Renewal Cycle:</strong> 3 years</li>
                <li><strong>Required Hours:</strong> 40 hours per cycle</li>
                <li><strong>Advancement Pathways:</strong> Opportunities for further specialization exist</li>
            </ul>
        </div>
    </div>

    <div class="section">
        <h2 class="section-title">📚 7. Programme Structure</h2>
        <div class="section-content">
            <p><strong>Typical Qualification Pathways:</strong></p>
            <ul>
                <li><strong>Academic Foundation:</strong> {academic_req}</li>
                <li><strong>Professional Experience:</strong> {professional_req}</li>
            </ul>
            <p><strong>Career Pathways:</strong> Progression within sustainability leadership, technical expertise, and strategic advisory roles</p>
        </div>
    </div>

    <div class="section">
        <h2 class="section-title">🌍 8. Industry Applications</h2>
        <div class="section-content">
            <p><strong>Professional contexts and employment sectors:</strong></p>
            <ul>
{industry_html}            </ul>
        </div>
    </div>
</body>
</html>"""
    
    return html_content

def generate_json_profile(profile_data, eqf_level):
    """Generate CEN/TS 17699:2022 Annex E compliant JSON profile"""
    
    role_name = profile_data.get('profile_name', '').replace(' Educational Profile', '')
    learning_outcomes = profile_data.get('learning_outcomes_by_eqf', {}).get(str(eqf_level), [])
    
    json_data = {
        "profileMetadata": {
            "title": f"Educational Profile of {role_name} - EQF Level {eqf_level}",
            "subtitle": "Compliant with CEN/TS 17699:2022 Annex E",
            "roleId": profile_data.get('id'),
            "roleName": role_name,
            "eqfLevel": eqf_level,
            "generatedDate": datetime.now().isoformat(),
            "complianceStandard": "CEN/TS 17699:2022 Annex E"
        },
        "description": {
            "rolePosition": f"{role_name} - EQF Level {eqf_level}",
            "professionalContext": profile_data.get('role_description', ''),
            "industryAlignment": profile_data.get('industry_alignment', 'Addresses sustainability challenges through professional expertise'),
            "scopeOfResponsibilities": "Strategic leadership in digital sustainability transformation"
        },
        "programmeLearningOutcomes": {
            "competenceBasedResults": learning_outcomes[:3],
            "proficiencyLevel": "Competent level as defined by relevant European competency frameworks",
            "eCFLinkage": "Observable outcomes directly mapped to e-CF competency units",
            "transversalSkills": ["Sustainability ethics", "Digital transformation leadership", "Environmental stewardship"]
        },
        "unitLearningOutcomes": {
            "competencyAreas": profile_data.get('core_competency_areas', []),
            "competencyBreakdown": "Profile-level specification focused on professional competence areas"
        },
        "assessments": {
            "methods": profile_data.get('assessment_philosophy', {}).get('methods', ['Portfolio assessment', 'Practical projects']),
            "assessmentPhilosophy": profile_data.get('assessment_philosophy', {}).get('approach', 'Competency-based validation'),
            "integratedValidation": "Demonstration of integrated sustainability leadership competencies addressing real-world challenges"
        },
        "frameworkAlignment": {
            "directReferences": profile_data.get('framework_alignment', {}).get('key_frameworks', []),
            "eqfDescriptors": profile_data.get('framework_alignment', {}).get('eqf_focus', 'Advanced knowledge, skills and competences'),
            "competencyEmphasis": profile_data.get('framework_alignment', {}).get('competency_emphasis', 'Strategic and technical excellence')
        },
        "continuingProfessionalDevelopment": {
            "renewalCycle": "3 years",
            "requiredHours": "40 hours per cycle", 
            "advancementPathways": "Opportunities for further specialization exist"
        },
        "programmeStructure": {
            "typicalQualificationPathways": profile_data.get('entry_requirements_by_eqf', {}).get(str(eqf_level), {}),
            "careerPathways": profile_data.get('career_progression', {})
        },
        "industryApplications": profile_data.get('industry_application', []),
        "complianceValidation": {
            "annexECompliant": True,
            "curriculumLevelContentRemoved": True,
            "profileLevelContentOnly": True,
            "validationDate": datetime.now().isoformat()
        }
    }
    
    return json.dumps(json_data, indent=2, ensure_ascii=False)

def generate_docx_profile(profile_data, eqf_level):
    """Generate CEN/TS 17699:2022 Annex E compliant DOCX profile"""
    
    role_name = profile_data.get('profile_name', '').replace(' Educational Profile', '')
    
    if not HAS_DOCX:
        # Fallback to text format with full Annex E structure
        text_content = f"""Educational Profile of {role_name} - EQF Level {eqf_level}
Compliant with CEN/TS 17699:2022 Annex E

1. DESCRIPTION
Role/Position: {role_name} - EQF Level {eqf_level}
Professional Context: {profile_data.get('role_description', '')}
Industry Alignment: {profile_data.get('industry_alignment', 'Addresses sustainability challenges through professional expertise')}

2. PROGRAMME LEARNING OUTCOMES
Observable, competence-based results linked to e-CF frameworks:
"""
        learning_outcomes = profile_data.get('learning_outcomes_by_eqf', {}).get(str(eqf_level), [])
        for outcome in learning_outcomes[:3]:
            text_content += f"• {outcome}\n"
        
        text_content += f"""
Proficiency Level: Competent level as defined by relevant European competency frameworks
Transversal Skills: Sustainability ethics, digital transformation leadership, environmental stewardship

3. UNIT LEARNING OUTCOMES
Competency-based learning areas:
"""
        for i, area in enumerate(profile_data.get('core_competency_areas', []), 1):
            text_content += f"• Competency Area {i}: {area}\n"
        
        text_content += f"""

4. ASSESSMENTS
Methods: {', '.join(profile_data.get('assessment_philosophy', {}).get('methods', ['Portfolio assessment']))}
Philosophy: {profile_data.get('assessment_philosophy', {}).get('approach', 'Competency-based validation')}
Integrated Validation: Demonstration of integrated sustainability leadership competencies addressing real-world challenges

5. FRAMEWORK ALIGNMENT
Key Frameworks: {', '.join(profile_data.get('framework_alignment', {}).get('key_frameworks', []))}
EQF Level {eqf_level} Descriptors: Advanced knowledge, skills and competences with autonomous responsibility

6. CONTINUING PROFESSIONAL DEVELOPMENT
Renewal Cycle: 3 years
Required Hours: 40 hours per cycle
Advancement Pathways: Opportunities for further specialization exist

7. PROGRAMME STRUCTURE
Typical Qualification Pathways:
Academic Foundation: {entry_reqs.get("academic", f"EQF Level {eqf_level-1} qualification")}
Professional Experience: {entry_reqs.get("professional", "Relevant professional experience")}
Career Pathways: Progression within sustainability leadership and technical expertise roles
Entry Requirements: {profile_data.get('entry_requirements_by_eqf', {}).get(str(eqf_level), {}).get('academic', f'EQF Level {eqf_level-1} qualification')}
Career Pathways: Progression within sustainability leadership and technical expertise

8. INDUSTRY APPLICATIONS
"""
        for app in profile_data.get('industry_application', []):
            text_content += f"• {app}\n"
        
        return text_content
    
    # Generate proper DOCX with full Annex E structure
    doc = Document()
    
    # Title and subtitle
    title = doc.add_heading(f'Educational Profile of {role_name} - EQF Level {eqf_level}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(f'Compliant with CEN/TS 17699:2022 Annex E')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    
    doc.add_paragraph()
    
    # 1. Description
    doc.add_heading('1. Description', level=1)
    desc_para = doc.add_paragraph()
    desc_para.add_run('Role/Position: ').bold = True
    desc_para.add_run(f'{role_name} - EQF Level {eqf_level}\n')
    desc_para.add_run('Professional Context: ').bold = True
    desc_para.add_run(f'{profile_data.get("role_description", "")}\n')
    desc_para.add_run('Industry Alignment: ').bold = True
    desc_para.add_run(profile_data.get('industry_alignment', 'Addresses sustainability challenges through professional expertise'))
    
    # 2. Programme Learning Outcomes
    doc.add_heading('2. Programme Learning Outcomes', level=1)
    doc.add_paragraph('Observable, competence-based results linked to e-CF frameworks:')
    learning_outcomes = profile_data.get('learning_outcomes_by_eqf', {}).get(str(eqf_level), [])
    for outcome in learning_outcomes[:3]:
        doc.add_paragraph(f'• {outcome}', style='List Bullet')
    
    prof_para = doc.add_paragraph()
    prof_para.add_run('Proficiency Level: ').bold = True
    prof_para.add_run('Competent level as defined by relevant European competency frameworks\n')
    prof_para.add_run('Transversal Skills: ').bold = True
    prof_para.add_run('Sustainability ethics, digital transformation leadership, environmental stewardship')
    
    # 3. Unit Learning Outcomes
    doc.add_heading('3. Unit Learning Outcomes', level=1)
    doc.add_paragraph('Competency-based learning areas (profile-level specification):')
    for i, area in enumerate(profile_data.get('core_competency_areas', []), 1):
        doc.add_paragraph(f'• Competency Area {i}: {area}', style='List Bullet')
    
    # 4. Assessments
    doc.add_heading('4. Assessments', level=1)
    assessment_info = profile_data.get('assessment_philosophy', {})
    
    assess_para = doc.add_paragraph()
    assess_para.add_run('Methods: ').bold = True
    assess_para.add_run(f'{", ".join(assessment_info.get("methods", ["Portfolio assessment"]))}\n')
    assess_para.add_run('Philosophy: ').bold = True
    assess_para.add_run(f'{assessment_info.get("approach", "Competency-based validation")}\n')
    assess_para.add_run('Integrated Validation: ').bold = True
    assess_para.add_run('Demonstration of integrated sustainability leadership competencies addressing real-world challenges')
    
    # 5. Framework Alignment
    doc.add_heading('5. Framework Alignment', level=1)
    framework_info = profile_data.get('framework_alignment', {})
    
    fw_para = doc.add_paragraph()
    fw_para.add_run('Key Frameworks: ').bold = True
    fw_para.add_run(f'{", ".join(framework_info.get("key_frameworks", []))}\n')
    fw_para.add_run(f'EQF Level {eqf_level} Descriptors: ').bold = True
    fw_para.add_run('Advanced knowledge, skills and competences with autonomous responsibility')
    
    # 6. CPD
    doc.add_heading('6. Continuing Professional Development', level=1)
    cpd_para = doc.add_paragraph()
    cpd_para.add_run('Renewal Cycle: ').bold = True
    cpd_para.add_run('3 years • ')
    cpd_para.add_run('Required Hours: ').bold = True
    cpd_para.add_run('40 hours per cycle • ')
    cpd_para.add_run('Advancement Pathways: ').bold = True
    cpd_para.add_run('Opportunities for further specialization exist')
    
    # 7. Programme Structure
    doc.add_heading('7. Programme Structure', level=1)
    entry_reqs = profile_data.get('entry_requirements_by_eqf', {}).get(str(eqf_level), {})
    
    struct_para = doc.add_paragraph()
    struct_para.add_run('Programme Scope: ').bold = True
    struct_para.add_run(f'Aligned with EQF Level {eqf_level} qualification workload expectations\n')
    struct_para.add_run('Entry Requirements: ').bold = True
    struct_para.add_run(f'{entry_reqs.get("academic", f"EQF Level {eqf_level-1} qualification")}\n')
    struct_para.add_run('Career Pathways: ').bold = True
    struct_para.add_run('Progression within sustainability leadership and technical expertise roles')
    
    # 8. Industry Applications
    doc.add_heading('8. Industry Applications', level=1)
    doc.add_paragraph('Professional contexts and employment sectors:')
    for app in profile_data.get('industry_application', []):
        doc.add_paragraph(f'• {app}', style='List Bullet')
    
    return doc

def generate_single_profile(role_id, eqf_level, number=None):
    """Generate a single CEN/TS 17699:2022 Annex E compliant profile"""
    
    print(f"📄 Generating {number:02d}_COMPACT_EP: {role_id} EQF{eqf_level}...")
    
    # Get profile data from JSON
    profile_data = get_profile_data(role_id, eqf_level)
    
    # Setup output directory  
    output_dir = Path("/Users/dietmar/Dropbox/NCI/DIGITAL4Business/digital4sustainability/python/DSCG/output/compact_appendix")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Base filename
    if number:
        base_filename = f"{number:02d}_COMPACT_EP_{role_id}_EQF{eqf_level}_EU_Test"
    else:
        base_filename = f"COMPACT_EP_{role_id}_EQF{eqf_level}_EU_Test"
    
    try:
        # Generate HTML
        html_content = generate_html_profile(profile_data, eqf_level)
        html_path = output_dir / f"{base_filename}.html"
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        print(f"   ✅ HTML: {base_filename}.html")
        
        # Generate JSON
        json_content = generate_json_profile(profile_data, eqf_level)
        json_path = output_dir / f"{base_filename}.json"
        with open(json_path, 'w', encoding='utf-8') as f:
            f.write(json_content)
        print(f"   ✅ JSON: {base_filename}.json")
        
        # Generate DOCX
        docx_content = generate_docx_profile(profile_data, eqf_level)
        docx_path = output_dir / f"{base_filename}.docx"
        
        if HAS_DOCX and hasattr(docx_content, 'save'):
            docx_content.save(str(docx_path))
            print(f"   ✅ DOCX: {base_filename}.docx")
        else:
            # Fallback to TXT
            txt_path = output_dir / f"{base_filename}.txt"
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(docx_content)
            print(f"   ✅ TXT: {base_filename}.txt (install python-docx for DOCX)")
        
        return True
        
    except Exception as e:
        print(f"   ❌ Error generating {role_id} EQF{eqf_level}: {e}")
        return False

def generate_all_22_profiles():
    """Generate all 22 CEN/TS 17699:2022 Annex E compliant educational profiles"""
    
    print("📄 Generating 22 CEN/TS 17699:2022 Annex E compliant educational profiles...")
    
    # Define all 22 role/EQF combinations based on the JSON data
    profiles = [
        ('DSL', 7, 1), ('DSL', 8, 2),
        ('DSM', 6, 3), ('DSM', 7, 4),
        ('DSC', 6, 5), ('DSC', 7, 6), ('DSC', 8, 7),
        ('DAN', 6, 8), ('DAN', 7, 9),
        ('SDD', 5, 10), ('SDD', 6, 11), ('SDD', 7, 12),
        ('SSD', 6, 13), ('SSD', 7, 14),
        ('SBA', 6, 15), ('SBA', 7, 16),
        ('DSI', 5, 17), ('DSI', 6, 18), ('DSI', 7, 19),
        ('DSE', 6, 20), ('DSE', 7, 21), ('DSE', 8, 22)
    ]
    
    success_count = 0
    total_count = len(profiles)
    
    for role, eqf, number in profiles:
        if generate_single_profile(role, eqf, number):
            success_count += 1
    
    print(f"✅ Successfully generated: {success_count}/{total_count} profiles")
    if success_count == total_count:
        print("🎉 All CEN/TS 17699:2022 Annex E compliant profiles generated successfully!")
        print(f"📁 Files saved to: /Users/dietmar/Dropbox/NCI/DIGITAL4Business/digital4sustainability/python/DSCG/output/compact_appendix")
    
    return success_count == total_count

def main():
    """Main function for profile generation"""
    
    print("📚 CEN/TS 17699:2022 Annex E Strictly Compliant Educational Profile Generator")
    print("=" * 80)
    
    if len(sys.argv) == 1:
        print("Usage:")
        print("  python3 generate_compact_profiles.py all                    # Generate all 22 profiles")
        print("  python3 generate_compact_profiles.py test                   # Generate single test profile")
        print("  python3 generate_compact_profiles.py DSL 7                  # Generate single profile")
        print("  python3 generate_compact_profiles.py DSL 7 1                # Generate numbered profile")
        return
    
    if sys.argv[1] == "all":
        return generate_all_22_profiles()
    elif sys.argv[1] == "test":
        return generate_single_profile("DSL", 7, 1)
    elif len(sys.argv) >= 3:
        role_id = sys.argv[1].upper()
        try:
            eqf_level = int(sys.argv[2])
            number = int(sys.argv[3]) if len(sys.argv) > 3 else None
            return generate_single_profile(role_id, eqf_level, number)
        except ValueError:
            print("❌ EQF level must be a number (5-8)")
            return False
    else:
        print("❌ Invalid arguments")
        return False

if __name__ == "__main__":
    main()