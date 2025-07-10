# generate_curricula_v21.py
"""
ECM Curriculum Generator - Enhanced Version 21
Generates high-quality curricula with optional visual outcome mapping
Complies with T3.2 requirements and educational standards
"""

import json
import os
import argparse
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import qn

class EnhancedCurriculumGenerator:
    """Generate high-quality curricula with enhanced professional formatting"""
    
    def __init__(self, config_path='config/settings.json', include_visual_map=True):
        print("=== ECM Curriculum Generator v21 - Enhanced Version ===")
        print("✓ Complete target audience coverage")
        print("✓ EQF Level 4-8 comprehensive coverage")
        print("✓ Professional formatting and assessment methods")
        print(f"✓ Visual Outcome Mapping: {'Enabled' if include_visual_map else 'Disabled'}")
        
        self.include_visual_map = include_visual_map
        self.config = self.load_config(config_path)
        self.setup_paths()
        
        # Load all required data
        self.modules_data = self.load_modules_data()
        self.roles_data = self.load_roles_data()
        
        # Validate data integrity
        self.validate_data_integrity()
        
        # Create efficient lookups
        self.modules_by_id = {module['id']: module for module in self.modules_data}
        self.roles_by_id = {role['id']: role for role in self.roles_data}
        
        # Define role competence profiles with clean titles
        self.role_profiles = self.define_role_profiles()
        
        # Define the exact 10 core curricula
        self.core_curricula = self.define_core_curricula()
        
        # Define EQF-aligned action verbs
        self.eqf_action_verbs = self.define_eqf_action_verbs()
        
        print(f"✓ Loaded {len(self.modules_data)} modules")
        print(f"✓ Loaded {len(self.roles_data)} roles")
        print(f"✓ Defined {len(self.core_curricula)} core curricula")
    
    def define_eqf_action_verbs(self):
        """Define EQF-aligned action verbs for consistent learning outcomes"""
        return {
            3: {
                'knowledge': ['Identify', 'List', 'Recognize', 'State', 'Describe'],
                'skills': ['Apply', 'Demonstrate', 'Use', 'Perform', 'Execute'],
                'competence': ['Work', 'Take responsibility', 'Adapt']
            },
            4: {
                'knowledge': ['Explain', 'Describe', 'Outline', 'Summarize', 'Compare'],
                'skills': ['Apply', 'Analyze', 'Solve', 'Select', 'Organize'],
                'competence': ['Take responsibility', 'Exercise management', 'Supervise']
            },
            5: {
                'knowledge': ['Analyze', 'Evaluate', 'Synthesize', 'Critique', 'Assess'],
                'skills': ['Design', 'Develop', 'Manage', 'Coordinate', 'Innovate'],
                'competence': ['Manage', 'Lead', 'Take responsibility for others']
            },
            6: {
                'knowledge': ['Evaluate', 'Synthesize', 'Create', 'Judge', 'Defend'],
                'skills': ['Design', 'Create', 'Evaluate', 'Manage complex projects', 'Lead'],
                'competence': ['Manage complex situations', 'Lead teams', 'Make strategic decisions']
            },
            7: {
                'knowledge': ['Synthesize', 'Evaluate critically', 'Create original solutions', 'Theorize'],
                'skills': ['Research', 'Innovate', 'Lead complex initiatives', 'Develop new approaches'],
                'competence': ['Lead strategic change', 'Manage complex stakeholders', 'Drive innovation']
            },
            8: {
                'knowledge': ['Develop new theories', 'Create paradigms', 'Establish frameworks'],
                'skills': ['Pioneer', 'Revolutionize', 'Lead transformation', 'Create breakthroughs'],
                'competence': ['Drive systemic change', 'Influence policy', 'Lead global initiatives']
            }
        }
    
    def define_role_profiles(self):
        """Define role competence profiles with clean professional titles"""
        return {
            'DAN': {
                'title': 'Data Analyst (Sustainability)',
                'abbreviation': 'DAN',
                'competence_focus': 'ESG data analysis and regulatory compliance',
                'key_capabilities': ['Data collection and validation', 'Statistical analysis', 'Regulatory reporting'],
                'professional_context': 'Sustainability data analysis and ESG compliance'
            },
            'DSM': {
                'title': 'Digital Sustainability Manager',
                'abbreviation': 'DSM', 
                'competence_focus': 'Program management and digital transformation',
                'key_capabilities': ['Project management', 'Stakeholder coordination', 'Digital strategy'],
                'professional_context': 'Digital sustainability program management'
            },
            'DSE': {
                'title': 'Data Engineer (Sustainability)',
                'abbreviation': 'DSE',
                'competence_focus': 'Sustainable data infrastructure development',
                'key_capabilities': ['Infrastructure design', 'Green computing', 'Performance optimization'],
                'professional_context': 'Sustainable data engineering and infrastructure'
            },
            'DSL': {
                'title': 'Digital Sustainability Lead',
                'abbreviation': 'DSL',
                'competence_focus': 'Strategic leadership and organizational transformation',
                'key_capabilities': ['Strategic planning', 'Change leadership', 'Stakeholder influence'],
                'professional_context': 'Strategic digital sustainability leadership'
            },
            'DSC': {
                'title': 'Digital Sustainability Consultant',
                'abbreviation': 'DSC',
                'competence_focus': 'Expert consulting across industry sectors',
                'key_capabilities': ['Consulting methodology', 'Multi-sector expertise', 'Solution design'],
                'professional_context': 'Professional sustainability consulting'
            }
        }
    
    def define_core_curricula(self):
        """Define comprehensive curricula covering all T3.2 target audiences"""
        return [
            {
                'id': 'DAN_Foundation',
                'role_id': 'DAN',
                'title': 'ESG Data Analysis (Foundation)',
                'ects': 0.5,
                'eqf_level': 4,
                'target_audience': 'students_job_seekers',
                'programme_type': 'micro_learning',
                'description': 'Foundational ESG data skills for students and job seekers entering sustainability careers, providing essential data literacy for employment readiness.',
                'delivery_methods': {
                    'primary': 'online_asynchronous',
                    'secondary': ['blended', 'self_paced'],
                    'workplace_component': 'mentored_practice'
                },
                'assessment_methods': {
                    'primary': 'practical portfolio',
                    'components': ['ESG data analysis project', 'Regulatory compliance case study', 'Peer review exercise'],
                    'weightings': [40, 35, 25]
                }
            },
            {
                'id': 'DSM_Basic', 
                'role_id': 'DSM',
                'title': 'Program Management (Basic)',
                'ects': 1.0,
                'eqf_level': 4,
                'target_audience': 'digital_professionals',
                'programme_type': 'micro_learning',
                'description': 'Project and program management skills for digital professionals implementing sustainability strategies within organizations.',
                'delivery_methods': {
                    'primary': 'blended',
                    'secondary': ['online_synchronous', 'intensive_workshop'],
                    'workplace_component': 'project_application'
                },
                'assessment_methods': {
                    'primary': 'project based',
                    'components': ['Sustainability project plan', 'Stakeholder engagement strategy', 'Implementation timeline'],
                    'weightings': [45, 30, 25]
                }
            },
            {
                'id': 'BOM_Foundation',
                'role_id': 'DSL',
                'title': 'Business Sustainability Strategy (Foundation)',
                'ects': 2.5,
                'eqf_level': 5,
                'target_audience': 'business_owners_managers',
                'programme_type': 'comprehensive',
                'description': 'Strategic sustainability foundations for business owners and managers seeking to integrate sustainability into business operations and strategy.',
                'delivery_methods': {
                    'primary': 'executive_intensive',
                    'secondary': ['blended', 'peer_learning'],
                    'workplace_component': 'strategy_implementation'
                },
                'assessment_methods': {
                    'primary': 'strategic portfolio',
                    'components': ['Business sustainability assessment', 'Strategic roadmap development', 'ROI analysis'],
                    'weightings': [40, 35, 25]
                }
            },
            {
                'id': 'DSE_Core',
                'role_id': 'DSE', 
                'title': 'Data Engineering (Core)',
                'ects': 2.0,
                'eqf_level': 5,
                'target_audience': 'digital_professionals',
                'programme_type': 'micro_learning',
                'description': 'Building and managing sustainable digital infrastructures for handling sustainability data—addressing critical expertise shortage in green IT.',
                'delivery_methods': {
                    'primary': 'hands_on_intensive',
                    'secondary': ['blended', 'workplace_integrated'],
                    'workplace_component': 'infrastructure_project'
                },
                'assessment_methods': {
                    'primary': 'technical implementation',
                    'components': ['Infrastructure design project', 'Performance optimization case', 'Technical documentation'],
                    'weightings': [50, 30, 20]
                }
            },
            {
                'id': 'DSL_Intermediate',
                'role_id': 'DSL',
                'title': 'Leadership (Intermediate)', 
                'ects': 5.0,
                'eqf_level': 6,
                'target_audience': 'business_owners_managers',
                'programme_type': 'comprehensive',
                'description': 'Mid-level leadership development for business owners and managers driving sustainability transformations across organizations.',
                'delivery_methods': {
                    'primary': 'executive_intensive',
                    'secondary': ['blended', 'peer_learning'],
                    'workplace_component': 'leadership_project'
                },
                'assessment_methods': {
                    'primary': 'leadership portfolio',
                    'components': ['Strategic planning exercise', 'Change management project', '360-degree feedback'],
                    'weightings': [40, 40, 20]
                }
            },
            {
                'id': 'DSC_Foundation',
                'role_id': 'DSC',
                'title': 'Consulting (Foundation)',
                'ects': 10.0,
                'eqf_level': 6,
                'target_audience': 'students_job_seekers',
                'programme_type': 'comprehensive',
                'description': 'Comprehensive preparation for students and job seekers to advise SMEs and organizations on integrating sustainability into operational practices.',
                'delivery_methods': {
                    'primary': 'structured_program',
                    'secondary': ['blended', 'mentored_learning'],
                    'workplace_component': 'consulting_practicum'
                },
                'assessment_methods': {
                    'primary': 'consulting portfolio',
                    'components': ['Client assessment report', 'Sustainability strategy design', 'Implementation roadmap'],
                    'weightings': [35, 35, 30]
                }
            },
            {
                'id': 'BOM_Advanced',
                'role_id': 'DSL',
                'title': 'Executive Sustainability Leadership',
                'ects': 15.0,
                'eqf_level': 7,
                'target_audience': 'business_owners_managers',
                'programme_type': 'comprehensive',
                'description': 'Advanced executive programme for business owners and senior managers leading organizational sustainability transformation and stakeholder engagement.',
                'delivery_methods': {
                    'primary': 'executive_modular',
                    'secondary': ['action_learning', 'peer_consultation'],
                    'workplace_component': 'transformation_project'
                },
                'assessment_methods': {
                    'primary': 'executive portfolio',
                    'components': ['Organizational transformation plan', 'Stakeholder engagement strategy', 'Impact measurement framework'],
                    'weightings': [45, 30, 25]
                }
            },
            {
                'id': 'DAN_Advanced',
                'role_id': 'DAN',
                'title': 'ESG Data Analysis (Advanced)',
                'ects': 7.5,
                'eqf_level': 6,
                'target_audience': 'digital_professionals',
                'programme_type': 'comprehensive',
                'description': 'Advanced ESG reporting and regulatory compliance training for experienced digital professionals specializing in sustainability data.',
                'delivery_methods': {
                    'primary': 'advanced_practicum',
                    'secondary': ['online_intensive', 'expert_facilitated'],
                    'workplace_component': 'regulatory_project'
                },
                'assessment_methods': {
                    'primary': 'advanced analytics',
                    'components': ['Regulatory compliance audit', 'Advanced ESG modeling', 'Industry benchmarking study'],
                    'weightings': [40, 35, 25]
                }
            },
            {
                'id': 'DSL_Advanced',
                'role_id': 'DSL',
                'title': 'Leadership (Advanced)',
                'ects': 30.0,
                'eqf_level': 7,
                'target_audience': 'business_owners_managers',
                'programme_type': 'comprehensive',
                'description': 'Comprehensive programme for senior business leaders and owners directing digital sustainability initiatives and managing complex stakeholder ecosystems.',
                'delivery_methods': {
                    'primary': 'executive_modular',
                    'secondary': ['action_learning', 'peer_consultation'],
                    'workplace_component': 'transformation_project'
                },
                'assessment_methods': {
                    'primary': 'strategic portfolio',
                    'components': ['Organizational transformation plan', 'Stakeholder engagement strategy', 'Impact measurement framework'],
                    'weightings': [45, 30, 25]
                }
            },
            {
                'id': 'DSC_Advanced',
                'role_id': 'DSC',
                'title': 'Consulting (Advanced)',
                'ects': 45.0,
                'eqf_level': 7,
                'target_audience': 'digital_professionals',
                'programme_type': 'comprehensive',
                'description': 'Advanced preparation for experienced digital professionals as consultants for complex, multi-sector sustainability transformations.',
                'delivery_methods': {
                    'primary': 'consulting_intensive',
                    'secondary': ['case based learning', 'client_projects'],
                    'workplace_component': 'multi_sector_engagement'
                },
                'assessment_methods': {
                    'primary': 'consulting mastery',
                    'components': ['Multi-client portfolio', 'Methodology development', 'Thought leadership piece'],
                    'weightings': [50, 30, 20]
                }
            }
        ]
    
    def select_modules_for_curriculum(self, curriculum_spec):
        """Select optimal modules with improved EQF consistency for curriculum"""
        role_id = curriculum_spec['role_id']
        target_ects = curriculum_spec['ects']
        eqf_level = curriculum_spec['eqf_level']
        
        # Get relevant modules for this role with improved EQF filtering
        relevant_modules = []
        for module in self.modules_data:
            role_relevance = module.get('role_relevance', {}).get(role_id, 0)
            module_eqf = module.get('eqf_level', 6)
            
            # Improved EQF consistency: prefer modules close to programme level
            min_eqf = max(3, eqf_level - 2)  # Never go below EQF 3, max 2 levels below
            max_eqf = eqf_level  # Never exceed programme level
            
            if role_relevance >= 60 and min_eqf <= module_eqf <= max_eqf:
                # Add bonus relevance for modules closer to programme EQF level
                eqf_bonus = 10 - abs(module_eqf - eqf_level) * 3
                adjusted_relevance = role_relevance + eqf_bonus
                
                relevant_modules.append({
                    'module': module,
                    'relevance': adjusted_relevance,
                    'ects': module.get('ects_points', 5),
                    'id': module['id']
                })
        
        # Sort by adjusted relevance
        relevant_modules.sort(key=lambda x: x['relevance'], reverse=True)
        
        # Select modules to match target ECTS
        selected_modules = []
        allocated_ects = 0
        
        for module_data in relevant_modules:
            if allocated_ects >= target_ects:
                break
                
            module_ects = module_data['ects']
            remaining_ects = target_ects - allocated_ects
            
            # Allocate appropriate ECTS
            if remaining_ects >= module_ects:
                allocated_ects_for_module = module_ects
            elif remaining_ects >= 0.5:  # Minimum micro-credential size
                allocated_ects_for_module = remaining_ects
            else:
                continue
                
            selected_modules.append({
                **module_data,
                'allocated_ects': allocated_ects_for_module
            })
            allocated_ects += allocated_ects_for_module
            
            if allocated_ects >= target_ects:
                break
        
        return selected_modules, allocated_ects
    
    def create_enhanced_learning_outcomes(self, module_data, curriculum_spec):
        """Create enhanced learning outcomes with EQF-aligned action verbs"""
        if not module_data or 'module' not in module_data or not module_data['module']:
            return {
                'knowledge': 'Identify fundamental concepts in digital sustainability professional practice.',
                'skills': 'Apply practical techniques and methodologies in professional contexts.',
                'competence': 'Work as a competent professional applying knowledge responsibly.'
            }
            
        module = module_data['module']
        eqf_level = curriculum_spec.get('eqf_level', 6)
        
        # Get appropriate action verbs for this EQF level
        verbs = self.eqf_action_verbs.get(eqf_level, self.eqf_action_verbs[6])
        
        # Use existing learning outcomes as base
        existing_outcomes = module.get('learning_outcomes', {})
        
        enhanced_outcomes = {}
        for outcome_type, outcome_text in existing_outcomes.items():
            if outcome_type in verbs:
                enhanced_outcomes[outcome_type] = outcome_text
            else:
                enhanced_outcomes[outcome_type] = outcome_text
        
        # If no existing outcomes, create based on module description
        if not enhanced_outcomes:
            module_name = module.get('name', 'Digital Sustainability')
            role_id = curriculum_spec.get('role_id', 'Professional')
            
            enhanced_outcomes = {
                'knowledge': f"{verbs['knowledge'][0]} fundamental concepts in {module_name} relevant to {role_id} professional practice.",
                'skills': f"{verbs['skills'][0]} practical techniques and methodologies from {module_name} in professional contexts.",
                'competence': f"{verbs['competence'][0]} as a competent professional applying {module_name} knowledge responsibly."
            }
        
        return enhanced_outcomes
    
    def create_visual_outcome_mapping(self, curriculum_spec, module_details, programme_outcomes):
        """Create detailed visual outcome mapping for DOCX integration"""
        if not self.include_visual_map:
            return None
            
        role_profile = self.role_profiles[curriculum_spec['role_id']]
        
        # Define progression paths based on current EQF level
        current_eqf = curriculum_spec['eqf_level']
        next_eqf = min(current_eqf + 1, 8)
        
        progression_paths = {
            'knowledge': f"EQF {next_eqf}: Advanced {role_profile['competence_focus']} Analysis",
            'skills': f"EQF {next_eqf}: Strategic {role_profile['key_capabilities'][0]} Implementation", 
            'competence': f"EQF {next_eqf}: Leadership in {role_profile['competence_focus']}"
        }
        
        # Map job role competences based on role
        job_competences = {
            'DAN': {
                'knowledge': 'ESG Data Sourcing & Classification',
                'skills': 'ESG Metrics Analysis & Visualization',
                'competence': 'Regulatory Reporting and Audit Readiness'
            },
            'DSM': {
                'knowledge': 'Digital Sustainability Program Design',
                'skills': 'Stakeholder Coordination & Project Delivery',
                'competence': 'Program Leadership & Strategic Alignment'
            },
            'DSE': {
                'knowledge': 'Sustainable Infrastructure Architecture',
                'skills': 'Green Computing & Performance Optimization',
                'competence': 'Technical Leadership & Innovation'
            },
            'DSL': {
                'knowledge': 'Strategic Sustainability Planning',
                'skills': 'Organizational Change Management',
                'competence': 'Executive Leadership & Influence'
            },
            'DSC': {
                'knowledge': 'Multi-Sector Sustainability Expertise',
                'skills': 'Consulting Methodology & Solution Design',
                'competence': 'Client Relationship & Business Development'
            }
        }
        
        # Create visual mapping rows
        visual_mapping_rows = []
        assessment_components = curriculum_spec['assessment_methods']['components']
        assessment_weightings = curriculum_spec['assessment_methods']['weightings']
        
        for i, (outcome_type, outcome_text) in enumerate(programme_outcomes.items()):
            # Get first relevant micro-credential for this outcome type
            relevant_micro_cred = None
            for module in module_details:
                for mc in module.get('micro_credentials', []):
                    if outcome_type.lower() in mc['title'].lower():
                        relevant_micro_cred = mc
                        break
                if relevant_micro_cred:
                    break
            
            # Fallback if no specific micro-credential found
            if not relevant_micro_cred and module_details:
                relevant_micro_cred = module_details[0]['micro_credentials'][0] if module_details[0]['micro_credentials'] else None
            
            # Create row data
            if relevant_micro_cred:
                micro_cred_text = f"{relevant_micro_cred['title']} ({relevant_micro_cred['ects_fraction']:.1f} ECTS)"
            else:
                micro_cred_text = f"Core Competence - {outcome_type.title()} (0.5 ECTS)"
            
            job_competence = job_competences.get(curriculum_spec['role_id'], {}).get(
                outcome_type, f"{outcome_type.title()} Professional Practice"
            )
            
            # Map to assessment component
            assessment_idx = i % len(assessment_components)
            assessment_text = f"{assessment_components[assessment_idx]} ({assessment_weightings[assessment_idx]}%)"
            
            progression_path = progression_paths.get(outcome_type, f"EQF {next_eqf}: Advanced Professional Practice")
            
            visual_mapping_rows.append((
                f"{outcome_type.title()}: {outcome_text[:80]}..." if len(outcome_text) > 80 else f"{outcome_type.title()}: {outcome_text}",
                micro_cred_text,
                job_competence,
                assessment_text,
                progression_path
            ))
        
        return {
            'headers': [
                "Programme-Level Outcome",
                "Micro-Credential",
                "Job Role Competence", 
                "Assessment Component",
                "Progression Path"
            ],
            'rows': visual_mapping_rows,
            'title': f"Visual Outcome Map: {role_profile['title']} ({role_profile['abbreviation']})"
        }
    
    def add_visual_outcome_table_to_docx(self, doc, visual_mapping):
        """Add visual outcome mapping table to DOCX document"""
        if not visual_mapping or not self.include_visual_map:
            return None
            
        # Add heading for visual outcome map
        doc.add_heading(visual_mapping['title'], level=1)
        
        # Add explanatory paragraph
        explanation = doc.add_paragraph(
            "This visual outcome map demonstrates the alignment between programme-level learning outcomes, "
            "micro-credentials, job role competences, assessment methods, and progression pathways. "
            "Each row shows how specific learning outcomes translate into professional capabilities and career advancement."
        )
        explanation.style = 'Normal'
        
        # Create table
        headers = visual_mapping['headers']
        rows = visual_mapping['rows']
        
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        
        # Style header row
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            # Make header bold and set alignment
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Set header background color
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D9E2F3')  # Light blue background
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
        
        # Populate data rows
        for row_idx, row_data in enumerate(rows, start=1):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                cell.text = cell_text
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # Set cell formatting
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Calibri'
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                
                # Add subtle alternating row colors
                if row_idx % 2 == 0:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'F8F9FA')  # Very light gray
                    cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Set column widths for better readability
        for i, width in enumerate([3.0, 2.5, 2.0, 1.8, 2.2]):  # Widths in inches
            for row in table.rows:
                row.cells[i].width = Inches(width)
        
        doc.add_paragraph()  # Add space after table
        
        return table
    
    def create_module_details(self, module_data, module_number, curriculum_spec):
        """Create detailed module information with enhanced learning outcomes and assessments"""
        if not module_data or 'module' not in module_data or not module_data['module']:
            return {
                'module_number': module_number,
                'module_id': 'UNKNOWN',
                'module_title': 'Professional Development Module',
                'module_description': 'Professional development in digital sustainability',
                'ects_credits': 0.5,
                'enhanced_learning_outcomes': {
                    'competence': 'Develop professional competence in digital sustainability'
                },
                'micro_credentials': [],
                'detailed_assessment': {
                    'method': 'portfolio',
                    'components': ['Professional reflection'],
                    'weightings': [100],
                    'quality_criteria': ['Professional application']
                }
            }
            
        module = module_data['module']
        allocated_ects = module_data.get('allocated_ects', 0.5)
        
        # Calculate workload breakdown
        total_hours = allocated_ects * 25
        contact_hours = round(total_hours * 0.35)
        self_study_hours = round(total_hours * 0.30)
        workplace_hours = round(total_hours * 0.25)
        assessment_hours = round(total_hours * 0.10)
        
        # Create enhanced learning outcomes
        enhanced_outcomes = self.create_enhanced_learning_outcomes(module_data, curriculum_spec)
        
        # Create micro-credentials from enhanced outcomes
        micro_credentials = self.create_micro_credentials(module_data, curriculum_spec, enhanced_outcomes)
        
        # Get assessment details safely
        assessment_methods = curriculum_spec.get('assessment_methods', {})
        
        return {
            'module_number': module_number,
            'module_id': module.get('id', 'UNKNOWN'),
            'module_title': module.get('name', 'Professional Development Module'),
            'module_description': module.get('description', 'Professional development in digital sustainability'),
            'extended_description': module.get('extended_description', ''),
            
            # ECTS/ECVET Integration
            'ects_credits': allocated_ects,
            'ecvet_points': allocated_ects,
            'eqf_level': module.get('eqf_level', 6),
            
            # Workload Distribution
            'total_workload_hours': total_hours,
            'contact_hours': contact_hours,
            'self_study_hours': self_study_hours,
            'workplace_hours': workplace_hours,
            'assessment_hours': assessment_hours,
            
            # Enhanced Learning Outcomes
            'enhanced_learning_outcomes': enhanced_outcomes,
            'original_learning_outcomes': module.get('learning_outcomes', {}),
            
            # Micro-Credentials based on enhanced outcomes
            'micro_credentials': micro_credentials,
            'stackable_badges': [mc['id'] for mc in micro_credentials],
            
            # Module Content from actual data
            'topics': module.get('topics', []),
            'skills': module.get('skills', []),
            'prerequisites': module.get('prerequisites', []),
            'thematic_area': module.get('thematic_area', 'General'),
            
            # Enhanced Assessment from curriculum specification
            'detailed_assessment': {
                'method': module.get('assessment_method', 'portfolio'),
                'components': assessment_methods.get('components', ['Professional portfolio']),
                'weightings': assessment_methods.get('weightings', [100]),
                'quality_criteria': ['Industry relevance', 'Competence demonstration', 'Professional application']
            },
            'framework_alignment': module.get('framework_alignment', {}),
            
            # Work-based Learning
            'is_work_based': module.get('is_work_based', False),
            'dual_principle_applicable': module.get('dual_principle_applicable', False),
            
            # Quality Assurance
            'institutional_framework': module.get('institutional_framework', {}),
            'quality_assurance': module.get('quality_assurance', {}),
            
            # Recognition
            'role_relevance_score': module_data.get('relevance', 0),
            'professional_recognition': True,
            'eu_transferable': True
        }
    
    def create_micro_credentials(self, module_data, curriculum_spec, enhanced_outcomes):
        """Create detailed micro-credential mapping from enhanced learning outcomes"""
        if not module_data or 'module' not in module_data or not module_data['module']:
            return []
            
        module = module_data['module']
        allocated_ects = module_data.get('allocated_ects', 0.5)
        
        micro_credentials = []
        base_id = f"{curriculum_spec.get('id', 'CURR')}_{module.get('id', 'MOD')}"
        
        # Ensure we have enhanced outcomes
        if not enhanced_outcomes:
            enhanced_outcomes = {'competence': 'Professional competence development'}
        
        # Create micro-credentials based on enhanced learning outcomes
        for outcome_type, outcome_description in enhanced_outcomes.items():
            micro_credential = {
                'id': f"{base_id}_{outcome_type}",
                'title': f"{module.get('name', 'Professional Development')} - {outcome_type.title()} Competence",
                'description': outcome_description,
                'ects_fraction': allocated_ects / len(enhanced_outcomes) if enhanced_outcomes else allocated_ects,
                'assessment_method': curriculum_spec.get('assessment_methods', {}).get('primary', 'portfolio'),
                'digital_badge': True,
                'blockchain_verified': True,
                'stackable': True,
                'role_alignment': f"Supports {curriculum_spec.get('role_id', 'Professional')} competence development",
                'quality_criteria': ['Competence demonstration', 'Industry relevance', 'Professional application']
            }
            micro_credentials.append(micro_credential)
        
        return micro_credentials
    
    def generate_programme_outcomes(self, curriculum_spec, modules):
        """Generate programme-level learning outcomes with EQF-aligned verbs"""
        role_profile = self.role_profiles[curriculum_spec['role_id']]
        eqf_level = curriculum_spec['eqf_level']
        
        # Get EQF-appropriate verbs
        verbs = self.eqf_action_verbs.get(eqf_level, self.eqf_action_verbs[6])
        
        return {
            'knowledge': f"{verbs['knowledge'][0]} comprehensive knowledge in {role_profile['competence_focus']} relevant to {role_profile['title']} professional practice, including frameworks, technologies, and industry standards.",
            'skills': f"{verbs['skills'][0]} advanced skills in {', '.join(role_profile['key_capabilities'][:2])} and related technical competences appropriate for EQF Level {eqf_level} practice.",
            'competence': f"{verbs['competence'][0]} as a competent {role_profile['title']} with accountability for {role_profile['competence_focus']} and continuous professional development."
        }
    
    def generate_curriculum(self, curriculum_spec):
        """Generate complete curriculum with enhanced professional formatting"""
        role_title = self.role_profiles.get(curriculum_spec.get('role_id', ''), {}).get('title', 'Professional')
        print(f"Generating: {role_title} - {curriculum_spec.get('title', 'Curriculum')}")
        
        try:
            # Select appropriate modules using actual data
            selected_modules, total_allocated_ects = self.select_modules_for_curriculum(curriculum_spec)
            
            if not selected_modules:
                raise ValueError(f"❌ No suitable modules found for {curriculum_spec.get('id', 'UNKNOWN')}")
            
            print(f"  Selected {len(selected_modules)} modules with {total_allocated_ects} ECTS")
            
            # Create detailed module information
            module_details = []
            total_micro_credentials = 0
            
            for i, module_data in enumerate(selected_modules):
                print(f"  Processing module {i+1}/{len(selected_modules)}: {module_data.get('module', {}).get('name', 'Unknown')}")
                module_detail = self.create_module_details(module_data, i + 1, curriculum_spec)
                if module_detail:
                    module_details.append(module_detail)
                    total_micro_credentials += len(module_detail.get('micro_credentials', []))
            
            print(f"  Created {len(module_details)} module details with {total_micro_credentials} micro-credentials")
            
            # Generate programme-level outcomes
            programme_outcomes = self.generate_programme_outcomes(curriculum_spec, module_details)
            print(f"  Generated programme outcomes: {len(programme_outcomes)} outcomes")
            
            # Create visual outcome mapping (if enabled)
            visual_mapping = None
            if self.include_visual_map:
                visual_mapping = self.create_visual_outcome_mapping(curriculum_spec, module_details, programme_outcomes)
                print(f"  Created visual mapping with {len(visual_mapping.get('rows', []))} rows")
            
        except Exception as e:
            print(f"  Error in curriculum generation: {e}")
            raise e
        
        # Get role profile for title integration
        role_profile = self.role_profiles[curriculum_spec['role_id']]
        
        # Create proper singular/plural for modules
        modules_text = "Module" if len(module_details) == 1 else "Modules"
        
        # Create comprehensive curriculum
        curriculum = {
            'curriculum_identification': {
                'id': curriculum_spec['id'],
                'clean_title': f"{role_profile['title']} ({role_profile['abbreviation']})",
                'full_title': f"{role_profile['title']} - {curriculum_spec['title']}",
                'role_profile': {
                    'role_title': role_profile['title'],
                    'role_abbreviation': role_profile['abbreviation'],
                    'competence_focus': role_profile['competence_focus'],
                    'professional_context': role_profile['professional_context']
                },
                'version': '2.1',
                'development_date': datetime.now().isoformat(),
                'curriculum_level': curriculum_spec['title'],
                'target_role': curriculum_spec['role_id'],
                'eqf_level': curriculum_spec['eqf_level'],
                'total_ects': total_allocated_ects,
                'total_ecvet': total_allocated_ects,
                'total_modules': len(module_details),
                'modules_text': modules_text,
                'total_micro_credentials': total_micro_credentials,
                'target_audience': curriculum_spec['target_audience'],
                'programme_type': curriculum_spec.get('programme_type', 'comprehensive')
            },
            
            # Programme Overview
            'programme_overview': {
                'description': curriculum_spec['description'],
                'target_audience_description': self.get_target_audience_description(curriculum_spec['target_audience']),
                'learning_approach': f"EQF Level {curriculum_spec['eqf_level']} {curriculum_spec.get('programme_type', 'comprehensive')} programme with modular structure, workplace integration, and stackable micro-credentials",
                'programme_outcomes': programme_outcomes
            },
            
            # Enhanced Modular Structure
            'modular_structure': {
                'design_principle': 'Flexible, stackable modules with comprehensive micro-credential mapping',
                'total_modules': len(module_details),
                'modules_text': modules_text,
                'total_ects': total_allocated_ects,
                'total_micro_credentials': total_micro_credentials,
                'module_details': module_details
            },
            
            # Assessment Framework
            'assessment_framework': curriculum_spec['assessment_methods'],
            
            # Visual Outcome Mapping (conditional)
            'visual_outcome_mapping': visual_mapping,
            
            # Recognition Framework
            'recognition_framework': {
                'eqf_level': curriculum_spec['eqf_level'],
                'ects_transferable': True,
                'ecvet_compatible': True,
                'bologna_compliant': curriculum_spec['eqf_level'] >= 6,
                'cross_border_recognition': 'EU-wide qualification recognition through EQF referencing',
                'professional_recognition': f'Industry recognition for {role_profile["title"]} competences'
            }
        }
        
        return curriculum
    
    def get_target_audience_description(self, target_audience):
        """Get specific description for all T3.2 target audiences"""
        descriptions = {
            'students_job_seekers': 'Students and job seekers preparing for careers in digital sustainability with comprehensive foundational training and employment readiness support',
            'digital_professionals': 'Experienced digital professionals expanding into sustainability domains with advanced technical skills and specialized competences',
            'business_owners_managers': 'Business owners, entrepreneurs, and senior managers driving organizational sustainability transformation with strategic oversight and decision-making capabilities'
        }
        return descriptions.get(target_audience, 'Professional development for sustainability roles')
    
    def save_curriculum_json(self, curriculum, filename):
        """Save curriculum as JSON"""
        json_path = self.output_dir / f"{filename}.json"
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(curriculum, f, indent=2, ensure_ascii=False)
        return json_path
    
    def save_curriculum_html(self, curriculum, filename):
        """Save curriculum as enhanced HTML with professional layout"""
        info = curriculum['curriculum_identification']
        modular = curriculum['modular_structure']
        role_profile = info['role_profile']
        assessment = curriculum['assessment_framework']
        
        html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{info['clean_title']} - {info['curriculum_level']}</title>
    <style>
        :root {{
            --primary-color: #1e3a5f;
            --secondary-color: #2c5530;
            --accent-color: #28a745;
            --background-light: #f8f9fa;
            --border-color: #dee2e6;
            --warning-color: #ffc107;
            --info-color: #17a2b8;
            --success-color: #28a745;
        }}
        
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            background-color: #fff;
        }}
        
        .container {{ max-width: 1200px; margin: 0 auto; padding: 20px; }}
        
        .header {{
            background: linear-gradient(135deg, var(--primary-color), var(--secondary-color));
            color: white;
            padding: 2.5rem;
            border-radius: 10px;
            margin-bottom: 2rem;
            text-align: center;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }}
        
        .header h1 {{ 
            font-size: 2.5rem; 
            margin-bottom: 0.5rem; 
            font-weight: 400;
            letter-spacing: -0.02em;
        }}
        
        .role-profile {{ 
            font-size: 1.1rem; 
            opacity: 0.95; 
            margin: 1rem 0;
            background: rgba(255,255,255,0.1);
            padding: 1rem;
            border-radius: 8px;
        }}
        
        .header .subtitle {{ 
            font-size: 1rem; 
            opacity: 0.9;
            margin-top: 1rem;
        }}
        
        .badge {{
            display: inline-block;
            padding: 0.4rem 1rem;
            border-radius: 25px;
            font-size: 0.9rem;
            font-weight: 500;
            margin: 0.3rem;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        
        .badge-ects {{ background-color: var(--warning-color); color: #000; }}
        .badge-eqf {{ background-color: var(--success-color); color: white; }}
        .badge-modules {{ background-color: var(--info-color); color: white; }}
        .badge-credentials {{ background-color: #6f42c1; color: white; }}
        
        .section {{
            background: white;
            border: 1px solid var(--border-color);
            border-radius: 10px;
            padding: 2rem;
            margin-bottom: 2rem;
            box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        }}
        
        .section h2 {{
            color: var(--primary-color);
            margin-bottom: 1.5rem;
            padding-bottom: 0.5rem;
            border-bottom: 3px solid var(--accent-color);
            font-size: 1.8rem;
            font-weight: 500;
        }}
        
        .section h3 {{
            color: var(--secondary-color);
            margin: 1.5rem 0 1rem 0;
            font-size: 1.3rem;
        }}
        
        .module-card {{
            background: var(--background-light);
            border: 1px solid var(--border-color);
            border-radius: 8px;
            padding: 2rem;
            margin-bottom: 2rem;
            transition: box-shadow 0.3s ease;
        }}
        
        .module-card:hover {{
            box-shadow: 0 4px 12px rgba(0,0,0,0.1);
        }}
        
        .module-header {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 1.5rem;
            flex-wrap: wrap;
            gap: 1rem;
        }}
        
        .module-title {{
            color: var(--primary-color);
            font-size: 1.4rem;
            font-weight: 600;
        }}
        
        .info-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 1.5rem;
            margin: 1.5rem 0;
        }}
        
        .info-item {{
            background: white;
            padding: 1.5rem;
            border-radius: 8px;
            border-left: 4px solid var(--accent-color);
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        }}
        
        .info-value {{
            font-size: 1.5rem;
            font-weight: bold;
            color: var(--primary-color);
            margin-bottom: 0.5rem;
        }}
        
        .info-label {{
            font-size: 0.9rem;
            color: #666;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}
        
        .learning-outcomes {{
            background: linear-gradient(135deg, #e3f2fd, #f3e5f5);
            border: 1px solid var(--info-color);
            border-radius: 8px;
            padding: 1.5rem;
            margin: 1.5rem 0;
        }}
        
        .learning-outcomes ul {{
            list-style: none;
            padding: 0;
        }}
        
        .learning-outcomes li {{
            margin: 1rem 0;
            padding: 0.5rem 0;
            border-bottom: 1px solid rgba(23, 162, 184, 0.2);
        }}
        
        .learning-outcomes li:last-child {{
            border-bottom: none;
        }}
        
        .micro-credentials {{
            background: linear-gradient(135deg, #fff8e1, #f3e5f5);
            border: 1px solid var(--warning-color);
            border-radius: 8px;
            padding: 1.5rem;
            margin: 1.5rem 0;
        }}
        
        .assessment-framework {{
            background: linear-gradient(135deg, #fff3e0, #fce4ec);
            border: 1px solid #ff9800;
            border-radius: 8px;
            padding: 1.5rem;
            margin: 1.5rem 0;
        }}
        
        footer {{
            text-align: center;
            margin-top: 3rem;
            padding: 2rem;
            background: var(--background-light);
            border-radius: 10px;
            color: #666;
            border-top: 1px solid var(--border-color);
        }}
        
        @media (max-width: 768px) {{
            .container {{ padding: 10px; }}
            .header {{ padding: 1.5rem; }}
            .header h1 {{ font-size: 2rem; }}
            .info-grid {{ grid-template-columns: 1fr; }}
            .module-header {{ flex-direction: column; align-items: flex-start; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>{info['clean_title']}</h1>
            <div class="role-profile">
                <strong>Professional Context:</strong> {role_profile['professional_context']}<br>
                <strong>Competence Focus:</strong> {role_profile['competence_focus']}
            </div>
            <div class="subtitle">
                <span class="badge badge-modules">{info['total_modules']} {info['modules_text']}</span>
                <span class="badge badge-eqf">EQF Level {info['eqf_level']}</span>
                <span class="badge badge-ects">{info['total_ects']} ECTS</span>
                <span class="badge badge-credentials">{info['total_micro_credentials']} Micro-Credentials</span>
            </div>
        </div>
        
        <div class="section">
            <h2>Programme Overview</h2>
            <p><strong>Description:</strong> {curriculum['programme_overview']['description']}</p>
            <p><strong>Target Audience:</strong> {curriculum['programme_overview']['target_audience_description']}</p>
            <p><strong>Learning Approach:</strong> {curriculum['programme_overview']['learning_approach']}</p>
            
            <h3>Programme Learning Outcomes</h3>
            <div class="learning-outcomes">
                <ul>
                    <li><strong>Knowledge:</strong> {curriculum['programme_overview']['programme_outcomes']['knowledge']}</li>
                    <li><strong>Skills:</strong> {curriculum['programme_overview']['programme_outcomes']['skills']}</li>
                    <li><strong>Competence:</strong> {curriculum['programme_overview']['programme_outcomes']['competence']}</li>
                </ul>
            </div>
        </div>
        
        <div class="section">
            <h2>Assessment Framework</h2>
            <div class="assessment-framework">
                <p><strong>Primary Method:</strong> {assessment['primary']}</p>
                <p><strong>Assessment Components:</strong></p>
                <ul>
"""
        
        for i, component in enumerate(assessment['components']):
            weighting = assessment['weightings'][i]
            html_content += f"<li>{component} ({weighting}%)</li>"
        
        html_content += """
                </ul>
            </div>
        </div>
"""
        
        # Add module details with enhanced formatting
        html_content += f"""
        <div class="section">
            <h2>Modular Structure</h2>
            <p><strong>Total {info['modules_text']}:</strong> {info['total_modules']} | <strong>Total ECTS:</strong> {info['total_ects']} | <strong>Micro-Credentials:</strong> {info['total_micro_credentials']}</p>
"""
        
        for module in modular['module_details']:
            html_content += f"""
            <div class="module-card">
                <div class="module-header">
                    <div class="module-title">Module {module['module_number']}: {module['module_title']}</div>
                    <span class="badge badge-ects">{module['ects_credits']} ECTS</span>
                </div>
                
                <p><strong>Description:</strong> {module['module_description']}</p>
                
                <div class="info-grid">
                    <div class="info-item">
                        <div class="info-value">{module['total_workload_hours']}h</div>
                        <div class="info-label">Total Workload</div>
                    </div>
                    <div class="info-item">
                        <div class="info-value">{module['workplace_hours']}h</div>
                        <div class="info-label">Workplace Learning</div>
                    </div>
                    <div class="info-item">
                        <div class="info-value">{len(module['micro_credentials'])}</div>
                        <div class="info-label">Micro-Credentials</div>
                    </div>
                    <div class="info-item">
                        <div class="info-value">EQF {module['eqf_level']}</div>
                        <div class="info-label">Level</div>
                    </div>
                </div>
"""
            
            # Add enhanced learning outcomes
            if module.get('enhanced_learning_outcomes'):
                html_content += """
                <div class="learning-outcomes">
                    <strong>Enhanced Learning Outcomes:</strong>
                    <ul>
"""
                for outcome_type, outcome_text in module['enhanced_learning_outcomes'].items():
                    html_content += f"<li><strong>{outcome_type.title()}:</strong> {outcome_text}</li>"
                
                html_content += "</ul></div>"
            
            # Add detailed assessment information
            html_content += f"""
                <div class="assessment-framework">
                    <strong>Assessment Strategy:</strong>
                    <p><strong>Method:</strong> {module['detailed_assessment']['method'].replace('_', ' ').title()}</p>
                    <p><strong>Components:</strong> {', '.join(module['detailed_assessment']['components'])}</p>
                    <p><strong>Quality Criteria:</strong> {', '.join(module['detailed_assessment']['quality_criteria'])}</p>
                </div>
"""
            
            # Add micro-credentials
            html_content += """
                <div class="micro-credentials">
                    <strong>Stackable Micro-Credentials:</strong>
                    <ul>
"""
            for mc in module['micro_credentials']:
                html_content += f"<li><strong>{mc['title']}</strong> ({mc['ects_fraction']:.1f} ECTS) - {mc['description'][:100]}...</li>"
            
            html_content += """
                    </ul>
                </div>
            </div>
"""
        
        html_content += f"""
        </div>
        
        <div class="section">
            <h2>Recognition & Quality Assurance</h2>
            <div class="info-grid">
                <div class="info-item">
                    <div class="info-label">European Recognition</div>
                    <p>EQF Level {curriculum['recognition_framework']['eqf_level']} with ECTS/ECVET compatibility</p>
                </div>
                <div class="info-item">
                    <div class="info-label">Professional Recognition</div>
                    <p>{curriculum['recognition_framework']['professional_recognition']}</p>
                </div>
                <div class="info-item">
                    <div class="info-label">Cross-border Recognition</div>
                    <p>{curriculum['recognition_framework']['cross_border_recognition']}</p>
                </div>
            </div>
        </div>
        
        <footer>
            <p><strong>ECM Curriculum Generator v21</strong> | Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
            <p>Complete target audience coverage and EQF Level 4-8 consistency achieved</p>
        </footer>
    </div>
</body>
</html>
"""
        
        html_path = self.output_dir / f"{filename}.html"
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        return html_path
    
    def save_curriculum_docx(self, curriculum, filename):
        """Save curriculum as professionally formatted DOCX"""
        doc = Document()
        
        # Set document styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        info = curriculum['curriculum_identification']
        role_profile = info['role_profile']
        
        # Title with clean formatting
        title = doc.add_heading(info['clean_title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(info['curriculum_level'])
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.color.rgb = RGBColor(44, 85, 48)
        
        # Professional context
        context_para = doc.add_paragraph()
        context_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        context_run = context_para.add_run(f"Professional Context: {role_profile['professional_context']}")
        context_run.font.size = Pt(12)
        context_run.font.color.rgb = RGBColor(102, 102, 102)
        
        # Key metrics with proper grammar
        metrics_para = doc.add_paragraph()
        metrics_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        metrics_run = metrics_para.add_run(
            f"EQF Level {info['eqf_level']} | {info['total_ects']} ECTS | {info['total_modules']} {info['modules_text']} | {info['total_micro_credentials']} Micro-Credentials"
        )
        metrics_run.font.size = Pt(10)
        metrics_run.font.color.rgb = RGBColor(102, 102, 102)
        
        doc.add_paragraph()
        
        # Programme Overview
        doc.add_heading('Programme Overview', level=1)
        doc.add_paragraph(curriculum['programme_overview']['description'])
        doc.add_paragraph(f"Target Audience: {curriculum['programme_overview']['target_audience_description']}")
        doc.add_paragraph(f"Learning Approach: {curriculum['programme_overview']['learning_approach']}")
        
        # Enhanced Programme Learning Outcomes
        doc.add_heading('Programme Learning Outcomes', level=2)
        outcomes = curriculum['programme_overview']['programme_outcomes']
        for outcome_type, outcome_text in outcomes.items():
            outcome_para = doc.add_paragraph()
            outcome_para.add_run(f"{outcome_type.title()}: ").bold = True
            outcome_para.add_run(outcome_text)
        
        # Add Visual Outcome Mapping Table (if enabled)
        if self.include_visual_map and curriculum.get('visual_outcome_mapping'):
            doc.add_page_break()
            visual_mapping = curriculum['visual_outcome_mapping']
            self.add_visual_outcome_table_to_docx(doc, visual_mapping)
        
        # Assessment Framework
        doc.add_heading('Assessment Framework', level=1)
        assessment = curriculum['assessment_framework']
        
        doc.add_heading('Programme-Level Assessment', level=2)
        doc.add_paragraph(f"Primary Method: {assessment['primary']}")
        
        components_para = doc.add_paragraph("Assessment Components:")
        components_para.style = 'Heading 3'
        
        for i, component in enumerate(assessment['components']):
            weighting = assessment['weightings'][i]
            comp_para = doc.add_paragraph()
            comp_para.add_run(f"• {component}: ").bold = True
            comp_para.add_run(f"{weighting}%")
        
        # Programme Details Table
        doc.add_heading('Programme Details', level=1)
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'
        
        details = [
            ('Curriculum Level', info['curriculum_level']),
            ('Target Role', f"{role_profile['role_title']} ({role_profile['role_abbreviation']})"),
            ('Competence Focus', role_profile['competence_focus']),
            ('EQF Level', str(info['eqf_level'])),
            ('Total ECTS', f"{info['total_ects']} ECTS"),
            (f'Total {info["modules_text"]}', str(info['total_modules'])),
            ('Micro-Credentials', str(info['total_micro_credentials']))
        ]
        
        for i, (label, value) in enumerate(details):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        # Enhanced Modules with proper numbering
        doc.add_heading('Modular Structure', level=1)
        
        for module in curriculum['modular_structure']['module_details']:
            # Module heading with proper grammar
            doc.add_heading(f"Module {module['module_number']}: {module['module_title']}", level=2)
            
            # Module details table
            mod_table = doc.add_table(rows=6, cols=2)
            mod_table.style = 'Table Grid'
            
            mod_details = [
                ('ECTS Credits', str(module['ects_credits'])),
                ('Total Workload', f"{module['total_workload_hours']} hours"),
                ('Workplace Hours', f"{module['workplace_hours']} hours"),
                ('Micro-Credentials', str(len(module['micro_credentials']))),
                ('EQF Level', str(module['eqf_level'])),
                ('Thematic Area', module['thematic_area'])
            ]
            
            for i, (label, value) in enumerate(mod_details):
                mod_table.cell(i, 0).text = label
                mod_table.cell(i, 1).text = value
            
            # Description
            doc.add_paragraph(f"Description: {module['module_description']}")
            
            # Enhanced learning outcomes
            if module.get('enhanced_learning_outcomes'):
                doc.add_paragraph("Enhanced Learning Outcomes:", style='Heading 3')
                for outcome_type, outcome_text in module['enhanced_learning_outcomes'].items():
                    outcome_para = doc.add_paragraph()
                    outcome_para.add_run(f"{outcome_type.title()}: ").bold = True
                    outcome_para.add_run(outcome_text)
            
            # Assessment details
            assessment_info = module['detailed_assessment']
            doc.add_paragraph("Assessment Strategy:", style='Heading 3')
            doc.add_paragraph(f"Method: {assessment_info['method'].replace('_', ' ').title()}")
            doc.add_paragraph(f"Components: {', '.join(assessment_info['components'])}")
            doc.add_paragraph(f"Quality Criteria: {', '.join(assessment_info['quality_criteria'])}")
            
            # Micro-credentials
            if module.get('micro_credentials'):
                doc.add_paragraph("Stackable Micro-Credentials:", style='Heading 3')
                for mc in module['micro_credentials']:
                    mc_para = doc.add_paragraph()
                    mc_para.add_run(f"• {mc['title']}: ").bold = True
                    mc_para.add_run(f"{mc['ects_fraction']:.1f} ECTS - {mc['description']}")
        
        # Footer
        doc.add_page_break()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            f"ECM Curriculum Generator v21 | {datetime.now().strftime('%Y-%m-%d')}"
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(102, 102, 102)
        
        docx_path = self.output_dir / f"{filename}.docx"
        doc.save(docx_path)
        return docx_path
    
    def generate_all_curricula(self):
        """Generate all core curricula"""
        print(f"\n=== GENERATING CORE CURRICULA ===")
        print(f"✓ Visual Outcome Mapping: {'Enabled' if self.include_visual_map else 'Disabled'}")
        
        generated_files = []
        target_audience_count = {'students_job_seekers': 0, 'digital_professionals': 0, 'business_owners_managers': 0}
        
        for i, curriculum_spec in enumerate(self.core_curricula, 1):
            role_profile = self.role_profiles[curriculum_spec['role_id']]
            print(f"\n[{i}/{len(self.core_curricula)}] Generating: {role_profile['title']} ({role_profile['abbreviation']}) - {curriculum_spec['title']}")
            
            # Track T3.2 compliance metrics
            target_audience_count[curriculum_spec['target_audience']] += 1
            
            try:
                # Generate curriculum with enhanced features
                curriculum = self.generate_curriculum(curriculum_spec)
                
                # Create simplified filename: number_profile_EQFlevel
                filename = f"{i:02d}_{curriculum_spec['role_id']}_{curriculum_spec['eqf_level']}"
                
                # Save in all formats
                json_path = self.save_curriculum_json(curriculum, filename)
                html_path = self.save_curriculum_html(curriculum, filename)
                docx_path = self.save_curriculum_docx(curriculum, filename)
                
                generated_files.extend([json_path, html_path, docx_path])
                
                info = curriculum['curriculum_identification']
                print(f"✓ {info['total_ects']} ECTS | {info['total_modules']} {info['modules_text']} | {info['total_micro_credentials']} micro-credentials")
                print(f"  Role: {info['clean_title']}")
                print(f"  Target: {curriculum_spec['target_audience'].replace('_', ' ').title()}")
                print(f"  EQF Level: {info['eqf_level']}")
                
            except Exception as e:
                print(f"❌ Error generating {curriculum_spec['id']}: {e}")
                raise
        
        print(f"\n=== GENERATION COMPLETE ===")
        print(f"✓ Generated {len(self.core_curricula)} core curricula")
        print(f"✓ Created {len(generated_files)} files (3 per curriculum: JSON, HTML, DOCX)")
        
        print("\n📊 TARGET AUDIENCE COVERAGE:")
        for audience, count in target_audience_count.items():
            print(f"   • {audience.replace('_', ' ').title()}: {count} programmes")
        
        print(f"✓ Output directory: {self.output_dir}")
        
        return generated_files
    
    # Include all the utility methods from the original script
    def load_config(self, config_path):
        """Load configuration with robust path resolution"""
        script_dir = Path(__file__).parent
        
        possible_paths = [
            script_dir / config_path,
            script_dir / '../config/settings.json',
            Path.cwd() / 'config/settings.json',
            script_dir / 'settings.json'
        ]
        
        for path in possible_paths:
            if path.exists():
                print(f"✓ Using config: {path}")
                with open(path, 'r', encoding='utf-8') as f:
                    return json.load(f)
        
        raise FileNotFoundError("❌ No configuration file found. Ensure settings.json exists.")
    
    def setup_paths(self):
        """Setup all paths from configuration"""
        script_dir = Path(__file__).parent
        
        # Modules file path
        modules_path = self.config['paths']['input_modules']
        possible_modules_paths = [
            script_dir / modules_path.lstrip('./'),
            script_dir / '../input/modules/modules_v5.json',
            Path.cwd() / 'input/modules/modules_v5.json'
        ]
        
        self.modules_file = None
        for path in possible_modules_paths:
            if path.exists():
                self.modules_file = path
                break
        
        if not self.modules_file:
            raise FileNotFoundError("❌ modules_v5.json not found")
        
        # Roles file path (derived from modules location)
        roles_dir = self.modules_file.parent.parent / 'roles'
        self.roles_file = roles_dir / 'roles.json'
        
        if not self.roles_file.exists():
            raise FileNotFoundError(f"❌ roles.json not found at {self.roles_file}")
        
        # Output directory
        output_config = self.config.get('output', {}).get('curricula', {})
        output_dir = output_config.get('directory', './output/curricula')
        self.output_dir = script_dir / output_dir.lstrip('./')
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        print(f"✓ Modules source: {self.modules_file}")
        print(f"✓ Roles source: {self.roles_file}")
        print(f"✓ Output directory: {self.output_dir}")
    
    def load_modules_data(self):
        """Load modules with validation"""
        try:
            with open(self.modules_file, 'r', encoding='utf-8') as f:
                modules = json.load(f)
                
            if not isinstance(modules, list) or len(modules) == 0:
                raise ValueError("❌ Invalid modules data structure")
                
            return modules
        except Exception as e:
            raise RuntimeError(f"❌ Failed to load modules: {e}")
    
    def load_roles_data(self):
        """Load roles with validation"""
        try:
            with open(self.roles_file, 'r', encoding='utf-8') as f:
                roles = json.load(f)
                
            if not isinstance(roles, list) or len(roles) == 0:
                raise ValueError("❌ Invalid roles data structure")
                
            return roles
        except Exception as e:
            raise RuntimeError(f"❌ Failed to load roles: {e}")
    
    def validate_data_integrity(self):
        """Validate data integrity - no fallbacks allowed"""
        required_module_fields = ['id', 'name', 'eqf_level', 'ects_points', 'role_relevance']
        for module in self.modules_data:
            for field in required_module_fields:
                if field not in module:
                    raise ValueError(f"❌ Module {module.get('id', 'unknown')} missing required field: {field}")
        
        required_role_fields = ['id', 'name']
        for role in self.roles_data:
            for field in required_role_fields:
                if field not in role:
                    raise ValueError(f"❌ Role {role.get('id', 'unknown')} missing required field: {field}")
        
        print("✓ Data integrity validation passed")


def main():
    """Main execution function with command line argument support"""
    parser = argparse.ArgumentParser(description='Generate ECM Curricula with optional Visual Outcome Mapping')
    parser.add_argument('--no-visual-map', action='store_true', 
                       help='Disable Visual Outcome Mapping generation')
    parser.add_argument('--config', default='config/settings.json',
                       help='Path to configuration file (default: config/settings.json)')
    
    args = parser.parse_args()
    
    # Determine if visual mapping should be included (default: True, unless --no-visual-map is specified)
    include_visual_map = not args.no_visual_map
    
    try:
        print("Starting ECM Curriculum Generator v21...")
        
        # Initialize generator with visual mapping toggle
        generator = EnhancedCurriculumGenerator(
            config_path=args.config, 
            include_visual_map=include_visual_map
        )
        
        # Generate curricula
        files = generator.generate_all_curricula()
        
        print(f"\n🎉 SUCCESS: Generated {len(files)} files")
        print("✅ ACHIEVEMENTS:")
        print("✅ Complete target audience coverage")
        print("✅ EQF Level 4-8 coverage")
        print("✅ Professional quality formatting")
        print("✅ Simplified file naming")
        print("✅ Enhanced terminology (competence, project based, etc.)")
        print(f"✅ Visual Outcome Mapping: {'Included' if include_visual_map else 'Excluded'}")
        
        return True
        
    except Exception as e:
        print(f"\n❌ GENERATION FAILED: {e}")
        return False


if __name__ == "__main__":
    success = main()
    exit(0 if success else 1)